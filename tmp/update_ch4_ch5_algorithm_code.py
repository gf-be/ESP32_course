# -*- coding: utf-8 -*-
from pathlib import Path
import shutil

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


ROOT = Path(r"F:\mechineSight\stm32\罗丹")
SRC = ROOT / "初稿.docx"
OUT = ROOT / "初稿_第四五章算法代码补充.docx"


def ensure_style(doc):
    styles = doc.styles
    if "CodeBlock" not in [s.name for s in styles]:
        style = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Consolas"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
        style.font.size = Pt(8.5)
        style.paragraph_format.left_indent = Pt(18)
        style.paragraph_format.space_before = Pt(2)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.0


def set_text(paragraph, text):
    paragraph.text = text


def insert_paragraph_after(paragraph, text="", style=None, bold=False):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        run = new_para.add_run(text)
        run.bold = bold
    return new_para


def insert_code_after(paragraph, code):
    p = insert_paragraph_after(paragraph, style="CodeBlock")
    lines = code.strip("\n").splitlines()
    run = p.add_run(lines[0] if lines else "")
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    run.font.size = Pt(8.5)
    for line in lines[1:]:
        run.add_break()
        run.add_text(line)
    return p


def first_para(doc, startswith):
    for p in doc.paragraphs:
        if p.text.strip().startswith(startswith):
            return p
    raise RuntimeError("Paragraph not found: %s" % startswith)


def replace_first(doc, startswith, text):
    p = first_para(doc, startswith)
    set_text(p, text)
    return p


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def delete_paragraphs_starting(doc, prefixes):
    for p in list(doc.paragraphs):
        text = p.text.strip()
        if any(text.startswith(prefix) for prefix in prefixes):
            delete_paragraph(p)


def insert_sequence(anchor, items):
    cur = anchor
    for kind, text in items:
        if kind == "h":
            cur = insert_paragraph_after(cur, text, bold=True)
        elif kind == "p":
            cur = insert_paragraph_after(cur, text)
        elif kind == "code":
            cur = insert_code_after(cur, text)
    return cur


def main():
    shutil.copy2(SRC, OUT)
    doc = Document(OUT)
    ensure_style(doc)

    # Chapter 4: fix misplaced/old data and add implementation snippets.
    replace_first(
        doc,
        "为进一步分析陀螺仪静态噪声",
        "本节采用最新六位置数据建立 12 参数仿射模型。六个姿态分别对应 +X、-X、+Y、-Y、+Z、-Z 朝上的静止放置，每个姿态取均值后与理想重力方向建立对应关系。模型采用 raw = M true + b 的形式，其中 b 表示三轴零偏，M 的对角项主要反映比例因子，非对角项反映轴间耦合、非正交和固定安装角。实时补偿时使用 true = inv(M)(raw - b)。",
    )
    set_text(
        first_para(doc, "静态统计结果表明，陀螺仪三轴均存在非零零偏"),
        "最新结果表明，未标定时六个姿态的平均均值向量误差为 12.935 mg，6 参数零偏/比例因子模型降至 3.691 mg，12 参数仿射模型进一步降至 0.755 mg；均值向量 RMS 误差由 14.296 mg 降至 0.973 mg，平均误差降低约 17.122 倍。这说明误差改善不仅来自零偏扣除，还来自比例因子和轴间耦合补偿。",
    )
    accel_meaning = first_para(doc, "Allan 方差分析显示，三轴 Allan 偏差随平均时间 tau 增大先下降")
    set_text(
        accel_meaning,
        "12 参数模型在工程上的意义是把固定安装误差纳入校正矩阵。若传感器只有零偏误差，M 或 A_inv 应接近对角矩阵；本实验矩阵存在非零非对角项，说明实际模块相对 PCB 坐标轴存在小角度不正交或安装偏差，因此仅做三轴零偏扣除不能达到最佳模长一致性。",
    )
    insert_sequence(
        accel_meaning,
        [
            ("h", "关键代码段 4-1：加速度计 12 参数仿射标定与实时补偿"),
            ("code", """
# 六位置均值 means 对应理想重力方向 targets，单位为 g
X = np.column_stack([targets, np.ones(len(targets))])
W, *_ = np.linalg.lstsq(X, means, rcond=None)
M = W[:3, :].T          # raw = M @ true + b
b = W[3, :]
A_inv = np.linalg.inv(M)

def calibrate_accel(a_raw):
    return A_inv @ (a_raw - b)
"""),
            ("p", "该代码把六位置真值方向和原始均值写成线性最小二乘问题，求得 M 与 b 后保存为 CSV。实时融合前先调用 calibrate_accel()，使加速度计输出进入统一的板体坐标系，从源头降低 roll/pitch 解算中的系统误差。"),
        ],
    )

    mag_anchor = first_para(doc, "将大量不同姿态下采集的磁力计数据代入椭球方程")
    insert_sequence(
        mag_anchor,
        [
            ("h", "关键代码段 4-2：磁力计椭球 SVD 拟合与鲁棒重拟合"),
            ("code", """
D = np.column_stack([x*x, y*y, z*z, 2*x*y, 2*x*z, 2*y*z,
                     2*x, 2*y, 2*z, np.ones_like(x)])
_, s, vt = np.linalg.svd(D, full_matrices=False)
q = vt[-1]                       # 最小奇异值方向，对应最小残差
center = -np.linalg.solve(A, b)  # hard-iron 偏置
corrected = (T @ (points - center).T).T

r = np.linalg.norm(corrected, axis=1)
gate = max(3.5 * mad(r), 0.08 * np.median(r))
inlier = np.abs(r - np.median(r)) <= gate
# 使用 inlier 点重新拟合，降低手机/金属物体瞬时靠近带来的离群影响
"""),
            ("p", "最新鲁棒分析在 3600 个磁力计旋转样本中识别出 47 个疑似离群点，离群比例为 1.31%。普通 SVD 校准后半径 CV 为 0.057233，MAD 重拟合后为 0.057303，说明本组数据离群影响较小；但该脚本给出了可复现实验流程，可用于答辩解释为什么标准最小二乘对异常点敏感。"),
        ],
    )

    set_text(
        first_para(doc, "陀螺仪静止采样包括 180 s"),
        "陀螺仪 Allan 方差实验采用最新文件 gyro_allan_20260616_205634.csv，采样率为 50.000 Hz，持续 1799.980 s，共 90000 组样本。实验温度从 42.759°C 到 43.527°C，变化 0.768°C，适合作为静止零偏和短时噪声分析来源，但不足以支撑完整温度补偿模型。",
    )
    set_text(
        first_para(doc, "室温 180 s 静态基线中"),
        "静止均值显示 X/Y/Z 三轴陀螺零偏分别为 0.228304 deg/s、0.964654 deg/s 和 -0.100939 deg/s，标准差分别为 0.052827 deg/s、0.071176 deg/s 和 0.060315 deg/s。后续姿态融合均先扣除该静态零偏，避免角速度积分产生持续漂移。",
    )
    set_text(
        first_para(doc, "进行了温度漂移观察"),
        "温度补偿在本文中作为算法方案说明：若需要建立 b_g(T)=alpha+beta T+gamma T^2，应在多个温区重复静止采样并做最小二乘拟合。当前 Allan 实验温度跨度较小，因此报告不把二次温度模型作为已完成定量结论，只将其作为后续改进边界。",
    )
    set_text(
        first_para(doc, "长时间静止采样文件为 allan_raw_1h.csv"),
        "长时间静止采样文件为 gyro_allan_20260616_205634.csv，平均采样率为 50.000 Hz，记录时长约 30.00 min。",
    )
    set_text(
        first_para(doc, "Allan 方差曲线已由 fig_allan_curve_gz.png"),
        "Allan 方差曲线由 analyze_gyro_allan.py 生成，可在终稿中插入 gyro_allan_deviation.png，并标注三轴 Allan 偏差最低点。",
    )
    allan_anchor = first_para(doc, "以 gz 轴为例")
    set_text(
        allan_anchor,
        "X/Y/Z 轴最小 Allan 偏差分别为 0.001481 deg/s、0.001963 deg/s 和 0.001939 deg/s，对应聚类时间分别为 157.740 s、106.160 s 和 36.940 s。短 tau 段主要反映角随机游走，可指导测量噪声 R；较长 tau 附近的低谷反映零偏不稳定性，可指导 ESKF 中陀螺零偏随机游走 Q 的设置。",
    )
    insert_sequence(
        allan_anchor,
        [
            ("h", "关键代码段 4-3：Allan 方差计算与噪声参数提取"),
            ("code", """
for m in cluster_sizes:
    n = (len(gyro) // m) * m
    y = gyro[:n].reshape(-1, m).mean(axis=1)
    adev = np.sqrt(0.5 * np.mean(np.diff(y) ** 2))
    tau = m / sample_rate
    allan_curve.append((tau, adev))
best_tau, bias_instability = min(allan_curve, key=lambda r: r[1])
"""),
            ("p", "该代码将长时间静止角速度按不同 tau 分组求均值，再计算相邻均值差。它不是单纯画图，而是把陀螺白噪声和零偏不稳定性转化为后续卡尔曼滤波 Q/R 参数的量化依据。"),
            ("p", "图题：静止状态下陀螺仪三轴输出与温度变化曲线。对应文件：gyro_allan_timeseries.png。"),
            ("p", "图题：静止状态下陀螺仪三轴零偏分布直方图。对应文件：gyro_allan_histogram.png。"),
            ("p", "图题：陀螺仪三轴 Allan 偏差曲线及最小 Allan 偏差标注。对应文件：gyro_allan_deviation.png。"),
            ("h", "4.6 BMP280 气压高度滤波与误差预算"),
            ("p", "BMP280 分析脚本读取压力与温度数据，以实验开始阶段的压力中位数作为 P0，将气压换算为相对高度，并使用一维高度/速度 Kalman 滤波进行平滑。最新数据 bmp280_height_change_20260617_084325.csv 共 900 行，持续 179.800 s，原始气压高度标准差为 0.179568 m，KF 输出高度标准差为 0.180744 m，高频残差标准差为 0.015165 m。该结果说明短时气压高度具有分米级相对变化分辨能力，但绝对高度主要受 P0 基准气压控制。"),
            ("code", """
h_baro = 44330.0 * (1.0 - (pressure_pa / p0_pa) ** 0.190294957)
x = np.array([h_baro[0], 0.0])       # state: height, vertical velocity
F = np.array([[1.0, dt], [0.0, 1.0]])
Q = sigma_a**2 * np.array([[dt**4/4, dt**3/2], [dt**3/2, dt**2]])
H = np.array([[1.0, 0.0]])
K = P @ H.T @ np.linalg.inv(H @ P @ H.T + R)
x = x + K @ (h_baro[k] - H @ x)
"""),
            ("p", "BMP280 的精度提升重点不在于把气压高度滤得更“平”，而是明确误差预算：短时噪声约 0.18 m，而 P0 误差 1 hPa 会造成约 8.4 m 高度偏差，因此报告中把 BMP280 定位为相对高度变化传感器，并建议用 GPS 或本地气象参考周期性更新 P0。"),
        ],
    )

    # Chapter 5: fix latest fusion metrics and add implementation snippets.
    set_text(
        first_para(doc, "姿态融合实验共采集 16500 组样本"),
        "姿态融合实验共采集 16500 组样本，整体有效采样率为 49.561 Hz。离线运行时，互补滤波、Mahony PI 和 Madgwick MARG 的单算法更新吞吐率分别约为 363093 Hz、193657 Hz 和 37846 Hz，均远高于课程 Spec 中姿态融合更新率不低于 100 Hz 的要求。需要注意，该频率为电脑端离线吞吐率；ESP32 实时演示中还应同时考虑 I2C 读取和串口输出开销。",
    )
    comp_anchor = first_para(doc, "theta = alpha")
    insert_sequence(
        comp_anchor,
        [
            ("h", "关键代码段 5-1：互补滤波实现"),
            ("code", """
roll_g  = roll  + gx * dt
pitch_g = pitch + gy * dt
roll_acc  = np.degrees(np.arctan2(ay, az))
pitch_acc = np.degrees(np.arctan2(-ax, np.sqrt(ay*ay + az*az)))
roll  = alpha * roll_g  + (1.0 - alpha) * roll_acc
pitch = alpha * pitch_g + (1.0 - alpha) * pitch_acc
"""),
            ("p", "互补滤波的实现体现了传感器互补思想：陀螺仪负责短时连续性，加速度计负责长期重力约束。磁力计经椭球标定和倾斜补偿后用于 yaw 低频校正，但磁环境异常时需要降低权重。"),
        ],
    )

    set_text(
        first_para(doc, "本项目最终至少实现 Madgwick 或 Mahony 中的一种"),
        "本项目已完成互补滤波、Mahony PI 与 Madgwick MARG 三种算法对比。对比指标包括静止状态姿态角标准差、动态倾斜响应、姿态更新频率和算法计算量。Mahony 的反馈项便于解释零偏抑制，Madgwick 的 beta 参数便于解释观测校正强度，互补滤波则作为低计算量基准算法。",
    )
    set_text(
        first_para(doc, "当前实测验证以互补滤波为主"),
        "因此，第五章不再把 Mahony/Madgwick 写作后续计划，而是作为已经完成的姿态融合结果进行分析；后续仅保留 ESKF 参数调节、动态真值对比和嵌入式资源占用统计作为改进方向。",
    )
    mad_anchor = first_para(doc, "互补滤波采用 alpha=0.98")
    insert_sequence(
        mad_anchor,
        [
            ("h", "关键代码段 5-2：Mahony PI 与 Madgwick MARG 的核心更新"),
            ("code", """
# Mahony: 观测方向与预测方向叉乘形成姿态误差
e = np.cross(g_pred, acc_norm) + np.cross(m_pred, mag_norm)
gyro_bias += ki * e * dt
omega = gyro - gyro_bias + kp * e
q = normalize(quat_mul(q, small_angle_quat(omega * dt)))

# Madgwick: 梯度下降校正四元数误差
q_dot = 0.5 * quat_mul(q, np.r_[0.0, gyro])
q_dot -= beta * gradient_objective(q, acc_norm, mag_norm)
q = normalize(q + q_dot * dt)
"""),
            ("p", "Mahony 和 Madgwick 都使用四元数避免欧拉角奇异，但误差构造方式不同。Mahony 把方向误差作为 PI 反馈修正陀螺积分，Madgwick 用目标函数梯度下降修正四元数；beta 或 Kp/Ki 的选择决定了“相信陀螺预测”与“相信加速度计/磁力计观测”的平衡。"),
        ],
    )
    set_text(
        first_para(doc, "由静态标准差可见，在水平静止状态下"),
        "由最新静态标准差可见，在水平静止状态下，互补滤波 roll/pitch/yaw 标准差分别为 0.0189°、0.0186° 和 0.0333°，Mahony PI 分别为 0.0299°、0.0280° 和 0.0518°，Madgwick MARG 分别为 0.0196°、0.0204° 和 0.0394°；在固定倾斜状态下，互补滤波三轴标准差分别为 0.0263°、0.0242° 和 0.0868°，Mahony PI 分别为 0.0385°、0.0350° 和 0.1263°，Madgwick MARG 分别为 0.0709°、0.0643° 和 0.1002°。三种算法在静态段均达到小于 0.2° 的短时抖动水平，其中互补滤波和 Madgwick 在本数据集上最平稳，Mahony 在实时 ESP32 演示中具有参数直观、零偏反馈清晰的优势。",
    )

    eskf_anchor = first_para(doc, "考虑到普通 NEO-6M/GPS 单点定位在静止时仍会出现米级漂移")
    insert_sequence(
        eskf_anchor,
        [
            ("h", "关键代码段 5-3：15 维 ESKF 预测、GPS 更新与静止约束"),
            ("code", """
# predict: IMU 高频预测
omega = np.radians(gyro_dps) - bg
q = normalize(quat_mul(q, small_angle_quat(omega * dt)))
a_nav = quat_to_rot(q) @ (acc_g * G) - np.array([0.0, 0.0, G]) - ba
p = p + v * dt + 0.5 * a_nav * dt * dt
v = v + a_nav * dt
P = F @ P @ F.T + Q

# update: GPS 位置观测
innov = z_gps_enu - p
K = P @ H.T @ np.linalg.inv(H @ P @ H.T + R)
dx = K @ innov
p += dx[0:3]; v += dx[3:6]
q = normalize(quat_mul(small_angle_quat(dx[6:9]), q))
bg += dx[9:12]; ba += dx[12:15]

# static constraint: 静止时约束水平速度
if is_static(acc_g, gyro_dps, gps_speed):
    update_velocity(np.array([0.0, 0.0]), R_zupt)
"""),
            ("p", "ESKF 的实现把非线性的姿态积分保留在名义状态中，把位置、速度、姿态小角度误差和传感器零偏放入 15 维误差状态。GPS 只在低频到来时修正位置和速度，IMU 在两次 GPS 之间负责连续预测；静止约束用于抑制普通单点 GPS 在静止时的随机游走。"),
            ("p", "最新 GPS/IMU 后处理结果中，ESP32 同步日志包含 5975 行、有效 GPS 点 1116 个，ESP GPS 轨迹距离约 632.48 m，鲁棒 ESKF 轨迹距离约 912.85 m；相对于手机 GPX 参考，ESP GPS 最近距离中位数为 2.80 m、95 分位为 15.75 m，ESKF 中位数为 4.88 m、95 分位为 15.37 m。创新量中位数为 0.318 m、95 分位为 2.344 m，说明该松耦合版本可用于低速步行轨迹展示和滤波一致性分析，但绝对轨迹精度仍受 GPS 单点定位、多路径和时间同步限制。"),
            ("p", "为避免把不稳定结果写成正向结论，本文还补充了 Q 参数敏感性扫描。若直接使用未加入鲁棒约束的同步数据，多个 Q 倍率组合都会被脚本标记为 diverged，表现为轨迹长度远大于 GPS 实测轨迹、位置创新达到异常量级。该结果说明 ESKF 不是简单调参即可稳定，必须同时保证坐标轴映射、时间同步、重力去除和静止约束正确。"),
        ],
    )

    # A compact data consistency note near the start of chapter 4.
    ch4 = first_para(doc, "4 标定与精度提升算法")
    insert_sequence(
        ch4,
        [
            ("p", "【数据核对说明】本章数值已按 sensor-final-project/data 中的最新分析结果修订：加速度计采用 12 参数仿射模型结果，磁力计采用 3600 点椭球标定与鲁棒重拟合结果，陀螺仪采用 gyro_allan_20260616_205634.csv 的 30 min Allan 方差结果，BMP280 采用 bmp280_height_change_20260617_084325.csv 的高度滤波结果。"),
        ],
    )

    ch5 = first_para(doc, "5 多传感器融合")
    insert_sequence(
        ch5,
        [
            ("p", "【数据核对说明】本章姿态融合标准差和更新率已按 attitude_fusion_static_std.csv 与 attitude_fusion_update_rate.csv 更新；GPS/IMU 结果采用 gps_esp32_phone_eskf_summary_20260701_223520.csv 与 ESKF 鲁棒后处理结果，不再使用旧版异常发散指标作为正向性能结论。"),
        ],
    )

    delete_paragraphs_starting(
        doc,
        [
            "图：静止状态下陀螺仪三轴输出与温度变化曲线",
            "图：静止状态下陀螺仪三轴零偏分布直方图",
            "图：陀螺仪三轴 Allan 偏差曲线",
        ],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
