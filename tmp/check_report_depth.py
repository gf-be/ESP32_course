import sys
from pathlib import Path
from docx import Document


report = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(r"F:\mechineSight\stm32\罗丹\大作业\sensor-final-project\docs\new_pcb_report_check_补充实时GPS_ESKF_ECharts.docx")
doc = Document(report)
text = "\n".join(par.text for par in doc.paragraphs)

keys = [
    "SVD", "奇异值", "最小奇异", "齐次方程", "桌面", "水平", "温度补偿", "二次多项式",
    "Runge", "离群点", "RANSAC", "Huber", "t_fine", "Madgwick", "beta", "Yaw",
    "可观测", "Q矩阵", "高度KF", "alpha=0.98", "互补滤波", "I2C", "400kHz",
    "中断", "mutex", "安装倾斜", "尼龙", "铺铜", "误差预算", "datasheet",
    "距离", "WiFi", "HMC5883L", "ARW", "BI", "P0",
]

print("report:", report)
print("paragraphs:", len(doc.paragraphs))
for key in keys:
    print(f"{key}: {text.lower().count(key.lower())}")

print("\nheadings:")
for par in doc.paragraphs:
    style = par.style.name if par.style else ""
    if style.startswith("Heading") or par.text.strip().startswith(("第", "附录")):
        t = par.text.strip()
        if t:
            print(f"- {style}: {t[:90]}")
