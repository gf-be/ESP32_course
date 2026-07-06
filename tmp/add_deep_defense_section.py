from __future__ import annotations

import csv
import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(r"F:\mechineSight\stm32\罗丹\大作业")
SRC = ROOT / "sensor-final-project" / "docs" / "new_pcb_report_check_补充实时GPS_ESKF_ECharts.docx"
DST = ROOT / "sensor-final-project" / "docs" / "new_pcb_report_check_补充答辩深度分析.docx"


def read_csv_rows(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def item_map(path: Path) -> dict[str, str]:
    out: dict[str, str] = {}
    for row in read_csv_rows(path):
        if "item" in row and "value" in row:
            out[row["item"]] = row["value"]
        elif "model" in row:
            out[row["model"]] = row
    return out


def style_run(run, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    run.bold = bold


def add_para(doc: Document, text: str = "", style: str | None = None, bold_prefix: str | None = None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        style_run(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        style_run(r2)
    else:
        r = p.add_run(text)
        style_run(r)
    return p


def set_cell_text(cell, text: str, bold=False, font_size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(font_size)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_cm: float):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(Cm(width_cm).twips)))
    tc_w.set(qn("w:type"), "dxa")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, font_size=8.5)
        shade_cell(hdr[i], "D9EAF7")
        set_cell_width(hdr[i], widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, font_size=8.5)
            set_cell_width(cells[i], widths[i])
    doc.add_paragraph()
    return table


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        style_run(r)


def f3(v: str | float) -> str:
    try:
        return f"{float(v):.3f}"
    except Exception:
        return str(v)


def main():
    shutil.copy2(SRC, DST)

    accel_rows = read_csv_rows(ROOT / "sensor-final-project" / "data" / "accel_6pos_model_compare.csv")
    accel = {row["model"]: row for row in accel_rows}
    mag = item_map(ROOT / "data" / "analysis" / "mag_ellipsoid_summary.csv")
    with (ROOT / "sensor-final-project" / "data" / "mag_wifi_compare_summary.json").open("r", encoding="utf-8") as f:
        wifi = json.load(f)
    allan = item_map(ROOT / "data" / "analysis" / "gyro_allan_summary.csv")
    attitude_rows = read_csv_rows(ROOT / "sensor-final-project" / "data" / "attitude_fusion_static_std.csv")
    attitude = {(r["phase"], r["algorithm"]): r for r in attitude_rows}
    update = item_map(ROOT / "sensor-final-project" / "data" / "attitude_fusion_update_rate.csv")
    bmp = item_map(ROOT / "data" / "analysis" / "bmp280_summary.csv")
    gps = item_map(ROOT / "sensor-final-project" / "data" / "analysis" / "gps_esp32_phone_eskf_summary_20260701_223520.csv")
    freq = item_map(ROOT / "sensor-final-project" / "data" / "performance" / "frequency_test_20260616_215936.csv")

    doc = Document(DST)
    doc.add_page_break()
    add_para(doc, "附录 E 答辩深化分析：物理含义、算法边界与误差预算", "Heading 1")
    add_para(
        doc,
        "本附录面向答辩追问，对前文实验结果进一步解释其数学来源、物理含义和工程限制。"
        "前文已经给出采集数据、图表和基本指标，但对评分标准中的高分问题，还需要回答“为什么这样做、"
        "误差从哪里来、参数改变后会发生什么、如果用于真实系统应该如何改进”。因此本节不重复罗列曲线，"
        "而是把标定、融合、总线调度、PCB 布局和系统误差预算串联起来。"
    )

    add_para(doc, "E.1 追问覆盖性检查", "Heading 2")
    coverage_rows = [
        ["Q6", "磁力计椭球 DLS+SVD", "已有公式和结果；本次补充为什么取最小奇异值、采样覆盖不足时的退化。"],
        ["Q7", "六位置 12 参数标定", "已有 12 参数结果；本次补充桌面倾斜 2° 如何进入矩阵，以及如何用重力模值约束消除。"],
        ["Q8", "温度补偿模型", "已有温度漂移讨论；本次明确二次多项式系数必须由多温点最小二乘确定，当前数据不支撑夸大结论。"],
        ["Q9", "离群点与鲁棒估计", "已有 RANSAC 提及；本次补充 SVD 最小二乘对离群点敏感、RANSAC/Huber 的实现办法。"],
        ["Q10", "BMP280 t_fine", "已有 BMP280 实验；本次补充 Bosch 补偿算法中温度与气压交叉敏感的物理意义。"],
        ["Q11", "Madgwick beta", "已有 Madgwick 指标；本次补充 beta=0.035 的含义和大机动时应自适应降低。"],
        ["Q12", "姿态可观测性", "已有静态标准差；本次补充 Roll/Pitch 与 Yaw 可观测性差异、无磁力计 yaw 漂移估算。"],
        ["Q13", "KF/ESKF 的 Q 矩阵", "已有 ESKF 程序；本次补充 Q 过大/过小的工程现象和当前 Q 的量级。"],
        ["Q14", "GPS/气压互补 alpha", "已有互补滤波；本次补充 alpha 对应时间常数和天气气压漂移导致的高度误差。"],
        ["Q15", "I2C 400 kHz 调度", "已有 I2C 扫描；本次补充单主机顺序读、IRQ 标志位、RTOS mutex 的处理。"],
        ["Q16", "安装倾斜与标定", "已有 PCB 实验；本次补充 12 参数矩阵吸收安装角的前提。"],
        ["Q17", "无人机磁干扰方案", "原文较少；本次补充延伸杆、电流补偿、在线重标定、双 GPS 测向等方案比较。"],
        ["Q18", "PCB Layout 批判", "已有硬件说明；本次补充地平面、去耦、I2C 参考地、磁力计距离的量化改进方向。"],
        ["Q19", "系统误差预算", "已有单项指标；本次补充 Roll/Pitch、Yaw、高度、GPS/ESKF 的瓶颈归因表。"],
        ["Q20", "真实无人机应用", "原文较少；本次补充上机前最优先的三个技术改进点和预期收益。"],
    ]
    add_table(doc, ["题号", "主题", "本报告体现与本次补充"], coverage_rows, [1.5, 4.1, 10.2])

    add_para(doc, "E.2 标定算法的数学深度", "Heading 2")
    add_para(doc, "1. 椭球拟合 SVD 为什么取最小奇异值。", "Heading 3")
    add_para(
        doc,
        "磁力计椭球拟合可以写成齐次线性方程 Phi q = 0，其中 q 是二次曲面参数。"
        "如果直接令 q=0，残差最小但没有物理意义，因此需要加约束 ||q||=1，把问题变成"
        " min ||Phi q||, s.t. ||q||=1。对 Phi 做 SVD 分解 Phi=U Sigma V^T 后，"
        "残差能量等于各奇异值平方加权。最小奇异值对应的右奇异向量，是在单位长度约束下使残差最小的方向，"
        "所以应取 V 的最后一列，而不是最大奇异值对应方向。最大奇异值方向代表数据变化最强的方向，"
        "会使残差最大，不能作为椭球参数解。"
    )
    add_para(
        doc,
        f"本次磁力计椭球数据样本数为 {mag['samples']} 行，硬铁中心约为 "
        f"({f3(mag['hard_iron_x'])}, {f3(mag['hard_iron_y'])}, {f3(mag['hard_iron_z'])}) raw count，"
        f"三轴主半径为 {f3(mag['ellipsoid_radius_1'])}, {f3(mag['ellipsoid_radius_2'])}, "
        f"{f3(mag['ellipsoid_radius_3'])} raw count，轴不均衡系数为 {f3(mag['axis_imbalance'])}。"
        "如果采集只覆盖半个球面，Phi 的某些列会高度相关，多个小奇异值接近，解会变得病态："
        "椭球中心可能被拉向未覆盖方向，软铁矩阵会过度补偿，校准后散点不再接近球面。"
    )

    add_para(doc, "2. 六位置标定中桌面倾斜的影响。", "Heading 3")
    add_para(
        doc,
        "六位置标定的真值不是“桌面法向一定等于重力方向”，而是每个姿态下加速度计应测得一个模长为 1 g 的重力向量。"
        "如果桌面倾斜 2°，水平误差分量约为 sin(2°)=0.0349 g，即 34.9 mg。"
        "若六个面都以同一倾斜桌面作为参考，12 参数仿射模型可能把这部分误差吸收到非正交矩阵中，"
        "看起来标定残差变小，但矩阵中会混入夹具误差。"
    )
    add_para(
        doc,
        f"本实验中原始均值向量 RMS 误差为 {accel['raw']['mean_vector_rms_error_mg']} mg，"
        f"只做零偏和比例后为 {accel['diag6_bias_scale']['mean_vector_rms_error_mg']} mg，"
        f"12 参数仿射模型后降至 {accel['affine12_bias_scale_misalignment']['mean_vector_rms_error_mg']} mg，"
        f"均值误差改善约 {accel['affine12_bias_scale_misalignment']['improvement_vs_raw_mean_vector']} 倍。"
        "矩阵中的非对角项对应轴间非正交、模块安装角和比例耦合；若只有零偏误差，A_inv 应退化为对角矩阵，"
        "对角元素为三个轴灵敏度倒数，非对角元素应接近 0。"
    )
    add_bullets(doc, [
        "消除桌面误差的做法：使用精密六面体或水平仪校正夹具；每个方向重复摆放并取相反面配对；在优化中加入 |a_cal|=1 的球面约束，而不是完全相信桌面法向。",
        "若模块焊死在 PCB 上，12 参数矩阵可以吸收固定安装角；若模块用杜邦线每次位置变化，标定矩阵不再适用。",
    ])

    add_para(doc, "3. 温度补偿二次多项式的边界。", "Heading 3")
    add_para(
        doc,
        "温度补偿模型 b(T)=alpha+beta T+gamma T^2 的系数，应由多个温度点下的静止零偏数据做最小二乘得到。"
        "这要求温度覆盖工作区间，例如 20°C、30°C、40°C、50°C 分段采样。当前 Allan 实验温度范围只有 "
        f"{f3(allan['temp_range'])} °C，适合说明短时温度稳定性，但不足以支撑完整二次温度标定。"
        "因此报告中将温度补偿作为算法实现方案，而不把二次系数作为已完成的精度结论。"
    )
    add_para(
        doc,
        "如果只在 20°C 到 50°C 内拟合，然后拿到 60°C 使用，就从插值变成外推。"
        "外推误差不再受训练数据约束，高阶多项式还可能在区间外出现振荡，即 Runge 现象。"
        "工程上应限制模型适用温区，或者采用分段线性/低阶模型，并在运行时记录温度是否超出标定范围。"
    )

    add_para(doc, "4. 离群点对椭球拟合的影响。", "Heading 3")
    add_para(
        doc,
        "标准 SVD 椭球拟合本质是 L2 最小二乘，对离群点敏感。若采集时手机、螺丝刀或电脑靠近磁力计，"
        "少量异常点会以平方残差形式强烈拉动椭球中心和轴长。"
        "本实验校准前半径标准差为 "
        f"{f3(mag['raw_radius_std'])} raw count，校准后为 {f3(mag['cal_radius_std'])} raw count，"
        f"CV 仅由 {f3(mag['raw_radius_cv'])} 降到 {f3(mag['cal_radius_cv'])}，改善倍数 {f3(mag['cv_improvement'])}。"
        "这说明椭球拟合能估计硬铁中心，但采样覆盖、外部磁干扰或板上电磁环境仍是主要限制。"
    )
    add_bullets(doc, [
        "可实现的鲁棒方法一：先用原始 SVD 得到初值，计算每个点校准后的半径残差，用 MAD 或 3 sigma 阈值剔除异常点后重新拟合。",
        "可实现的鲁棒方法二：RANSAC 随机抽取最小样本拟合椭球，选择内点最多的模型，再用内点做精细拟合。",
        "可实现的鲁棒方法三：用 Huber 损失或迭代加权最小二乘，让大残差点权重下降，避免单个异常点主导结果。",
    ])

    add_para(doc, "5. BMP280 t_fine 的物理含义。", "Heading 3")
    add_para(
        doc,
        "Bosch datasheet 中的 t_fine 不是普通临时变量，而是由温度 ADC 和温度校准系数计算得到的细化温度状态。"
        "BMP280 的压力敏感元件存在温度交叉敏感，压力补偿公式需要知道当前芯片温度，所以 t_fine 在温度补偿中产生，"
        "又被传递到气压补偿公式中。若温度和气压读取相差 100 ms，在温度变化很快的情况下会引入交叉误差，"
        "量级可估为 1 Pa/°C × 温变速率 × 0.1 s。"
    )
    add_para(
        doc,
        f"实测 BMP280 静止压力标准差为 {f3(bmp['static_pressure_std'])} Pa，折算高度标准差约 "
        f"{f3(bmp['static_altitude_std'])} m；静止温度标准差为 {f3(bmp['static_temperature_std'])} °C。"
        "相比传感器分辨率，绝对高度的主导误差通常不是短时噪声，而是 P0 基准气压、温度交叉敏感和天气漂移。"
        "例如 P0 误差 1 hPa 会带来约 8.4 m 高度误差。"
    )

    add_para(doc, "E.3 融合算法的工程理解", "Heading 2")
    add_para(doc, "1. Madgwick beta 与大机动。", "Heading 3")
    add_para(
        doc,
        "本报告离线 Madgwick MARG 滤波器 beta 取 0.035。beta 可以理解为梯度下降校正项的权重，"
        "越大越相信加速度计/磁力计给出的重力和地磁方向，姿态收敛更快，但在强机动和磁干扰下也更容易被错误观测拉偏。"
        "若飞机以 200 deg/s 急转 0.5 s，陀螺积分对应约 100° 姿态变化；此时加速度计中可能包含明显离心加速度，"
        "测到的不再是纯重力方向。工程上应根据 |a|-1g、磁场模长残差或角速度大小自适应降低 beta，"
        "让短时姿态主要由陀螺积分维持，机动结束后再恢复较大 beta 进行收敛。"
    )
    lvl_mad = attitude[("level_static", "madgwick")]
    add_para(
        doc,
        f"在本次静止水平数据中，Madgwick 的 Roll/Pitch/Yaw 标准差分别为 "
        f"{f3(lvl_mad['roll_std_deg'])}°、{f3(lvl_mad['pitch_std_deg'])}°、{f3(lvl_mad['yaw_std_deg'])}°；"
        f"离线计算更新率约 {f3(update['madgwick_update_hz'])} Hz，说明该算法在 PC 端计算裕量充足。"
    )

    add_para(doc, "2. Roll/Pitch 与 Yaw 的可观测性。", "Heading 3")
    lvl_mah = attitude[("level_static", "mahony")]
    add_para(
        doc,
        f"Mahony 水平静止 Roll/Pitch/Yaw 标准差分别为 {f3(lvl_mah['roll_std_deg'])}°、"
        f"{f3(lvl_mah['pitch_std_deg'])}°、{f3(lvl_mah['yaw_std_deg'])}°。短时标准差中 yaw 看起来并不大，"
        "但这只说明当前室内磁场在短时间内较稳定，不代表绝对航向精度同样高。"
        "加速度计只能观测重力方向，因此能约束 Roll 和 Pitch；绕重力轴旋转时，加速度计读数不变，Yaw 对加速度计不可观，"
        "必须依赖磁力计、GPS 航向或其他外部观测。"
    )
    yaw_drift_10min = abs(float(allan["gyro_z_bias"])) * 600.0
    add_para(
        doc,
        f"若去掉磁力计且不做陀螺零偏补偿，Z 轴静止均值偏置约 {f3(allan['gyro_z_bias'])} deg/s，"
        f"10 分钟理论 yaw 漂移可达约 {yaw_drift_10min:.1f}°。"
        "即使做了静态零偏扣除，温度变化和零偏随机游走仍会让 yaw 随时间漂移，这就是 yaw 精度通常差于 Roll/Pitch 的根本原因。"
    )

    add_para(doc, "3. Q 矩阵与滤波器自信程度。", "Heading 3")
    add_para(
        doc,
        "卡尔曼滤波中的 Q 矩阵描述系统模型本身的不确定性：IMU 积分、未建模加速度、陀螺零偏随机游走、加速度计零偏随机游走等。"
        "本项目 15 维简化 ESKF 的误差状态为 dx=[dp,dv,dtheta,dbg,dba]，名义状态为 [p,v,q,bg,ba]。"
        "程序中使用 q_pos=0.0004，q_acc=0.45^2，q_gyr=(1.5 deg/s)^2，q_bg=(0.08 deg/s)^2，q_ba=0.015^2。"
        "这些数值不是传感器 datasheet 的直接复制，而是结合低速步行数据和 GPS 更新频率调出的工程量级。"
    )
    add_para(
        doc,
        "如果 Q 比真实过程噪声小 100 倍，滤波器会过于相信自己的预测，协方差 P 偏小，GPS 或气压观测来了也不愿修正，"
        "表现为过度自信、滞后甚至发散；如果 Q 大 100 倍，滤波器会过度相信观测，轨迹抖动变大，退化成带权平均，"
        "平滑效果消失。答辩时可以把这两种现象分别称为 over-confidence/filter divergence 和过噪声驱动的 measurement chasing。"
    )

    add_para(doc, "4. GPS/气压互补滤波 alpha 的时间常数。", "Heading 3")
    add_para(
        doc,
        "互补滤波中 alpha=0.98 时，若采样周期 dt=0.02 s，则时间常数近似 tau=dt/(1-alpha)=1 s。"
        "这意味着约 1 s 后低频观测已经明显渗入融合结果。气压高度对天气很敏感，台风前夕 1 小时下降 3 hPa 时，"
        "高度等效变化约为 3×8.4=25.2 m。如果不及时用 GPS 或本地气象站重置 P0，融合高度会把天气变化误认为真实高度变化。"
        "因此用于长时间户外运行时，应提高 GPS/P0 校准频率，或者把 P0 作为慢变状态纳入 KF，而不是固定常数。"
    )

    add_para(doc, "5. I2C 总线调度。", "Heading 3")
    add_para(
        doc,
        f"本项目 I2C 使用 SDA=GPIO21、SCL=GPIO22、400 kHz，总线上挂载 MPU6050(0x68)、HMC5883L(0x1E) 和 BMP280(0x76)。"
        f"频率测试中 IMU 读取约 {f3(freq['imu_sampling_rate_hz'])} Hz，融合更新约 {f3(freq['fusion_update_rate_hz'])} Hz，"
        "满足课程对实时性的要求。当前 MicroPython 主循环采用单线程顺序调度：先读 GPS 串口缓冲，再读 IMU 和磁力计，最后更新融合。"
        "I2C 是主从总线，主机控制时序；即使 MPU6050 中断在 BMP280 读取途中触发，中断也只是置位或提示数据已就绪，"
        "不会打断正在进行的 I2C 事务。若迁移到 FreeRTOS 多任务，应把 I2C burst read 包在 mutex/critical section 内，"
        "避免任务切换导致一个传感器事务被另一个任务插入。"
    )

    add_para(doc, "E.4 系统工程与批判性思维", "Heading 2")
    add_para(doc, "1. PCB 安装、磁干扰距离与 WiFi 对比。", "Heading 3")
    shift = wifi["delta_on_minus_off"]
    add_para(
        doc,
        f"WiFi 关闭与开启扫描对比中，磁场均值向量变化为 "
        f"({shift['dx']:.3f}, {shift['dy']:.3f}, {shift['dz']:.3f}) raw count，"
        f"向量位移 {shift['vector_shift']:.3f} raw count，约占磁场模长 {shift['vector_shift_percent_of_field']:.3f}%。"
        "这说明当前摆放和距离下，ESP32 WiFi 对磁力计均值的影响较小，但并不意味着布局可以任意靠近。"
    )
    add_para(
        doc,
        "若把磁力计到 ESP32 天线/高频电流路径的距离从 15 mm 减到 5 mm，准静态磁场耦合可按 1/r^3 做保守估计，"
        "干扰幅值可能增加 (15/5)^3=27 倍；若按远场辐射近似，则也会有约 3 倍增大。"
        "以当前 1.93 count 位移估算，27 倍后可能达到约 52 count，约为地磁模长的 12% 左右，"
        "对应航向误差可进入数度量级。因此磁力计应尽量远离 WiFi 天线、DC-DC、电机线和大电流回路。"
    )

    add_para(doc, "2. PCB Layout 可改进点。", "Heading 3")
    layout_rows = [
        ["磁力计位置", "远离 ESP32 天线和大电流路径，优先放在板边并预留禁布区", "干扰按距离 1/r 到 1/r^3 下降；15 mm 增至 45 mm 时，准静态耦合可下降约 27 倍"],
        ["去耦电容", "每个传感器 VCC 旁放 0.1 uF + 1 uF，走线短且直接回地", "可把供电尖峰和 I2C 翻转引起的局部电源噪声降低约 10-20 dB"],
        ["连续地平面", "I2C 和 UART 下方保持连续参考地，避免跨分割地", "减小环路面积和共模噪声，降低边沿振铃，提升 400 kHz I2C 可靠性"],
        ["传感器固定", "MPU6050/BMP280/HMC5883L 尽量采用焊接固定或螺柱固定，不使用会晃动的杜邦线", "固定安装角可由 12 参数标定吸收；可变安装角无法由一次标定补偿"],
    ]
    add_table(doc, ["改进项", "具体做法", "预期收益"], layout_rows, [3.0, 6.2, 6.6])

    add_para(doc, "3. 完整误差预算。", "Heading 3")
    err_rows = [
        ["Roll/Pitch", f"静止标准差约 0.02-0.04°；12 参数加速度计均值 RMS 残差 {accel['affine12_bias_scale_misalignment']['mean_vector_rms_error_mg']} mg", "瓶颈为安装角、标定姿态真值和动态加速度假设；纯软件可通过更好六面夹具和自适应加速度权重再提升"],
        ["Yaw/航向", f"短时 yaw 标准差约 0.04-0.13°；磁力计校准后半径 std {f3(mag['cal_radius_std'])} count", "瓶颈为磁干扰、软铁残差和采样覆盖；纯软件可做鲁棒椭球、磁场一致性门限，但硬件距离更关键"],
        ["高度", f"BMP280 静止高度 std {f3(bmp['static_altitude_std'])} m", "短时噪声小，绝对高度瓶颈为 P0；1 hPa P0 误差约 8.4 m，天气 3 hPa 可带来 25 m 级漂移"],
        ["GPS/ESKF 轨迹", f"相对手机 GPX：ESKF 中位误差 {f3(gps['eskf_to_phone_median'])} m，P95 {f3(gps['eskf_to_phone_p95'])} m", "瓶颈为 NEO-6M 单点定位、HDOP、多路径和 IMU 低速可观测性；纯软件主要改善平滑和异常剔除，难突破 GPS 观测本身"],
        ["实时性", f"IMU 采样 {f3(freq['imu_sampling_rate_hz'])} Hz，融合 {f3(freq['fusion_update_rate_hz'])} Hz", "频率满足课程要求；后续瓶颈在 MicroPython 抖动和串口输出，可迁移 C/FreeRTOS 降低延迟"],
    ]
    add_table(doc, ["指标", "本项目实测/估计", "瓶颈与改进方向"], err_rows, [2.8, 5.0, 8.0])

    add_para(doc, "4. 面向真实无人机的优先改进。", "Heading 3")
    add_bullets(doc, [
        "磁力计系统：采用远离电机/电调的延伸杆，配合电流补偿（类似 COMPASS_MOT）和磁场一致性门限。预期可把磁干扰导致的航向误差从数度降低到约 1-2°量级。",
        "惯性与振动系统：更换低噪声 IMU、增加机械减振和硬件同步时间戳。预期可降低姿态高频抖动和积分误差，提升 ESKF 预测段可靠性。",
        "导航观测系统：引入 RTK GNSS 或双天线 GNSS 测向，并把气压 P0 作为慢变状态估计。预期 GPS 轨迹误差由米级降到分米级，航向不再完全依赖磁力计。",
    ])

    add_para(doc, "E.5 后续实现建议", "Heading 2")
    add_bullets(doc, [
        "磁力计脚本增加 coverage score、SVD condition number、RANSAC/Huber 两种鲁棒拟合选项，并输出内点率和校准前后半径残差表。",
        "温度补偿实验至少采集 20°C、30°C、40°C、50°C 四个温区，每个温区静止 5-10 分钟，分别拟合加速度计和陀螺仪零偏温度模型。",
        "ESKF 参数实验保留三组 Q：保守、当前、激进，对比 innovation、P95 轨迹误差和静止漂移，用曲线说明 Q 调参效果。",
        "实时网页演示中保留三轴姿态板、GPS/ESKF 轨迹、innovation 和 sigma 曲线，演示时先静止再移动，说明静止约束如何抑制 GPS 漂移。",
    ])

    doc.save(DST)
    print(DST)


if __name__ == "__main__":
    main()
