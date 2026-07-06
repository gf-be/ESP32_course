from pathlib import Path
from docx import Document

docx = Path(r"F:\mechineSight\stm32\罗丹\大作业\sensor-final-project\docs\new_pcb_report_check_补充答辩深度分析.docx")
doc = Document(docx)
parts = [p.text for p in doc.paragraphs]
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            parts.append(cell.text)
text = "\n".join(parts)
required = [f"Q{i}" for i in range(6, 21)] + [
    "附录 E",
    "SVD",
    "最小奇异值",
    "Runge",
    "RANSAC",
    "Huber",
    "t_fine",
    "Madgwick",
    "beta",
    "可观测",
    "mutex",
    "误差预算",
    "P0",
]
missing = [x for x in required if x not in text]
headings = [p.text for p in doc.paragraphs if p.style and p.style.name.startswith("Heading")]
appendix = [h for h in headings if "E." in h or "附录 E" in h]
print("docx:", docx)
print("paragraphs:", len(doc.paragraphs))
print("tables:", len(doc.tables))
print("missing:", missing)
print("appendix headings:")
for h in appendix:
    print("-", h)
