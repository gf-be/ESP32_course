from pathlib import Path
from docx import Document

docx = Path(r"F:\mechineSight\stm32\罗丹\大作业\sensor-final-project\docs\new_pcb_report_check_补充算法实现说明.docx")
doc = Document(docx)
parts = [p.text for p in doc.paragraphs]
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            parts.append(cell.text)
text = "\n".join(parts)
required = [
    "附录 F",
    "软件总体数据流",
    "传感器驱动",
    "12 参数仿射",
    "椭球标定",
    "Allan",
    "互补滤波",
    "Mahony",
    "Madgwick",
    "15 维简化 ESKF",
    "高度融合",
    "AI 去噪",
    "实时演示",
    "后续代码改进路线",
]
missing = [x for x in required if x not in text]
headings = [p.text for p in doc.paragraphs if p.style and p.style.name.startswith("Heading")]
appendix = [h for h in headings if "F." in h or "附录 F" in h]
print("docx:", docx)
print("paragraphs:", len(doc.paragraphs))
print("tables:", len(doc.tables))
print("missing:", missing)
print("appendix F headings:")
for h in appendix:
    print("-", h)
