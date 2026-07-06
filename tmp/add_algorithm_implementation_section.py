from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(r"F:\mechineSight\stm32\罗丹\大作业")
SRC = ROOT / "sensor-final-project" / "docs" / "new_pcb_report_check_补充答辩深度分析.docx"
DST = ROOT / "sensor-final-project" / "docs" / "new_pcb_report_check_补充算法实现说明.docx"


def style_run(run, bold=False, size=10.5):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def add_para(doc: Document, text: str = "", style: str | None = None, bold_prefix: str | None = None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        style_run(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        style_run(r2)
    else:
        r = p.add_run(text)
        style_run(r)
    return p


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        style_run(r)


def set_cell_text(cell, text: str, bold=False, font_size=8.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(font_size)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_cm: float):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(Cm(width_cm).twips)))
    tc_w.set(qn("w:type"), "dxa")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, bold=True)
        shade_cell(cell, "D9EAF7")
        set_cell_width(cell, widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            set_cell_width(cells[i], widths[i])
    doc.add_paragraph()
    return table


def main():
    shutil.copy2(SRC, DST)
    doc = Document(DST)
    doc.add_page_break()

    add_para(doc, "附录 F 关键模块与融合算法实现说明", "Heading 1")
    add_para(
        doc,
        "本附录补充说明系统各模块在程序中的实现方式。前文以实验数据和图表为主，本节从代码执行流程、数学模型、"
        "参数含义和工程取舍角度解释“算法如何落地”。报告答辩时可以按本节顺序说明：先讲传感器数据如何变成统一坐标系下的物理量，"
        "再讲单传感器标定如何降低系统误差，最后讲多传感器融合如何把互补信息合成为姿态、位置和高度状态。"
    )

    add_para(doc, "F.1 软件总体数据流", "Heading 2")
    add_para(
        doc,
        "系统的软件数据流分为五层：驱动读取层、单位换算层、标定补偿层、融合估计层和可视化/记录层。"
        "ESP32 端负责以固定顺序读取 I2C 与 UART 数据，完成基础单位换算和实时融合；PC 端负责串口接收、CSV 保存、"
        "离线分析和 ECharts/Web 姿态演示。这样的划分保证演示时算法可以实时运行，同时保留完整数据用于论文复现实验。"
    )
    add_table(doc, ["层级", "输入", "处理", "输出"], [
        ["驱动读取层", "I2C: MPU6050/HMC5883L/BMP280；UART: GPS NMEA", "按寄存器地址读取原始计数；GPS 解析 GGA/RMC 等语句", "raw ax/ay/az/gx/gy/gz、mx/my/mz、pressure/temp、lat/lon/alt"],
        ["单位换算层", "传感器原始计数", "MPU6050 按量程换算为 g 和 deg/s；磁力计保留 raw count 或按比例换算；BMP280 用 datasheet 补偿", "统一物理量"],
        ["标定补偿层", "单位换算后的物理量", "加速度计 12 参数仿射、陀螺零偏扣除、磁力计硬铁/软铁校正、气压 P0 修正", "校准后的 IMU/MAG/BARO/GPS 数据"],
        ["融合估计层", "校准后的多源数据", "互补滤波、Mahony/Madgwick、15 维简化 ESKF、静止约束和 GPS 门限更新", "roll/pitch/yaw、四元数、ENU 位置、速度、协方差"],
        ["可视化记录层", "融合输出和原始观测", "CSV 落盘、ECharts 曲线、3D 姿态板、folium 轨迹叠加", "实验图表、实时演示和报告依据"],
    ], [2.6, 4.0, 6.0, 4.0])

    add_para(doc, "F.2 传感器驱动与预处理实现", "Heading 2")
    add_para(doc, "1. MPU6050/MPU6500 IMU 读取。", "Heading 3")
    add_para(
        doc,
        "IMU 读取采用 I2C burst read，从加速度计起始寄存器连续读取 14 字节，解析为 ax, ay, az, temp, gx, gy, gz。"
        "这种一次性连续读取比逐寄存器读取更适合实时融合，因为同一帧中加速度和角速度时间更一致。"
        "读取后按量程换算：加速度约除以 16384 得到 g，陀螺约除以 131 得到 deg/s，再扣除 Allan 静止实验得到的陀螺零偏。"
    )
    add_bullets(doc, [
        "实现要点：先完成 I2C scan 确认 0x68 存在，再唤醒 MPU6050，最后进入循环读取。",
        "工程意义：陀螺负责短时角速度积分，动态响应快但会漂移；加速度计提供重力方向，能约束 Roll/Pitch 低频漂移。",
        "注意边界：运动中加速度计不只包含重力，也包含线加速度，因此姿态融合中必须降低异常加速度的权重。",
    ])

    add_para(doc, "2. HMC5883L 磁力计读取。", "Heading 3")
    add_para(
        doc,
        "磁力计初始化为连续测量模式，主循环读取三轴磁场 raw count。磁力计用于观测地磁水平分量，为 yaw 提供绝对参考。"
        "在进入融合前，先减去硬铁中心，再乘以软铁校正矩阵，把椭球散点尽量拉回球面。"
        "融合时还会结合 Roll/Pitch 做倾斜补偿，即先把磁场旋转到水平面，再用 atan2 求航向角。"
    )
    add_bullets(doc, [
        "实现要点：地磁观测不能直接等于 yaw，必须先做硬铁/软铁校正和倾斜补偿。",
        "工程意义：磁力计解决 yaw 不可观问题，但对电流、铁磁物体和 WiFi/电机干扰敏感。",
        "门限策略：可用磁场模长残差、椭球半径残差或短时跳变检测来拒绝异常磁观测。",
    ])

    add_para(doc, "3. BMP280 气压计读取。", "Heading 3")
    add_para(
        doc,
        "BMP280 驱动按照 Bosch datasheet 的补偿公式实现。温度原始 ADC 先计算 t_fine，气压补偿再使用 t_fine 修正压力读数。"
        "得到压力后，用国际标准大气近似把气压转换为相对高度。由于绝对高度强依赖 P0，程序中把实验开始时的平均气压作为局部基准，"
        "报告中只把 BMP280 作为相对高度变化传感器，而不是绝对海拔传感器。"
    )

    add_para(doc, "4. GPS NMEA 解析。", "Heading 3")
    add_para(
        doc,
        "GPS 模块通过 UART 输出 NMEA 语句。程序主要解析 GGA/RMC 中的定位质量、卫星数、HDOP、经纬度和高度。"
        "经纬度先转换为十进制度，再以首次有效 GPS 点为原点，转换到局部 ENU 坐标。ESKF 更新时使用 HDOP 自适应设置 GPS 观测噪声，"
        "卫星数不足或 HDOP 过大时跳过更新，避免差观测破坏轨迹。"
    )

    add_para(doc, "F.3 标定算法实现说明", "Heading 2")
    add_para(doc, "1. 加速度计 12 参数仿射标定。", "Heading 3")
    add_para(
        doc,
        "加速度计标定采用 12 参数模型：a_raw = M a_true + b，或 a_cal = A_inv (a_raw - c)。"
        "其中 b/c 是三轴零偏，M/A_inv 的对角元素对应比例因子，非对角元素对应轴间耦合、非正交和固定安装角。"
        "六位置实验每个面取均值，构造最小二乘方程，求得能够把六个均值映射到 ±1 g 方向的仿射模型。"
    )
    add_bullets(doc, [
        "代码实现：读取六个姿态 CSV，按文件名或人工标注绑定真值方向；对每个姿态取均值；构造线性方程求 M 和 b；保存矩阵 CSV。",
        "算法作用：不仅消除零偏和比例误差，还能补偿固定安装角和轴间非正交。",
        "边界条件：如果六个姿态摆放不准，模型可能吸收夹具误差；若模块安装后移动，标定矩阵失效。",
    ])

    add_para(doc, "2. 磁力计椭球标定。", "Heading 3")
    add_para(
        doc,
        "磁力计标定的目标是把硬铁偏置和软铁拉伸造成的椭球点云恢复成球面。实现流程为：采集全方位旋转数据；"
        "建立二次曲面方程；用 SVD 求齐次最小二乘初值；从椭球中心得到硬铁偏置；对二次型矩阵特征分解得到主轴和比例因子；"
        "最后输出校正矩阵和校正前后 3D 散点图。"
    )
    add_bullets(doc, [
        "代码实现重点：SVD 取最小奇异值对应的右奇异向量，因为它是在 ||q||=1 约束下残差最小的参数方向。",
        "数据质量要求：旋转覆盖必须接近完整球面，否则椭球某些方向不可观，奇异值退化，中心和轴长会不稳定。",
        "鲁棒改进：先用半径残差剔除离群点，再二次拟合；后续可加入 RANSAC 或 Huber 迭代加权。",
    ])

    add_para(doc, "3. Allan 方差与噪声参数提取。", "Heading 3")
    add_para(
        doc,
        "Allan 方差实现上先对静止陀螺数据按不同聚合时间 tau 分组，计算相邻组均值差的均方根。"
        "短 tau 斜率约 -1/2 对应白噪声/角随机游走，中间最低点反映零偏不稳定性，长 tau 上升段对应慢变漂移。"
        "这些结果用于解释 KF/ESKF 中 R 和 Q 的选取：短时测量噪声进入 R，零偏随机游走进入 Q 中的 bg 状态。"
    )

    add_para(doc, "F.4 姿态融合算法实现说明", "Heading 2")
    add_para(doc, "1. 互补滤波。", "Heading 3")
    add_para(
        doc,
        "互补滤波把陀螺积分作为高频通道，把加速度计/磁力计解算角作为低频通道。"
        "实现形式为 angle = alpha*(angle + gyro*dt) + (1-alpha)*angle_meas。"
        "alpha 越接近 1，越相信陀螺，动态响应更平滑但漂移更慢被校正；alpha 越小，越相信加速度计/磁力计，静止收敛更快但动态时容易受线加速度影响。"
    )

    add_para(doc, "2. Mahony PI 滤波。", "Heading 3")
    add_para(
        doc,
        "Mahony 算法用四元数表示姿态，避免欧拉角奇异。程序先根据当前四元数预测重力和地磁方向，"
        "再与归一化后的加速度计/磁力计观测做叉乘，得到姿态误差 e。比例项 Kp*e 用于快速修正姿态，积分项 Ki∫e dt 用于估计并抵消陀螺零偏。"
        "ESP32 实时程序采用 Mahony PI 作为姿态内环，再把姿态四元数提供给 15 维 ESKF 做位置预测。"
    )
    add_bullets(doc, [
        "实现步骤：归一化 acc/mag；计算预测重力和磁场；求误差叉乘；修正 gyro；积分四元数；归一化四元数；输出 roll/pitch/yaw。",
        "参数含义：Kp 控制收敛速度，Ki 控制零偏估计速度；Kp 过大易抖动，Ki 过大易积分饱和。",
        "工程策略：当 |a|-1g 过大或磁场模长异常时，应降低对应观测权重，避免错误观测拉偏姿态。",
    ])

    add_para(doc, "3. Madgwick MARG 滤波。", "Heading 3")
    add_para(
        doc,
        "Madgwick 算法同样使用四元数，但校正项来自目标函数的梯度下降。目标函数描述“当前姿态下预测的重力/地磁方向”与"
        "加速度计/磁力计观测之间的误差，梯度给出姿态修正方向。程序中 beta 控制梯度校正强度：beta 大，观测校正强；beta 小，陀螺积分占主导。"
        "本项目离线对比中 beta=0.035，适合低速手持和桌面实验；若用于无人机急转，应根据角速度和 |a|-1g 自适应降低 beta。"
    )

    add_para(doc, "F.5 15 维简化 ESKF 实现说明", "Heading 2")
    add_para(
        doc,
        "15 维 ESKF 的核心思想是：名义状态用非线性形式正常积分，误差状态用小量线性化更新。"
        "本项目状态定义为位置 p(3)、速度 v(3)、姿态误差 dtheta(3)、陀螺零偏 bg(3)、加速度计零偏 ba(3)。"
        "名义姿态由 Mahony 四元数提供，IMU 预测把机体系加速度旋转到 ENU 坐标并扣除重力，GPS 到来时用位置观测更新 p，"
        "静止时加入零速度和弱位置保持约束，防止 GPS 静态漂移被误认为真实运动。"
    )
    add_table(doc, ["ESKF 步骤", "实现内容", "工程意义"], [
        ["预测", "读取 ax/ay/az 和 Mahony 四元数，将机体系加速度旋转到 ENU，积分 v 和 p", "在 GPS 两次更新之间保持连续轨迹"],
        ["协方差传播", "按 q_pos、q_acc、q_gyr、q_bg、q_ba 增大 P", "表达预测不确定性随时间增长"],
        ["GPS 更新", "HDOP 转换为 sigma_xy/sigma_z，构造位置残差 innovation，按 Kalman 增益修正状态", "让轨迹回到真实 GPS 观测附近"],
        ["静止约束", "当 |a|-1g 小且 gyro norm 小，触发 ZUPT 和位置保持伪观测", "抑制静止时 GPS 抖动造成的 ENU 漂移"],
        ["门限与保护", "卫星数/HDOP 不达标时跳过；innovation 过大时拒绝或放大协方差", "防止坏 GPS 点破坏滤波器"],
    ], [3.0, 7.0, 5.8])
    add_para(
        doc,
        "该 ESKF 是为课程硬件和低速步行数据设计的松耦合版本，不等同于飞控级紧耦合 GNSS/INS。"
        "它的优势是能在 ESP32/MicroPython 上实时运行，且状态维度和协方差意义清晰；限制是低速手持场景下 IMU 位置预测可观测性弱，"
        "轨迹精度最终仍主要受 GPS 观测质量、多路径和时间同步影响。"
    )

    add_para(doc, "F.6 高度融合与气压补偿实现说明", "Heading 2")
    add_para(
        doc,
        "高度估计采用气压计短时平滑与 GPS 低频约束的思想。BMP280 短时噪声小，适合观察相对高度变化；GPS 高度噪声大但长期不受天气气压漂移影响。"
        "互补滤波或一维 KF 可将气压高度作为高频相对变化，将 GPS 高度或人工 P0 校准作为低频基准。"
        "实现时需要把 P0 作为可更新参数，而不是固定常数；当天气变化明显时，P0 漂移会成为高度误差主导项。"
    )

    add_para(doc, "F.7 AI 去噪实现说明", "Heading 2")
    add_para(
        doc,
        "AI 去噪实验用已有 IMU 数据构造窗口样本，输入为 ax, ay, az, gx, gy, gz 六通道序列，输出为去噪后的中心点或窗口序列。"
        "1D-CNN 通过一维卷积学习局部时间模式，相当于一种数据驱动的非线性滤波器。报告中将其与原始数据、低通滤波、一阶卡尔曼进行对比，"
        "指标包括标准差、RMSE、SNR 或姿态角抖动。"
    )
    add_bullets(doc, [
        "实现流程：滑动窗口切片；训练/验证划分；用低噪声参考或平滑信号构造标签；训练 1D-CNN；输出训练曲线和去噪前后信号对比。",
        "工程边界：AI 去噪不能凭空恢复真实运动，只能利用数据统计规律抑制噪声；若训练样本只包含静止数据，动态泛化能力有限。",
        "报告表达：强调其作为增强算法对比，而不是替代物理标定和传感器融合。",
    ])

    add_para(doc, "F.8 实时演示与数据保存实现说明", "Heading 2")
    add_para(
        doc,
        "实时演示程序在 PC 端直接读取 COM4 串口，不依赖 Thonny Shell 文本复制。ESP32 输出 ESKF15 CSV 行，PC 程序一边解析并保存到 CSV，"
        "一边通过本地 HTTP 服务推送到浏览器。网页端使用 ECharts 绘制 roll/pitch/yaw 曲线、GPS/ESKF 轨迹、innovation 和 sigma 曲线，"
        "并用 3D 姿态板显示当前姿态。这样答辩时既能展示实时运行，也能保留数据文件作为论文分析来源。"
    )

    add_para(doc, "F.9 后续代码改进路线", "Heading 2")
    add_bullets(doc, [
        "在 ESP32 端增加统一的 sensor_frame 结构，明确时间戳、传感器坐标系、校准状态和有效标志，减少融合时的隐式假设。",
        "把 Mahony/Madgwick 中的 acc/mag 权重改为自适应：当 |a|-1g 或磁场模长残差超限时自动降低观测权重。",
        "在 ESKF 中加入更严格的时间同步：GPS 更新使用最近 IMU 时间戳插值，避免低速轨迹中因时延产生位置残差。",
        "为磁力计标定增加 condition number、coverage score 和鲁棒拟合开关，让报告能量化说明采样质量。",
        "将关键参数集中到 config.py 或 JSON，答辩时可现场展示 beta、Kp/Ki、Q/R 改变后曲线如何变化。",
    ])

    doc.save(DST)
    print(DST)


if __name__ == "__main__":
    main()
