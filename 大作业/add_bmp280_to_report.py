# -*- coding: utf-8 -*-
from pathlib import Path
import csv
import shutil

from docx import Document
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parent
REPORT_IN = ROOT / "多传感器融合扩展板课程论文初稿_补充频率测试结果.docx"
REPORT_OUT = ROOT / "多传感器融合扩展板课程论文初稿_补充BMP280实验结果.docx"
SUMMARY_CSV = ROOT / "data" / "analysis" / "bmp280_summary.csv"
STATIC_CSV = ROOT / "data" / "analysis" / "bmp280_static_stats.csv"


def read_key_value(path):
    rows = {}
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        for row in reader:
            rows[row["item"]] = row
    return rows


def value(rows, key, field="value", as_float=True):
    raw = rows[key][field]
    return float(raw) if as_float else raw


def find_para_by_prefixes(doc, prefixes):
    for prefix in prefixes:
        for p in doc.paragraphs:
            if p.text.strip().startswith(prefix):
                return p
    return doc.paragraphs[-1]


def insert_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def add_table_after(paragraph, rows, headers):
    tbl = paragraph._parent.add_table(rows=1, cols=len(headers), width=Inches(6.5))
    paragraph._p.addnext(tbl._tbl)
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    for row in rows:
        cells = tbl.add_row().cells
        for i, h in enumerate(headers):
            cells[i].text = str(row[h])
    after_p = OxmlElement("w:p")
    tbl._tbl.addnext(after_p)
    return tbl, Paragraph(after_p, paragraph._parent)


def main():
    summary = read_key_value(SUMMARY_CSV)
    static = read_key_value(STATIC_CSV)

    static_file = value(summary, "static_source_file", as_float=False)
    motion_file = value(summary, "motion_source_file", as_float=False)
    chip_id = value(summary, "chip_id", as_float=False)
    static_samples = int(value(summary, "static_samples"))
    motion_samples = int(value(summary, "motion_samples"))
    pressure_std = value(summary, "static_pressure_std")
    temp_std = value(summary, "static_temperature_std")
    alt_std = value(summary, "static_altitude_std")
    motion_alt_min = value(summary, "motion_altitude_min")
    motion_alt_max = value(summary, "motion_altitude_max")
    motion_alt_range = value(summary, "motion_altitude_range")
    motion_pressure_range = value(summary, "motion_pressure_range")

    pressure_mean = value(static, "pressure", field="mean")
    pressure_ptp = value(static, "pressure", field="peak_to_peak")
    temp_mean = value(static, "temperature", field="mean")
    temp_ptp = value(static, "temperature", field="peak_to_peak")
    alt_ptp = value(static, "relative_altitude", field="peak_to_peak")

    work_dir = ROOT / "_doc_work" / "bmp280"
    if work_dir.exists():
        shutil.rmtree(work_dir, ignore_errors=True)
    work_dir.mkdir(parents=True, exist_ok=True)

    try:
        tmp_in = work_dir / "report.docx"
        tmp_out = work_dir / "report_out.docx"
        shutil.copy2(REPORT_IN, tmp_in)
        doc = Document(str(tmp_in))

        ch34 = find_para_by_prefixes(doc, ["3.4", "3、4", "3 4", "3. 4", "3."])
        anchor = insert_after(
            ch34,
            "BMP280 气压计通过 I2C 总线接入系统，地址为 0x76。本实验读取芯片 ID 得到 %s，与 BMP280 常见 ID 0x58 一致，说明板载气压计模块焊接和通信正常。驱动程序读取 0x88 起始的温度、气压校准系数，并按照 BMP280 数据手册中的整数补偿公式计算温度与气压，再由气压相对变化换算相对高度。"
            % chip_id,
            "Normal",
        )

        ch62 = find_para_by_prefixes(doc, ["6.2", "6、2", "6 2", "6. 2", "6."])
        anchor = insert_after(
            ch62,
            "为验证第三类传感器 BMP280 的可用性，本项目补充进行了气压计静止噪声和高度变化实验。静止实验数据文件为 %s，共 %d 个样本；高度变化实验数据文件为 %s，共 %d 个样本。采样过程中 BMP280 通过 I2C 总线读取温度和气压，并以实验开始阶段的平均气压作为参考气压计算相对高度。"
            % (static_file, static_samples, motion_file, motion_samples),
            "Normal",
        )
        anchor = insert_after(
            anchor,
            "静止实验中，平均气压为 %.3f Pa，气压标准差为 %.3f Pa，峰峰值为 %.3f Pa；温度平均值为 %.3f°C，温度标准差为 %.3f°C，峰峰值为 %.3f°C。由气压换算的相对高度标准差为 %.3f m，说明 BMP280 在静止条件下能够提供亚米级高度变化观测，但仍会受到气压噪声和环境扰动影响。"
            % (pressure_mean, pressure_std, pressure_ptp, temp_mean, temp_std, temp_ptp, alt_std),
            "Normal",
        )
        anchor = insert_after(
            anchor,
            "高度变化实验中，相对高度最小值为 %.3f m，最大值为 %.3f m，高度变化范围为 %.3f m；对应气压变化范围为 %.3f Pa。该结果说明 BMP280 能够感知手动抬升/放低或楼层微小变化带来的气压差，可作为系统中的高度变化辅助观测量。"
            % (motion_alt_min, motion_alt_max, motion_alt_range, motion_pressure_range),
            "Normal",
        )

        rows = [
            {"项目": "芯片 ID", "结果": chip_id, "单位": "-", "说明": "BMP280 常见 ID 为 0x58"},
            {"项目": "静止样本数", "结果": str(static_samples), "单位": "rows", "说明": "5 Hz 静止采集"},
            {"项目": "静止气压标准差", "结果": "%.3f" % pressure_std, "单位": "Pa", "说明": "越小越稳定"},
            {"项目": "静止温度标准差", "结果": "%.3f" % temp_std, "单位": "°C", "说明": "温度波动"},
            {"项目": "静止相对高度标准差", "结果": "%.3f" % alt_std, "单位": "m", "说明": "由气压换算"},
            {"项目": "高度变化样本数", "结果": str(motion_samples), "单位": "rows", "说明": "高度变化采集"},
            {"项目": "相对高度变化范围", "结果": "%.3f" % motion_alt_range, "单位": "m", "说明": "高度响应幅度"},
            {"项目": "气压变化范围", "结果": "%.3f" % motion_pressure_range, "单位": "Pa", "说明": "与高度变化对应"},
        ]
        p = insert_after(anchor, "")
        _, anchor = add_table_after(p, rows, ["项目", "结果", "单位", "说明"])

        anchor = insert_after(anchor, "图：BMP280 静止状态下气压、温度与相对高度噪声曲线（对应文件 bmp280_static_pressure_temp.png）", "Normal")
        anchor = insert_after(anchor, "图：BMP280 高度变化实验中的气压与相对高度响应曲线（对应文件 bmp280_height_change.png）", "Normal")

        ch67 = find_para_by_prefixes(doc, ["6.7", "6、7", "6 7", "6. 7", "6.6"])
        insert_after(
            ch67,
            "BMP280 实验进一步证明当前 PCB 扩展板不仅完成 IMU 和磁力计集成，也完成了气压计传感器的硬件接入与数据验证。实测结果中 BMP280 静止气压标准差为 %.3f Pa，相对高度标准差为 %.3f m，高度变化实验可观测约 %.3f m 的相对高度变化，因此该模块可作为后续多传感器高度估计或楼层变化检测的扩展基础。"
            % (pressure_std, alt_std, motion_alt_range),
            "Normal",
        )

        doc.save(str(tmp_out))
        shutil.copy2(tmp_out, REPORT_OUT)
    finally:
        shutil.rmtree(work_dir, ignore_errors=True)

    print("wrote", REPORT_OUT)
    print("pressure_std", pressure_std)
    print("alt_std", alt_std)
    print("motion_alt_range", motion_alt_range)


if __name__ == "__main__":
    main()
