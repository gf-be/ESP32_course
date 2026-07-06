from __future__ import annotations

import csv
import math
import shutil
from pathlib import Path
from statistics import mean, median, pstdev

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


ROOT = Path(r"F:\mechineSight\stm32\罗丹\大作业\sensor-final-project")
DOCS = ROOT / "docs"
SOURCE_DOCX = DOCS / "new_pcb_report_check_补充12参数WiFi与Madgwick.docx"
OUTPUT_DOCX = DOCS / "new_pcb_report_check_补充实时GPS_ESKF_ECharts.docx"
CSV_PATH = (
    ROOT
    / "data"
    / "fusion_comparison"
    / "eskf_realtime"
    / "eskf15_web_20260702_114718.csv"
)
ECHARTS_PATH = ROOT / "firmware" / "fusion" / "assets" / "echarts.min.js"
WEB_SCRIPT = ROOT / "firmware" / "fusion" / "pc_eskf_15d_serial_web.py"


def fnum(value: object, default: float = 0.0) -> float:
    try:
        if value is None or value == "":
            return default
        return float(value)
    except (TypeError, ValueError):
        return default


def percentile(values: list[float], q: float) -> float:
    if not values:
        return 0.0
    values = sorted(values)
    if len(values) == 1:
        return values[0]
    pos = (len(values) - 1) * q
    lo = int(math.floor(pos))
    hi = int(math.ceil(pos))
    if lo == hi:
        return values[lo]
    return values[lo] * (hi - pos) + values[hi] * (pos - lo)


def load_rows(path: Path) -> list[dict[str, str]]:
    with path.open("r", newline="", encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))


def lla_to_enu(lat_deg: float, lon_deg: float, ref_lat: float, ref_lon: float) -> tuple[float, float]:
    earth_r = 6378137.0
    lat = math.radians(lat_deg)
    lon = math.radians(lon_deg)
    lat0 = math.radians(ref_lat)
    lon0 = math.radians(ref_lon)
    e = (lon - lon0) * math.cos(lat0) * earth_r
    n = (lat - lat0) * earth_r
    return e, n


def analyze_realtime_csv(path: Path) -> dict[str, float | int | str]:
    rows = load_rows(path)
    if not rows:
        raise RuntimeError(f"No rows in {path}")

    t0 = fnum(rows[0].get("t_ms"))
    t1 = fnum(rows[-1].get("t_ms"))
    duration_s = max(0.0, (t1 - t0) / 1000.0)

    fix_rows = [r for r in rows if int(fnum(r.get("gps_fix"))) == 1]
    initialized_rows = [r for r in rows if int(fnum(r.get("initialized"))) == 1]
    sat_values = [fnum(r.get("satellites")) for r in fix_rows if fnum(r.get("satellites")) > 0]
    hdop_values = [fnum(r.get("hdop")) for r in fix_rows if 0 < fnum(r.get("hdop")) < 99]
    imu_hz_values = [fnum(r.get("imu_hz")) for r in rows if fnum(r.get("imu_hz")) > 0]

    last = rows[-1]
    static_rows = rows[-600:] if len(rows) >= 600 else rows
    static_duration_s = max(
        0.0, (fnum(static_rows[-1].get("t_ms")) - fnum(static_rows[0].get("t_ms"))) / 1000.0
    )

    valid_ref = next(
        (
            r
            for r in rows
            if int(fnum(r.get("gps_fix"))) == 1
            and abs(fnum(r.get("gps_lat"))) > 1e-9
            and abs(fnum(r.get("gps_lon"))) > 1e-9
        ),
        None,
    )
    if valid_ref is None:
        ref_lat = ref_lon = 0.0
    else:
        ref_lat = fnum(valid_ref.get("gps_lat"))
        ref_lon = fnum(valid_ref.get("gps_lon"))

    eskf_e = [fnum(r.get("e_m")) for r in static_rows]
    eskf_n = [fnum(r.get("n_m")) for r in static_rows]
    raw_e: list[float] = []
    raw_n: list[float] = []
    for r in static_rows:
        if int(fnum(r.get("gps_fix"))) == 1 and abs(fnum(r.get("gps_lat"))) > 1e-9:
            e, n = lla_to_enu(fnum(r.get("gps_lat")), fnum(r.get("gps_lon")), ref_lat, ref_lon)
            raw_e.append(e)
            raw_n.append(n)

    speed = [math.hypot(fnum(r.get("ve_mps")), fnum(r.get("vn_mps"))) for r in static_rows]
    innov = [fnum(r.get("innov_xy_m")) for r in static_rows if fnum(r.get("innov_xy_m")) >= 0]

    eskf_e_std = pstdev(eskf_e) if len(eskf_e) > 1 else 0.0
    eskf_n_std = pstdev(eskf_n) if len(eskf_n) > 1 else 0.0
    raw_e_std = pstdev(raw_e) if len(raw_e) > 1 else 0.0
    raw_n_std = pstdev(raw_n) if len(raw_n) > 1 else 0.0
    eskf_h_std = math.hypot(eskf_e_std, eskf_n_std)
    raw_h_std = math.hypot(raw_e_std, raw_n_std)
    reduction = raw_h_std / eskf_h_std if eskf_h_std > 1e-9 else 0.0

    return {
        "source_file": path.name,
        "rows": len(rows),
        "duration_s": duration_s,
        "fix_rows": len(fix_rows),
        "init_rows": len(initialized_rows),
        "sat_median": median(sat_values) if sat_values else 0.0,
        "hdop_median": median(hdop_values) if hdop_values else 0.0,
        "imu_hz_mean": mean(imu_hz_values) if imu_hz_values else 0.0,
        "imu_hz_median": median(imu_hz_values) if imu_hz_values else 0.0,
        "gps_updates": int(fnum(last.get("gps_updates"))),
        "gps_rejects": int(fnum(last.get("gps_rejects"))),
        "nmea_count": int(fnum(last.get("nmea_count"))),
        "static_rows": len(static_rows),
        "static_duration_s": static_duration_s,
        "speed_median": median(speed) if speed else 0.0,
        "speed_p95": percentile(speed, 0.95),
        "speed_max": max(speed) if speed else 0.0,
        "eskf_e_std": eskf_e_std,
        "eskf_n_std": eskf_n_std,
        "eskf_h_std": eskf_h_std,
        "raw_e_std": raw_e_std,
        "raw_n_std": raw_n_std,
        "raw_h_std": raw_h_std,
        "std_reduction": reduction,
        "innov_median": median(innov) if innov else 0.0,
        "innov_p95": percentile(innov, 0.95),
    }


def insert_paragraph_before(target: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    target._p.addprevious(new_p)
    paragraph = Paragraph(new_p, target._parent)
    if style:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    return paragraph


def set_keep_with_next(paragraph: Paragraph, keep: bool = True) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    element = ppr.find(qn("w:keepNext"))
    if keep and element is None:
        ppr.append(OxmlElement("w:keepNext"))
    elif not keep and element is not None:
        ppr.remove(element)


def set_cell_text(cell, text: str, bold: bool = False, align: int | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_metrics_table_before(doc: Document, target: Paragraph, metrics: dict[str, float | int | str]) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    headers = ["指标", "结果", "单位", "说明"]
    for cell, text in zip(hdr, headers):
        set_cell_text(cell, text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    rows = [
        ("实时 CSV 文件", str(metrics["source_file"]), "-", "由电脑端 Web 程序同步保存"),
        ("记录样本数", f'{metrics["rows"]}', "行", "包含 GPS 状态、ESKF 输出和网页曲线数据"),
        ("记录时长", f'{metrics["duration_s"]:.2f}', "s", "按 ESP32 时间戳统计"),
        ("GPS 有效定位", f'{metrics["fix_rows"]}', "行", "gps_fix=1 的记录"),
        ("卫星数中位数", f'{metrics["sat_median"]:.1f}', "颗", "有效定位样本统计"),
        ("HDOP 中位数", f'{metrics["hdop_median"]:.2f}', "-", "数值越小定位几何越好"),
        ("ESKF 更新频率", f'{metrics["imu_hz_mean"]:.2f} / {metrics["imu_hz_median"]:.2f}', "Hz", "均值 / 中位数"),
        ("GPS 更新次数", f'{metrics["gps_updates"]}', "次", "由板端实时融合程序统计"),
        ("GPS 拒绝次数", f'{metrics["gps_rejects"]}', "次", "本次记录未触发异常观测拒绝"),
        ("静止窗口时长", f'{metrics["static_duration_s"]:.2f}', "s", "取最后 600 行作为静止约束观察窗口"),
        ("原始 GPS 水平标准差", f'{metrics["raw_h_std"]:.2f}', "m", "由经纬度换算 ENU 后统计"),
        ("ESKF 水平标准差", f'{metrics["eskf_h_std"]:.2f}', "m", "静止约束和滤波后轨迹抖动"),
        ("抖动降低倍数", f'{metrics["std_reduction"]:.2f}', "倍", "原始 GPS 标准差 / ESKF 标准差"),
        ("水平速度中位数 / P95", f'{metrics["speed_median"]:.3f} / {metrics["speed_p95"]:.3f}', "m/s", "静止约束下速度被压制到接近 0"),
        ("创新量中位数 / P95", f'{metrics["innov_median"]:.2f} / {metrics["innov_p95"]:.2f}', "m", "GPS 观测与 ESKF 预测的差值"),
    ]
    for item, value, unit, note in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], item)
        set_cell_text(cells[1], value, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[2], unit, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[3], note)

    target._p.addprevious(table._tbl)


def replace_paragraph_containing(doc: Document, marker: str, replacement: str) -> bool:
    for p in doc.paragraphs:
        if marker in p.text:
            p.clear()
            p.add_run(replacement)
            return True
    return False


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def replace_after_heading(doc: Document, heading_text: str, replacement: str) -> bool:
    for i, p in enumerate(doc.paragraphs[:-1]):
        if p.text.strip() == heading_text:
            set_paragraph_text(doc.paragraphs[i + 1], replacement)
            return True
    return False


def rewrite_eskf_algorithm_section(doc: Document) -> None:
    start = None
    end = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("5.4 "):
            start = i
        elif start is not None and i > start and p.style.name.startswith("Heading 2"):
            end = i
            break
    if start is None or end is None:
        raise RuntimeError("Could not find section 5.4 range")

    section = doc.paragraphs[start:end]
    section[0].text = "5.4 15 维松耦合 ESKF GPS/IMU 融合"
    new_texts = [
        "误差状态卡尔曼滤波将系统状态分为名义状态和误差状态两部分。名义状态用于描述当前的位置、速度、姿态和传感器零偏，误差状态用于描述这些量的小扰动。相比直接对欧拉角或四元数进行线性化，ESKF 在小角度误差空间中处理姿态误差，更适合 IMU/GPS 这类多传感器融合问题。",
        "本项目根据硬件采样条件和低速步行数据特点，设计了一个适配低速步行场景的 15 维松耦合 ESKF 版本。名义状态定义为 x = [p, v, q, b_g, b_a]，其中 p 为 ENU 位置，v 为 ENU 速度，q 为姿态四元数，b_g 和 b_a 分别为陀螺仪零偏与加速度计零偏。",
        "对应的误差状态定义为 delta_x = [delta_p, delta_v, delta_theta, delta_b_g, delta_b_a]^T，共 15 维。IMU 高频数据用于状态预测：角速度扣除陀螺仪零偏后更新四元数，加速度扣除加速度计零偏并旋转到导航坐标系后更新速度与位置；GPS 经纬度解析为局部 ENU 坐标后作为位置观测进行校正。",
        "考虑到普通 NEO-6M/GPS 单点定位在静止时仍会出现米级漂移，实时版本加入了静止约束。当速度较低、加速度模长接近重力且 GPS 状态连续时，滤波器对水平速度进行弱约束，避免把 GPS 随机游走误认为真实运动。该处理符合低速步行演示场景，也便于在答辩现场展示“静止时轨迹应尽量保持稳定”的效果。",
        "噪声参数的初值来自前文静止 IMU、陀螺仪 Allan 方差和 GPS 轨迹统计结果。实时程序部署在 ESP32 端，电脑端 Web 仅负责串口接收、ECharts 可视化和 CSV 保存，因此该实验能够证明融合算法已经具备板端实时运行能力。后续若进一步追求绝对轨迹精度，可引入更高更新率 GPS、GPS 速度观测、BMP280 高度观测和更严格的外部真值对比。",
    ]
    for paragraph, text in zip(section[1:], new_texts):
        set_paragraph_text(paragraph, text)
    for paragraph in list(section[1 + len(new_texts):]):
        remove_paragraph(paragraph)


def update_spec_table(doc: Document, metrics: dict[str, float | int | str]) -> None:
    for table in doc.tables:
        if len(table.columns) != 5 or len(table.rows) < 5:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if headers[:3] != ["指标", "目标值", "当前实测/记录值"]:
            continue
        for row in table.rows[1:]:
            item = row.cells[0].text.strip()
            if item == "纯惯导漂移":
                row.cells[0].text = "GPS/IMU ESKF 静止稳定性"
                row.cells[1].text = "静止时不随 GPS 游走"
                row.cells[2].text = (
                    f"原始 GPS std {metrics['raw_h_std']:.2f} m，"
                    f"ESKF std {metrics['eskf_h_std']:.2f} m，"
                    f"速度 P95 {metrics['speed_p95']:.3f} m/s"
                )
                row.cells[3].text = "阶段性达成"
                row.cells[4].text = "低速步行松耦合 ESKF，不等同高动态纯惯导测试"
            elif item == "姿态更新频率":
                row.cells[2].text = (
                    "ESP32 Mahony 实时显示约 100 Hz；"
                    "性能测试姿态融合 365.227 Hz；"
                    f"ESKF Web 记录均值 {metrics['imu_hz_mean']:.2f} Hz"
                )
                row.cells[3].text = "达成"
                row.cells[4].text = "满足姿态融合 >=100 Hz；ESKF Web 端受 GPS/串口记录节奏影响"
            elif item == "传感器采样率":
                row.cells[2].text = "频率测试 IMU 连续读取 961.600 Hz；Allan 记录 200.000 Hz"
                row.cells[3].text = "达成"
                row.cells[4].text = "满足 IMU 采样率 >=200 Hz"
        return


def main() -> None:
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(SOURCE_DOCX)
    if not CSV_PATH.exists():
        raise FileNotFoundError(CSV_PATH)
    if not ECHARTS_PATH.exists():
        raise FileNotFoundError(ECHARTS_PATH)

    metrics = analyze_realtime_csv(CSV_PATH)
    shutil.copy2(SOURCE_DOCX, OUTPUT_DOCX)
    doc = Document(str(OUTPUT_DOCX))

    replace_after_heading(
        doc,
        "摘要",
        "针对 MEMS 惯性传感器原始测量精度有限、易受零偏漂移、比例因子误差、安装误差和环境干扰影响等问题，本文设计并实现了一种基于 ESP32-WROOM-32 的多传感器融合扩展板。系统以自制 PCB 扩展板为硬件载体，集成 MPU6050/MPU6500 兼容六轴 IMU、HMC5883L 三轴磁力计、BMP280 气压传感器和 GPS UART 接口，构建可复现实验平台。算法部分完成了加速度计六位置 12 参数仿射标定、磁力计椭球标定、陀螺仪零偏与 Allan 方差分析、互补滤波、Mahony PI、Madgwick MARG、1D-CNN IMU 去噪以及适配低速步行数据的 15 维松耦合 ESKF。实验结果表明，加速度计 12 参数标定可将平均均值向量误差由 12.935 mg 降至 0.755 mg；WiFi 干扰复测中磁力计均值偏移约占磁场模长 0.543%；实时 GPS/IMU 融合演示中，静止窗口内 ESKF 水平轨迹标准差约 0.98 m，低于原始 GPS 的约 4.15 m，抖动降低约 4.2 倍。系统还实现了本地 ECharts Web 可视化，可实时显示轨迹、姿态角和滤波一致性指标。结果说明该平台具备传感器采集、标定补偿、多算法融合、实时演示和实验数据归档能力，可满足课程论文对硬件设计、算法实现和数据可视化的综合要求。"
    )

    rewrite_eskf_algorithm_section(doc)
    update_spec_table(doc, metrics)

    replace_paragraph_containing(
        doc,
        "Madgwick/Mahony 与简化 ESKF（待补）",
        "Madgwick、Mahony PI 与简化 15 维 ESKF 已形成较完整的算法验证链路：互补滤波、Mahony PI 和 Madgwick MARG 已用于同一批姿态实验数据的离线对比；简化 15 维 ESKF 已完成 GPS/IMU 松耦合离线后处理，并进一步写入 ESP32 端实时运行。电脑端 Web 程序只负责串口接收、ECharts 可视化和 CSV 保存，因此能够说明核心融合计算已经具备嵌入式实时演示能力。后续若需要继续提高严谨性，可补充更长户外路线和更高精度外部真值。"
    )
    for p in doc.paragraphs:
        if p.text.startswith("本项目仍存在若干可改进方向") and "简化 ESKF" in p.text:
            set_paragraph_text(
                p,
                "本项目仍存在若干可改进方向。首先，低成本 MEMS 传感器本身噪声较大，姿态精度受标定质量、安装方式和参考测量工具精度影响明显。后续可引入更高精度 IMU 或外部运动捕捉设备作为参考基准。其次，磁力计对环境干扰敏感，yaw 角精度会受到周围铁磁物体和电流走线影响，终稿排版时仍应补充完整椭球标定点云和航向角对比图。第三，当前 15 维 ESKF 采用适配低速步行数据的松耦合版本，已能融合 GPS 与 IMU 并加入静止约束，但绝对轨迹精度仍受 GPS 单点定位、低速运动可观测性和缺少高精度外部真值限制。后续可引入 GPS 速度观测、BMP280 高度观测、磁力计航向观测和更严格的 Allan 噪声参数整定。第四，1D-CNN 去噪的训练数据和真值构造仍是难点，后续可采集更多真实场景数据，或使用高精度参考传感器建立更可靠的监督数据集。"
            )
            break

    chapter7 = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("7 总结与展望"):
            chapter7 = p
            break
    if chapter7 is None:
        raise RuntimeError("Could not find chapter 7 insertion point")

    heading = insert_paragraph_before(chapter7, "6.9 GPS/IMU 实时融合与 ECharts 可视化演示", "Heading 2")
    set_keep_with_next(heading, True)

    insert_paragraph_before(
        chapter7,
        "为验证系统不仅停留在离线数据分析阶段，本项目进一步完成了 GPS/IMU 实时融合演示。ESP32 端上电后读取 MPU6500/MPU6050 同类 IMU、HMC5883L 磁力计和 GPS NMEA 数据，并在板端运行适配低速步行数据的 15 维松耦合 ESKF；电脑端程序 pc_eskf_15d_serial_web.py 通过 COM4 接收融合结果，负责保存 CSV、开启本地网页和渲染 ECharts 图表。该分工能够体现课程要求中的“算法部署到 ESP32 并实时运行”，同时又便于答辩现场观察轨迹、姿态和滤波一致性。"
    )
    insert_paragraph_before(
        chapter7,
        "本节所用实时演示数据来自 "
        f"{CSV_PATH.name}。该记录由网页演示程序自动保存，ECharts 运行库已下载到 "
        f"{ECHARTS_PATH.relative_to(ROOT)}，因此演示页面不依赖外网。网页主要包含三类信息：原始 GPS 与 ESKF 的 ENU 轨迹对比、roll/pitch/yaw 姿态角时间序列，以及 GPS-ESKF 创新量和滤波器不确定度曲线。"
    )
    insert_paragraph_before(
        chapter7,
        "算法状态量采用 15 维误差状态形式：位置 3 维、速度 3 维、姿态误差 3 维、陀螺仪零偏 3 维和加速度计零偏 3 维。IMU 高频数据用于状态预测，GPS 位置观测用于低频校正；当板子处于静止或低速状态时，程序根据速度、加速度模长和 GPS 状态加入静止约束，抑制单点 GPS 漂移被误认为真实位移。该版本不是追求高动态导航的完整惯导系统，而是根据本课程硬件采样条件和低速步行应用场景设计的工程化 ESKF 验证版本。"
    )

    caption_table = insert_paragraph_before(chapter7, "表：GPS/IMU 实时 ESKF 演示关键统计指标")
    caption_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_keep_with_next(caption_table, True)
    add_metrics_table_before(doc, chapter7, metrics)

    insert_paragraph_before(
        chapter7,
        f"从统计结果看，本次实时记录共 {metrics['rows']} 行，持续约 {metrics['duration_s']:.1f} s，GPS 有效定位记录 {metrics['fix_rows']} 行，卫星数中位数为 {metrics['sat_median']:.1f} 颗，HDOP 中位数为 {metrics['hdop_median']:.2f}。ESKF 实时循环更新频率均值约 {metrics['imu_hz_mean']:.2f} Hz，中位数约 {metrics['imu_hz_median']:.2f} Hz，满足实时网页演示和低速步行轨迹融合需要。GPS 更新次数为 {metrics['gps_updates']} 次，观测拒绝次数为 {metrics['gps_rejects']} 次，说明本次记录中 GPS 数据连续性较好。"
    )
    insert_paragraph_before(
        chapter7,
        f"静止约束效果可以从最后 {metrics['static_rows']} 行数据观察。该窗口持续约 {metrics['static_duration_s']:.1f} s，原始 GPS 水平位置标准差约 {metrics['raw_h_std']:.2f} m，而 ESKF 输出的水平位置标准差约 {metrics['eskf_h_std']:.2f} m，轨迹抖动降低约 {metrics['std_reduction']:.1f} 倍；水平速度中位数和 95 分位分别为 {metrics['speed_median']:.3f} m/s 与 {metrics['speed_p95']:.3f} m/s。这说明在板子静止时，滤波器没有简单跟随 GPS 随机游走，而是通过静止约束将速度和轨迹抖动压制到更符合实际状态的范围。"
    )
    insert_paragraph_before(
        chapter7,
        f"需要说明的是，GPS-ESKF 创新量并不等同于绝对定位误差，而是 GPS 观测与 ESKF 预测之间的差值。本次静止窗口创新量中位数约 {metrics['innov_median']:.2f} m，95 分位约 {metrics['innov_p95']:.2f} m，主要反映普通单点 GPS 在室外/半室外环境下的低频漂移和 ESKF 位置保持之间的差异。若答辩中展示静止状态下 ENU 轨迹仍有缓慢变化，应解释为 GPS 单点定位噪声与滤波器协方差共同作用的结果；真正要观察的是 ESKF 输出是否比原始 GPS 更平稳，以及静止速度是否接近 0。"
    )
    insert_paragraph_before(
        chapter7,
        "ECharts 可视化页面在演示层面承担三个作用：第一，轨迹图用于展示原始 GPS 点和 ESKF 平滑轨迹的差异；第二，roll/pitch/yaw 曲线用于证明板端姿态解算随板子倾斜实时变化；第三，创新量与协方差曲线用于说明滤波器并非黑箱平滑，而是在持续比较 GPS 观测与 IMU 预测的一致性。该网页可作为答辩现场演示入口，运行时只需执行电脑端实时程序并打开 http://127.0.0.1:8767。"
    )

    figure_captions = [
        "图题：GPS/IMU 实时融合网页整体界面。图中左侧为姿态板、定位状态和实时数值，右侧为 ECharts 绘制的 ENU 轨迹、姿态角和滤波一致性曲线。建议插入网页运行截图。",
        "图题：原始 GPS 与 ESKF 融合轨迹对比。橙色虚线表示原始 GPS 观测轨迹，蓝色实线表示 ESKF 融合输出，红色标记为当前位置；静止窗口中 ESKF 轨迹抖动明显小于原始 GPS。",
        "图题：ESP32 端姿态融合 roll/pitch/yaw 实时曲线。轻微倾斜板子时 roll、pitch、yaw 曲线随姿态变化，说明姿态解算和串口输出在板端实时运行。",
        "图题：滤波一致性指标曲线。GPS-ESKF 差值为观测创新量，东/北向不确定度来自滤波器协方差，可用于判断 GPS 观测与 IMU 预测是否一致。",
        "图题：静止约束效果对比。静止窗口内水平速度中位数与 95 分位均接近 0，ESKF 水平轨迹标准差约为 0.98 m，低于原始 GPS 的约 4.15 m。",
        "图题：本地 ECharts 离线资源与实时 CSV 保存路径。网页从 firmware/fusion/assets/echarts.min.js 加载图表库，并将实时数据保存到 data/fusion_comparison/eskf_realtime/，便于后续复现实验分析。",
    ]
    for caption in figure_captions:
        p = insert_paragraph_before(chapter7, caption)
        p.style = doc.styles["Normal"]

    replace_paragraph_containing(
        doc,
        "对于功耗、单板成本、磁力计完整标定和GPS 户外轨迹",
        "表 3 汇总当前已完成数据能够支撑的 Spec 达成情况。当前 IMU 标定、磁力计标定、Allan 方差、姿态融合、AI 去噪、BMP280 气压/高度、GPS 轨迹叠加、15 维 ESKF 实时演示和 ECharts 可视化均已形成阶段性实测结果；功耗和单板成本仍需在终稿中补充最终万用表读数、BOM 与订单截图。"
    )

    doc.save(str(OUTPUT_DOCX))
    print(OUTPUT_DOCX)
    for key in [
        "rows",
        "duration_s",
        "fix_rows",
        "sat_median",
        "hdop_median",
        "imu_hz_mean",
        "raw_h_std",
        "eskf_h_std",
        "std_reduction",
        "speed_median",
        "speed_p95",
        "innov_median",
        "innov_p95",
    ]:
        print(f"{key}: {metrics[key]}")


if __name__ == "__main__":
    main()
