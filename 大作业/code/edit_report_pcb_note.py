# -*- coding: utf-8 -*-
from pathlib import Path
import sys

from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def insert_after(paragraph, text, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    new_para.add_run(text)
    if style == "Normal" or style is None:
        new_para.paragraph_format.first_line_indent = Pt(21)
        new_para.paragraph_format.line_spacing = 1.25
    return new_para


def find_heading(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise ValueError("heading not found: " + prefix)


def main():
    if len(sys.argv) != 3:
        raise SystemExit("usage: edit_report_pcb_note.py input.docx output.docx")

    infile = Path(sys.argv[1])
    outfile = Path(sys.argv[2])
    doc = Document(str(infile))

    anchor = find_heading(doc, "2.5")
    items = [
        (
            "2.6 PCB 实物设置差异与可实验性验证",
            "Heading 2",
        ),
        (
            "本次扩展板实物调试中，PCB 板载器件和接口实现与课程给出的推荐报告模板或原始设计设想并非完全一致。实际板卡采用 ESP32 的 GPIO21/GPIO22 作为共享 I2C 总线，传感器地址扫描结果为 0x68、0x1E 和 0x76；其中 0x68 对应 MPU 系列六轴 IMU，WHO_AM_I 读数为 0x70，说明实际器件可能为 MPU6050 兼容型号或同系列 IMU；0x1E 对应 GY-273/HMC5883L 磁力计模块；0x76 的芯片 ID 为 0x58，可确认为 BMP280 气压计。上述差异不会改变本课程要求的核心实验对象，即 IMU、磁力计和气压计的多传感器采集、标定与融合验证。",
            "Normal",
        ),
        (
            "为验证该 PCB 仍具备可实验性，本文补充了硬件连通性和传感器可用性测试。I2C 扫描能够同时发现 0x68、0x1E、0x76 三个设备；传感器冒烟测试中，加速度计、陀螺仪、磁力计数据均能随板子姿态变化产生合理响应，BMP280 也能返回有效芯片 ID。随后完成了 120 s IMU 静止采样，得到 6000 组数据，采样频率为 50.000 Hz，加速度模长均值为 1.025397 g，模长标准差为 4.275 mg；陀螺仪 X/Y/Z 三轴零偏分别为 0.221958 deg/s、0.982664 deg/s 和 -0.123877 deg/s。这说明当前板卡能够稳定输出可用于零偏估计、静态噪声分析和后续标定的数据。",
            "Normal",
        ),
        (
            "针对磁力计可能受 ESP32 和周边器件干扰的问题，本文进行了 90 s 三维旋转采样和 WiFi 开关对比测试。旋转采样共获得 1800 组数据，估计硬铁偏置为 (63.0, -104.5, 32.5) raw count，三轴半径不均衡比为 1.139，中心化后半径变异系数为 0.0736，说明磁场点云整体覆盖较均衡，具备进一步椭球标定条件。WiFi 干扰对比中，WiFi 开启并扫描后磁场均值合成偏移约为 15.50 raw count，占关闭状态平均磁场模长的 3.23%；同时三轴标准差未增大，反而下降至关闭状态的 0.15-0.20 倍。因此，本次测试未观察到 WiFi 活动导致的明显随机噪声增强，GY-273 在当前安装位置受 ESP32 WiFi 射频活动的直接干扰不明显。",
            "Normal",
        ),
        (
            "综上，虽然 PCB 实物设置与模板描述或原计划存在器件型号、模块形式和局部布局差异，但从地址扫描、寄存器识别、动态读数、静态采样和磁干扰对比结果看，当前板卡已经满足课程大作业所需的实验闭环：能够完成多传感器读取、静态误差统计、加速度计标定、磁力计标定、气压扩展验证以及后续姿态融合算法对比。后续报告中将把该差异作为工程调试事实说明，并以实测数据支撑系统可用性，而不将模板设计与实际实物混同。",
            "Normal",
        ),
    ]
    for text, style in items:
        anchor = insert_after(anchor, text, style)

    anchor = find_heading(doc, "6.8")
    insert_after(
        anchor,
        "硬件实物与模板设计存在差异是本项目调试中的主要工程事实之一。本文不将该差异简单视为失败，而是通过 I2C 扫描、传感器冒烟测试、IMU 静止采样、磁力计旋转采样和 WiFi 干扰对比逐项验证实物板卡的实验可用性。结果表明，当前 PCB 虽与推荐模板不完全一致，但能够稳定支撑课程要求的传感器驱动、标定和融合算法验证。后续改进可进一步统一器件丝印、原理图标注和报告描述，并使用固定夹具重复 WiFi 干扰测试以降低姿态微小变化带来的均值偏移。",
        "Normal",
    )

    doc.save(str(outfile))
    print(outfile)


if __name__ == "__main__":
    main()
