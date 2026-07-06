from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parent
PROJECT = ROOT / "sensor-final-project"
DOCS = PROJECT / "docs"
DATA = PROJECT / "data" / "new_pcb_bringup"

INPUT = DOCS / "多传感器融合扩展板课程论文初稿1_口径统一占位保留.docx"
OUTPUT = DOCS / "多传感器融合扩展板课程论文初稿1_补充新PCB硬件验证.docx"

SUMMARY = DATA / "new_pcb_bringup_summary_COM4_20260629_141804.csv"
I2C_LOG = DATA / "i2c_sensor_smoke_COM4_20260629_141804.txt"
GPS_LOG = DATA / "gps_uart_smoke_COM4_20260629_141804.txt"


def read_summary(path):
    rows = {}
    for line in path.read_text(encoding="utf-8").splitlines()[1:]:
        if "," in line:
            k, v = line.split(",", 1)
            rows[k] = v
    return rows


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def add_table_after(paragraph, doc, rows):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, value in enumerate(rows[0]):
        hdr[i].text = value
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    paragraph._p.addnext(table._tbl)
    return table


def replace_gps_placeholder(doc):
    old_markers = [
        "本项目保留 GPS UART 接口，但由于前期 GNSS 实验已经验证 NMEA 解析和轨迹记录流程",
        "本项目目前 GPS 接口主要作为挑战算法和系统扩展预留",
    ]
    for para in doc.paragraphs:
        text = para.text
        if any(marker in text for marker in old_markers):
            para.text = (
                "本项目在新版 PCB 上对 GPS6MV2/NEO-6M 兼容模块进行了 UART 冒烟测试。"
                "ESP32 采用 UART2，GPS TX 接 GPIO16，GPS RX 接 GPIO17，波特率 9600。"
                "实测能够连续接收 RMC、GGA、GSA、GSV、GLL 等 NMEA 语句，其中 RMC 语句状态位为 A，"
                "GGA 语句 fix quality=1，说明 GPS 串口通信和单点定位均已正常。"
                "因此，GPS 模块已由原先的接口预留状态更新为可用于后续 GPS 轨迹叠加和 GPS/IMU 融合实验的数据来源。"
            )


def find_insert_anchor(doc):
    # Prefer placing the new subsection right after "6.1 实验平台..." if present.
    for para in doc.paragraphs:
        if para.text.strip().startswith("6.1"):
            return para
    for para in doc.paragraphs:
        if "实验平台" in para.text:
            return para
    return doc.paragraphs[-1]


def main():
    s = read_summary(SUMMARY)
    doc = Document(str(INPUT))
    replace_gps_placeholder(doc)

    anchor = find_insert_anchor(doc)
    p = insert_paragraph_after(anchor, "6.1.1 新版 PCB 硬件可用性复测", style=anchor.style)
    p = insert_paragraph_after(
        p,
        (
            "为保证最终提交硬件与实验数据之间的口径一致性，本文在更换新版 PCB 后重新进行了硬件冒烟测试。"
            "新版 PCB 与前一版相比主要调整了板子尺寸和模块空间位置，核心传感器型号、ESP32 引脚连接、"
            "I2C/UART 通信方式和采样程序保持一致。因此，前期完成的 IMU 静态噪声、六位置标定、"
            "磁力计椭球标定、陀螺仪 Allan 方差、姿态融合与 AI 去噪实验数据仍可作为同一系统方案的算法验证依据。"
            "同时，新版 PCB 需要通过独立冒烟测试证明最终硬件本身可正常读取各模块数据。"
        ),
    )
    p = insert_paragraph_after(
        p,
        (
            "2026 年 6 月 29 日，ESP32 通过 COM4 连接电脑，分别运行 I2C 传感器冒烟测试和 GPS UART 冒烟测试。"
            "I2C 扫描结果为 0x1E、0x68、0x76，分别对应 HMC5883L 磁力计、MPU6050/MPU6500 兼容 IMU 和 BMP280 气压计。"
            "IMU 的 WHO_AM_I 返回 0x70，符合 MPU6500/兼容器件特征；BMP280 芯片 ID 返回 0x58。"
            "实时数据中加速度、陀螺仪、温度和磁力计读数稳定输出，说明 I2C 总线和三类传感器均可用。"
        ),
    )
    p = insert_paragraph_after(
        p,
        (
            "GPS 冒烟测试采用 UART2，RX=GPIO16，TX=GPIO17，波特率 9600。"
            "60 s 测试期间共接收 "
            + s.get("gps_nmea_count", "")
            + " 条 NMEA 语句，其中有效定位计数 "
            + s.get("gps_fix_count", "")
            + " 次；GGA 语句显示 fix quality="
            + s.get("gps_fix_quality", "")
            + "、卫星数 "
            + s.get("gps_satellites", "")
            + "、HDOP="
            + s.get("gps_hdop", "")
            + "、海拔约 "
            + s.get("gps_alt_m", "")
            + " m。该结果说明新版 PCB 的 GPS 供电、串口连接和定位功能均正常，可支撑后续 GPS 轨迹叠加与 GPS/IMU 融合实验。"
        ),
    )

    table_rows = [
        ["测试项", "实测结果", "结论"],
        ["I2C 地址扫描", "0x1E, 0x68, 0x76", "HMC5883L、IMU、BMP280 均在线"],
        ["IMU 标识", "WHO_AM_I = " + s.get("imu_who_am_i", ""), "MPU6050/MPU6500 兼容 IMU 可读取"],
        ["BMP280 标识", "chip ID = " + s.get("bmp280_chip_id", ""), "BMP280 可读取"],
        ["I2C 实时样本", s.get("i2c_live_samples", "") + " 组 ACC/GYRO/TEMP/MAG", "动态数据输出正常"],
        ["GPS UART", "NMEA = " + s.get("gps_nmea_count", "") + ", fix = " + s.get("gps_fix_count", ""), "GPS 串口与定位正常"],
        ["GPS 定位质量", "fix quality=" + s.get("gps_fix_quality", "") + ", satellites=" + s.get("gps_satellites", "") + ", HDOP=" + s.get("gps_hdop", ""), "具备轨迹实验条件"],
    ]
    add_table_after(p, doc, table_rows)

    p2 = doc.add_paragraph()
    p2.add_run(
        "综上，新版 PCB 已完成最终硬件冒烟验证。由于新旧 PCB 的核心模块、接口定义和采样程序保持一致，"
        "本文沿用前期标定与融合数据作为算法性能分析来源，并使用新版 PCB 的 I2C/GPS 冒烟测试结果证明最终硬件平台具备可实验性。"
    )

    doc.save(str(OUTPUT))
    print("Wrote", OUTPUT)
    print("Source logs:")
    print(I2C_LOG)
    print(GPS_LOG)


if __name__ == "__main__":
    main()
