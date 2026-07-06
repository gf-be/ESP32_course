# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent / "sensor-final-project"
DOC_DIR = ROOT / "docs"
SRC = DOC_DIR / "new_pcb_report_check.docx"
OUT = DOC_DIR / "new_pcb_report_check_补充12参数加速度计与WiFi干扰.docx"


ACCEL_SETUP = (
    "为提高重力加速度标定精度，本文在六位置实验基础上进一步建立 12 参数仿射误差模型。"
    "实验中板子按 +X、-X、+Y、-Y、+Z、-Z 六个姿态静置，每个姿态采集 1500 组样本，"
    "原始数据文件分别为 pos_x_up_20260616_191706.csv、neg_x_up_20260616_191706.csv、"
    "pos_y_up_20260616_191706.csv、neg_y_up_20260616_191706.csv、"
    "pos_z_up_20260616_191706.csv 和 neg_z_up_20260616_191706.csv。"
)

ACCEL_RESULT = (
    "六位置标定结果表明：未标定时六个姿态的平均均值向量模长误差为 12.935 mg，"
    "仅考虑三轴零偏和比例因子的 6 参数模型可将其降至 3.691 mg；进一步引入轴间耦合和安装不正交误差后，"
    "12 参数仿射模型可将平均均值向量模长误差降至 0.755 mg，最大均值向量模长误差由 24.704 mg 降至 1.812 mg。"
    "相对于原始数据，12 参数模型的平均误差降低倍数约为 17.122 倍。"
)

ACCEL_MODEL = (
    "本次 12 参数模型采用 raw = M * true + b 的形式，其中 b=(0.0104550107, 0.0087837300, 0.0051916032) g，"
    "M 的三行分别为 [0.9966041110, -0.0796001057, -0.0779581687]、"
    "[0.0697239533, 0.9966398483, -0.0058036150]、"
    "[0.0544679540, -0.0059822600, 1.0172163930]。固件端使用反算形式 calibrated = C * raw + d，"
    "其中 C 的三行分别为 [0.9936363144, 0.0798200426, 0.0766064259]、"
    "[-0.0698260446, 0.9977966301, 0.0003414386]、"
    "[-0.0536159808, 0.0015940015, 0.9789750285]，"
    "d=(-0.0114873061, -0.0080361168, -0.0045358955) g。"
)

WIFI_RESULT = (
    "针对磁力计可能受 ESP32 与 WiFi 工作状态影响的问题，本文重新进行了 WiFi 关闭/开启扫描对比测试。"
    "新数据采集于 2026 年 6 月 30 日，数据文件为 mag_wifi_off_20260630_094409.csv 和 "
    "mag_wifi_on_scan_20260630_094409.csv。两组数据均采集 1200 点，持续 59.950 s，采样率为 20 Hz；"
    "WiFi 开启组在采样期间完成 8 次扫描。结果显示，WiFi 开启后三轴均值变化分别为 "
    "Mx +1.017 raw count、My -0.442 raw count、Mz -1.574 raw count，磁场模长均值变化为 -0.569 raw count；"
    "三轴均值偏移量为 1.926 raw count，仅占 WiFi 关闭状态磁场模长均值的 0.543%。"
    "去除前 5 s 启动瞬态后重复统计，偏移比例约为 0.581%，结论保持一致。"
)

WIFI_CONCLUSION = (
    "因此，本次新 PCB/新布置下 WiFi 工作状态对 GY-273/HMC5883L 磁力计的静态磁场测量影响较小，"
    "未观察到会破坏航向角估计的显著磁干扰。报告中将该结果作为 PCB 磁力计布局可用性的实测依据，"
    "后续姿态融合中仍应避免靠近铁磁物体、电机和大电流导线。"
)

SUMMARY_ACCEL = (
    "实验结果表明，加速度计六位置 12 参数仿射标定能够显著降低模长误差，"
    "平均均值向量误差由 12.935 mg 降至 0.755 mg，最大均值向量误差由 24.704 mg 降至 1.812 mg；"
    "WiFi 干扰复测中，磁力计三轴均值偏移仅为 1.926 raw count，占磁场模长 0.543%，说明当前 PCB 上磁力计位置基本满足后续航向融合实验需求。"
)


def replace_matching_paragraphs(doc):
    replacements = [
        ("针对磁力计可能受 ESP32 和周边器件干扰的问题", WIFI_RESULT),
        ("为提高重力加速度标定精度，我们新增执行", ACCEL_SETUP),
        ("六位置结果：校准前平均模长误差为", ACCEL_RESULT),
        ("六位置标定每个姿态采集 1200 组静态数据", ACCEL_SETUP),
        ("加速度计标定得到偏置向量 c=", ACCEL_MODEL),
        ("标定后六个方向的平均模长误差由 10.401 mg", ACCEL_RESULT),
        ("六位置标定已完成。实验数据表明", ACCEL_RESULT),
        ("加速度计六位置标定使用 accel_6pos_means_final_group.csv", ACCEL_MODEL),
        ("实验结果表明，加速度计六位置标定能显著降低模长误差", SUMMARY_ACCEL),
    ]
    for para in doc.paragraphs:
        text = para.text.strip()
        for key, value in replacements:
            if text.startswith(key) or key in text:
                para.text = value
                break


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    return table


def append_update_section(doc):
    doc.add_page_break()
    doc.add_heading("补充实验结果：12 参数加速度计标定与 WiFi 磁干扰复测", level=1)
    doc.add_paragraph(ACCEL_SETUP)
    doc.add_paragraph(ACCEL_RESULT)
    doc.add_paragraph(ACCEL_MODEL)
    add_table(
        doc,
        ["模型", "平均均值向量误差/mg", "最大均值向量误差/mg", "样本 MAE/mg", "相对原始提升倍数"],
        [
            ["原始数据", "12.935", "24.704", "13.421", "1.000"],
            ["6 参数零偏/比例因子", "3.691", "5.272", "4.891", "3.505"],
            ["12 参数仿射模型", "0.755", "1.812", "3.125", "17.122"],
        ],
    )
    doc.add_paragraph("图题：加速度计六位置标定 6 参数与 12 参数模型误差对比。对应文件：data/figures/accel_6pos_12param_error_compare.png。")

    doc.add_heading("WiFi 开关对磁力计静态输出影响", level=2)
    doc.add_paragraph(WIFI_RESULT)
    add_table(
        doc,
        ["指标", "WiFi 关闭", "WiFi 开启扫描", "变化"],
        [
            ["样本数", "1200", "1200", "59.950 s"],
            ["Mx 均值/raw count", "173.178", "174.195", "+1.017"],
            ["My 均值/raw count", "261.266", "260.824", "-0.442"],
            ["Mz 均值/raw count", "165.836", "164.262", "-1.574"],
            ["磁场模长均值/raw count", "354.627", "354.057", "-0.569"],
            ["磁场模长标准差/raw count", "1.532", "1.156", "0.754 倍"],
            ["三轴均值偏移/raw count", "-", "-", "1.926"],
            ["偏移占磁场模长比例", "-", "-", "0.543%"],
        ],
    )
    doc.add_paragraph(WIFI_CONCLUSION)
    doc.add_paragraph("图题：ESP32 WiFi 关闭与开启扫描状态下磁力计三轴原始数据时间序列。对应文件：data/figures/mag_wifi_compare_timeseries.png。")
    doc.add_paragraph("图题：ESP32 WiFi 开启前后磁场均值及均值偏移对比。对应文件：data/figures/mag_wifi_compare_shift.png。")


def main():
    doc = Document(SRC)
    replace_matching_paragraphs(doc)
    append_update_section(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
