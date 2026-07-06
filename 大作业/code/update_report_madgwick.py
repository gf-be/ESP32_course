# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent / "sensor-final-project"
DOC_DIR = ROOT / "docs"
SRC = DOC_DIR / "new_pcb_report_check_补充12参数加速度计与WiFi干扰.docx"
OUT = DOC_DIR / "new_pcb_report_check_补充12参数WiFi与Madgwick.docx"


INTRO = (
    "本章在完成加速度计、磁力计和陀螺仪基础标定后，进一步实现姿态融合实验。"
    "实验采用四组典型运动状态：水平静止姿态、固定倾斜姿态、晃动后回到水平姿态、连续旋转或手动倾斜过程。"
    "每组实验均保存原始 IMU 与磁力计数据，并在电脑端离线运行互补滤波、Mahony PI 和 Madgwick MARG 三种姿态融合算法，"
    "以比较静态稳定性、动态响应和计算更新频率。新版分析在每个实验段开始时重置滤波器状态，并使用加速度计 12 参数仿射标定结果。"
)

PARAMS = (
    "互补滤波采用 alpha=0.98；Mahony PI 采用比例项 kp=0.08、积分项 ki=0.004；"
    "Madgwick MARG 采用四元数姿态表示，并利用陀螺仪积分、加速度计重力方向和磁力计地磁方向进行梯度下降修正，"
    "本次实验 beta 参数取 0.035。"
)

RESULT_TEXT = (
    "姿态融合实验现已完成互补滤波、Mahony PI 和 Madgwick MARG 三种算法对比。"
    "在水平静止状态下，Madgwick 的 roll/pitch/yaw 标准差分别为 0.0196°、0.0204° 和 0.0394°，"
    "与互补滤波的 0.0189°、0.0186° 和 0.0333° 接近，优于本次参数下 Mahony PI 的 0.0299°、0.0280° 和 0.0518°。"
    "在固定倾斜状态下，Madgwick 的 roll/pitch/yaw 标准差分别为 0.0709°、0.0643° 和 0.1002°；"
    "其 yaw 抖动小于 Mahony PI，但 roll/pitch 抖动大于互补滤波和 Mahony PI，说明 Madgwick 的 beta 参数及磁力计权重仍有进一步整定空间。"
)

RATE_TEXT = (
    "离线计算频率统计显示，互补滤波、Mahony PI 和 Madgwick MARG 的单算法更新吞吐率分别约为 "
    "363093 Hz、193657 Hz 和 37846 Hz。Madgwick 由于包含四元数梯度下降和磁力计目标函数，计算量明显高于互补滤波和 Mahony PI，"
    "但仍远高于课程 Spec 中姿态融合更新率不低于 100 Hz 的要求。需要注意，该频率为电脑端离线吞吐率，ESP32 端实时部署时还应结合串口输出、I2C 读取和主循环调度重新测试。"
)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    return table


def replace_old_paragraphs(doc):
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("第三，实现并对比多种姿态估计算法。"):
            para.text = (
                "第三，实现并对比多种姿态估计算法。本文已设计并验证互补滤波、Mahony PI 与 Madgwick MARG 三种姿态融合算法，"
                "实现对 roll、pitch、yaw 姿态角的估计，并进一步准备简化 ESKF 作为高分扩展算法。"
            )
            continue
        if text.startswith("本章在完成加速度计、磁力计和陀螺仪基础标定后，进一步实现姿态融合实验"):
            para.text = INTRO
        elif text.startswith("互补滤波采用 alpha=0.98"):
            para.text = PARAMS
        elif text.startswith("姿态融合实验完成了互补滤波与 Mahony PI 两种算法对比"):
            para.text = RESULT_TEXT
        elif "Madgwick 待补" in text:
            para.text = text.replace("Mahony PI 滤波（Madgwick 待补）器", "Mahony PI 与 Madgwick 滤波器").replace("Madgwick 待补", "Madgwick MARG")
        elif "两种算法均使用 Allan 方差实验得到的陀螺仪零偏进行预补偿" in text:
            para.text = text.replace("两种算法均使用", "三种算法均使用")


def append_section(doc):
    doc.add_page_break()
    doc.add_heading("补充实验结果：Madgwick 姿态融合算法", level=1)
    doc.add_paragraph(INTRO)
    doc.add_paragraph(PARAMS)
    doc.add_paragraph(RESULT_TEXT)
    add_table(
        doc,
        ["实验状态", "算法", "Roll std/deg", "Pitch std/deg", "Yaw std/deg"],
        [
            ["水平静止", "Complementary", "0.0189", "0.0186", "0.0333"],
            ["水平静止", "Mahony PI", "0.0299", "0.0280", "0.0518"],
            ["水平静止", "Madgwick", "0.0196", "0.0204", "0.0394"],
            ["固定倾斜", "Complementary", "0.0263", "0.0242", "0.0868"],
            ["固定倾斜", "Mahony PI", "0.0385", "0.0350", "0.1263"],
            ["固定倾斜", "Madgwick", "0.0709", "0.0643", "0.1002"],
        ],
    )
    doc.add_paragraph(RATE_TEXT)
    add_table(
        doc,
        ["指标", "数值", "单位"],
        [
            ["IMU 数据采样率", "49.561", "Hz"],
            ["Complementary 离线更新吞吐率", "363093", "Hz"],
            ["Mahony PI 离线更新吞吐率", "193657", "Hz"],
            ["Madgwick 离线更新吞吐率", "37846", "Hz"],
        ],
    )
    doc.add_paragraph("图题：水平静止状态下三种姿态融合算法的 roll、pitch、yaw 时间序列对比。对应文件：data/figures/attitude_level_static_rpy.png。")
    doc.add_paragraph("图题：固定倾斜状态下三种姿态融合算法的 roll、pitch、yaw 时间序列对比。对应文件：data/figures/attitude_tilt_static_rpy.png。")
    doc.add_paragraph("图题：晃动后回到水平过程中的三种姿态融合算法动态响应曲线。对应文件：data/figures/attitude_shake_return_response.png。")
    doc.add_paragraph("图题：连续手动运动过程中的三种姿态融合算法姿态角输出对比。对应文件：data/figures/attitude_continuous_motion_rpy.png。")


def clean_residual_old_wording(doc):
    for para in doc.paragraphs:
        text = para.text
        if "Madgwick 待补" in text or "两种算法均使用" in text:
            text = text.replace("Mahony PI 滤波（Madgwick 待补）器", "Mahony PI 与 Madgwick 滤波器")
            text = text.replace("Madgwick 待补", "Madgwick MARG")
            text = text.replace("两种算法均使用", "三种算法均使用")
            para.text = text


def main():
    doc = Document(SRC)
    replace_old_paragraphs(doc)
    append_section(doc)
    clean_residual_old_wording(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
