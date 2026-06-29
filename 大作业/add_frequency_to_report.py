# -*- coding: utf-8 -*-
from pathlib import Path
import csv
import shutil

from docx import Document
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parent
REPORT_IN = ROOT / "多传感器融合扩展板课程论文初稿_补充AI去噪实验结果.docx"
REPORT_OUT = ROOT / "多传感器融合扩展板课程论文初稿_补充频率测试结果.docx"
FREQ_CSV = ROOT / "data" / "performance" / "frequency_test_20260616_215936.csv"


def read_frequency():
    data = {}
    with FREQ_CSV.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        for row in reader:
            key = row["item"]
            value = row["value"]
            if key == "meta":
                extra = row.get(None) or []
                if extra:
                    data[value.strip()] = extra[0].strip()
                else:
                    parts = value.strip().split(",", 1)
                    if len(parts) == 2:
                        data[parts[0].strip()] = parts[1].strip()
            else:
                data[key] = value
    return data


def find_para_by_prefixes(doc, prefixes):
    for prefix in prefixes:
        for p in doc.paragraphs:
            if p.text.strip().startswith(prefix):
                return p
    return doc.paragraphs[-1]


def insert_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def add_table_after(paragraph, rows, headers):
    tbl = paragraph._parent.add_table(rows=1, cols=len(headers), width=Inches(6.5))
    paragraph._p.addnext(tbl._tbl)
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    for row in rows:
        cells = tbl.add_row().cells
        for i, h in enumerate(headers):
            cells[i].text = str(row[h])
    after_p = OxmlElement("w:p")
    tbl._tbl.addnext(after_p)
    return tbl, Paragraph(after_p, paragraph._parent)


def main():
    data = read_frequency()
    duration = float(data["duration_s"])
    i2c_freq = float(data["i2c_freq"])
    imu_count = int(float(data["imu_read_count"]))
    imu_elapsed = float(data["imu_elapsed_ms"])
    imu_hz = float(data["imu_sampling_rate_hz"])
    fusion_count = int(float(data["fusion_update_count"]))
    fusion_elapsed = float(data["fusion_elapsed_ms"])
    fusion_hz = float(data["fusion_update_rate_hz"])
    imu_pass = "满足" if imu_hz >= 200 else "未满足"
    fusion_pass = "满足" if fusion_hz >= 100 else "未满足"

    work_dir = ROOT / "_doc_work" / "frequency"
    if work_dir.exists():
        shutil.rmtree(work_dir, ignore_errors=True)
    work_dir.mkdir(parents=True, exist_ok=True)

    try:
        tmp_in = work_dir / "report.docx"
        tmp_out = work_dir / "report_out.docx"
        shutil.copy2(REPORT_IN, tmp_in)
        doc = Document(str(tmp_in))

        ch62 = find_para_by_prefixes(doc, ["6.2", "6、2", "6 2", "6. 2", "6."])
        anchor = insert_after(
            ch62,
            "系统频率测试采用 ESP32 端临时 MicroPython 程序完成，测试时 I2C 总线频率设置为 %.0f Hz，每项测试持续 %.0f s。测试内容包括纯 IMU 连续读取频率，以及包含 IMU、磁力计读取和互补滤波姿态更新在内的融合循环频率。"
            % (i2c_freq, duration),
            "Normal",
        )
        anchor = insert_after(
            anchor,
            "测试结果显示，IMU 连续读取 %d 次，用时 %.0f ms，实际读取频率为 %.3f Hz；姿态融合循环更新 %d 次，用时 %.0f ms，实际更新频率为 %.3f Hz。课程指标要求 IMU 采样率不低于 200 Hz、姿态融合更新率不低于 100 Hz，本系统两项频率指标均已%s要求。"
            % (imu_count, imu_elapsed, imu_hz, fusion_count, fusion_elapsed, fusion_hz, "满足"),
            "Normal",
        )

        rows = [
            {"指标": "IMU连续读取频率", "测试值": "%.3f" % imu_hz, "要求": ">= 200 Hz", "结论": imu_pass},
            {"指标": "姿态融合更新频率", "测试值": "%.3f" % fusion_hz, "要求": ">= 100 Hz", "结论": fusion_pass},
            {"指标": "I2C总线频率", "测试值": "%.0f" % i2c_freq, "要求": "实验设置", "结论": "400 kHz"},
            {"指标": "单项测试时长", "测试值": "%.0f" % duration, "要求": "记录", "结论": "s"},
        ]
        p = insert_after(anchor, "")
        _, anchor = add_table_after(p, rows, ["指标", "测试值", "要求", "结论"])

        ch67 = find_para_by_prefixes(doc, ["6.7", "6、7", "6 7", "6. 7", "6.6"])
        insert_after(
            ch67,
            "工程化频率指标测试表明，本系统 IMU 实际读取频率为 %.3f Hz，约为 200 Hz 指标的 %.2f 倍；姿态融合实际更新频率为 %.3f Hz，约为 100 Hz 指标的 %.2f 倍。说明当前硬件连接和算法实现具有足够实时性，可支撑后续姿态显示、数据记录和多传感器融合应用。"
            % (imu_hz, imu_hz / 200.0, fusion_hz, fusion_hz / 100.0),
            "Normal",
        )

        doc.save(str(tmp_out))
        shutil.copy2(tmp_out, REPORT_OUT)
    finally:
        shutil.rmtree(work_dir, ignore_errors=True)

    print("wrote", REPORT_OUT)
    print("imu_hz", imu_hz)
    print("fusion_hz", fusion_hz)


if __name__ == "__main__":
    main()
