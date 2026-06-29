# -*- coding: utf-8 -*-
from pathlib import Path
import csv
import shutil

from docx import Document
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parent
REPORT_IN = ROOT / "多传感器融合扩展板课程论文初稿_补充陀螺仪Allan方差结果.docx"
REPORT_OUT = ROOT / "多传感器融合扩展板课程论文初稿_补充姿态融合实验结果.docx"
STATIC_CSV = ROOT / "data" / "analysis" / "attitude_fusion_static_std.csv"
RATE_CSV = ROOT / "data" / "analysis" / "attitude_fusion_update_rate.csv"


def read_csv(path):
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def find_para_by_prefixes(doc, prefixes):
    for prefix in prefixes:
        for p in doc.paragraphs:
            if p.text.strip().startswith(prefix):
                return p
    return doc.paragraphs[-1]


def insert_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def add_table_after(paragraph, rows, headers):
    tbl = paragraph._parent.add_table(rows=1, cols=len(headers), width=Inches(6.5))
    paragraph._p.addnext(tbl._tbl)
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    for row in rows:
        cells = tbl.add_row().cells
        for i, h in enumerate(headers):
            cells[i].text = str(row[h])
    after_p = OxmlElement("w:p")
    tbl._tbl.addnext(after_p)
    return tbl, Paragraph(after_p, paragraph._parent)


def main():
    static_rows = read_csv(STATIC_CSV)
    rate_rows = read_csv(RATE_CSV)
    rate = {r["item"]: float(r["value"]) for r in rate_rows}

    level_comp = next(r for r in static_rows if r["phase"] == "level_static" and r["algorithm"] == "complementary")
    level_mahony = next(r for r in static_rows if r["phase"] == "level_static" and r["algorithm"] == "mahony")
    tilt_comp = next(r for r in static_rows if r["phase"] == "tilt_static" and r["algorithm"] == "complementary")
    tilt_mahony = next(r for r in static_rows if r["phase"] == "tilt_static" and r["algorithm"] == "mahony")

    work_dir = ROOT / "_doc_work" / "attitude_fusion"
    if work_dir.exists():
        shutil.rmtree(work_dir, ignore_errors=True)
    work_dir.mkdir(parents=True, exist_ok=True)

    try:
        tmp_in = work_dir / "report.docx"
        tmp_out = work_dir / "report_out.docx"
        shutil.copy2(REPORT_IN, tmp_in)
        doc = Document(str(tmp_in))

        ch5 = find_para_by_prefixes(doc, ["5.", "5 ", "第5", "第五"])
        anchor = insert_after(
            ch5,
            "本章在完成加速度计、磁力计和陀螺仪基础标定后，进一步实现姿态融合实验。实验采用四组典型运动状态：水平静止姿态、固定倾斜姿态、晃动后回到水平姿态、连续旋转或手动倾斜过程。每组实验均保存原始 IMU 与磁力计数据，并在电脑端离线运行两种姿态融合算法，以比较静态稳定性、动态响应和计算更新频率。",
            "Normal",
        )
        anchor = insert_after(
            anchor,
            "第一种算法为互补滤波。该算法利用陀螺仪积分获得短时间姿态变化，并用加速度计的重力方向和磁力计航向角对低频漂移进行修正。实验中互补系数取 0.98，因此短时间响应主要由陀螺仪决定，长期漂移由加速度计和磁力计缓慢拉回。",
            "Normal",
        )
        anchor = insert_after(
            anchor,
            "第二种算法为 Mahony 风格 PI 姿态融合。该方法在陀螺仪角速度积分基础上，引入由加速度计和磁力计构造的姿态误差，并通过比例项和积分项修正角速度积分漂移。与互补滤波相比，该方法对零偏和低频漂移具有更强抑制能力，但计算量略高。两种算法均使用 Allan 方差实验得到的陀螺仪零偏进行预补偿。",
            "Normal",
        )
        anchor = insert_after(
            anchor,
            "姿态融合实验共采集 16500 组样本，整体有效采样率为 %.3f Hz。离线运行时，互补滤波更新频率约 %.1f Hz，Mahony 风格 PI 融合更新频率约 %.1f Hz，二者均远高于 50 Hz 采样频率，说明在 ESP32 或上位机环境中均具备实时运行基础。"
            % (rate["sample_rate_hz"], rate["complementary_update_hz"], rate["mahony_update_hz"]),
            "Normal",
        )

        table_rows = []
        phase_name = {"level_static": "水平静止", "tilt_static": "固定倾斜"}
        alg_name = {"complementary": "互补滤波", "mahony": "Mahony PI"}
        for r in static_rows:
            table_rows.append({
                "实验状态": phase_name.get(r["phase"], r["phase"]),
                "算法": alg_name.get(r["algorithm"], r["algorithm"]),
                "roll标准差(°)": "%.3f" % float(r["roll_std_deg"]),
                "pitch标准差(°)": "%.3f" % float(r["pitch_std_deg"]),
                "yaw标准差(°)": "%.3f" % float(r["yaw_std_deg"]),
                "roll均值(°)": "%.3f" % float(r["roll_mean_deg"]),
                "pitch均值(°)": "%.3f" % float(r["pitch_mean_deg"]),
                "yaw均值(°)": "%.3f" % float(r["yaw_mean_deg"]),
            })
        p = insert_after(anchor, "")
        _, anchor = add_table_after(
            p,
            table_rows,
            ["实验状态", "算法", "roll标准差(°)", "pitch标准差(°)", "yaw标准差(°)", "roll均值(°)", "pitch均值(°)", "yaw均值(°)"],
        )

        rate_table = [
            {"项目": "样本数", "结果": "%.0f" % rate["samples"], "单位": "rows"},
            {"项目": "数据时长", "结果": "%.3f" % rate["data_duration_s"], "单位": "s"},
            {"项目": "采样率", "结果": "%.3f" % rate["sample_rate_hz"], "单位": "Hz"},
            {"项目": "互补滤波更新频率", "结果": "%.1f" % rate["complementary_update_hz"], "单位": "Hz"},
            {"项目": "Mahony PI 更新频率", "结果": "%.1f" % rate["mahony_update_hz"], "单位": "Hz"},
        ]
        p = insert_after(anchor, "")
        _, anchor = add_table_after(p, rate_table, ["项目", "结果", "单位"])

        anchor = insert_after(anchor, "图：水平静止状态下互补滤波与 Mahony PI 算法姿态角输出（对应文件 attitude_level_static_rpy.png）", "Normal")
        anchor = insert_after(anchor, "图：固定倾斜状态下互补滤波与 Mahony PI 算法姿态角输出（对应文件 attitude_tilt_static_rpy.png）", "Normal")
        anchor = insert_after(anchor, "图：晃动后回到水平过程中的姿态动态响应曲线（对应文件 attitude_shake_return_response.png）", "Normal")
        anchor = insert_after(anchor, "图：连续手动倾斜过程中的 roll、pitch、yaw 时间序列（对应文件 attitude_continuous_motion_rpy.png）", "Normal")

        anchor = insert_after(
            anchor,
            "由静态标准差可见，在水平静止状态下，互补滤波 roll/pitch/yaw 标准差分别为 %.3f°、%.3f° 和 %.3f°，Mahony PI 分别为 %.3f°、%.3f° 和 %.3f°；在固定倾斜状态下，互补滤波三轴标准差分别为 %.3f°、%.3f° 和 %.3f°，Mahony PI 分别为 %.3f°、%.3f° 和 %.3f°。整体上 Mahony PI 在静态段输出更平稳，尤其对 roll、pitch 角抖动抑制更明显；yaw 角仍受磁力计噪声和环境磁场影响，稳定性弱于 roll/pitch。"
            % (
                float(level_comp["roll_std_deg"]),
                float(level_comp["pitch_std_deg"]),
                float(level_comp["yaw_std_deg"]),
                float(level_mahony["roll_std_deg"]),
                float(level_mahony["pitch_std_deg"]),
                float(level_mahony["yaw_std_deg"]),
                float(tilt_comp["roll_std_deg"]),
                float(tilt_comp["pitch_std_deg"]),
                float(tilt_comp["yaw_std_deg"]),
                float(tilt_mahony["roll_std_deg"]),
                float(tilt_mahony["pitch_std_deg"]),
                float(tilt_mahony["yaw_std_deg"]),
            ),
            "Normal",
        )

        ch65 = find_para_by_prefixes(doc, ["6.5", "6、5", "6 5", "6. 5"])
        insert_after(
            ch65,
            "姿态融合实验完成了互补滤波与 Mahony PI 两种算法对比。实验数据覆盖水平静止、固定倾斜、晃动后回水平和连续手动倾斜四类状态。结果表明，互补滤波结构简单、更新频率最高，适合资源受限场景；Mahony PI 在引入比例-积分误差反馈后，静态姿态角抖动更小，水平静止状态下 roll/pitch 标准差由互补滤波的 %.3f°/%.3f° 降至 %.3f°/%.3f°。不过 yaw 角仍明显受磁力计环境扰动影响，因此后续若继续优化，应重点改善磁力计安装位置、磁场标定和航向角异常值抑制。"
            % (
                float(level_comp["roll_std_deg"]),
                float(level_comp["pitch_std_deg"]),
                float(level_mahony["roll_std_deg"]),
                float(level_mahony["pitch_std_deg"]),
            ),
            "Normal",
        )

        doc.save(str(tmp_out))
        shutil.copy2(tmp_out, REPORT_OUT)
    finally:
        shutil.rmtree(work_dir, ignore_errors=True)

    print("wrote", REPORT_OUT)
    print("sample_rate", rate["sample_rate_hz"])
    print("comp_update_hz", rate["complementary_update_hz"])
    print("mahony_update_hz", rate["mahony_update_hz"])


if __name__ == "__main__":
    main()
