# -*- coding: utf-8 -*-
from pathlib import Path
import csv
import shutil

from docx import Document
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parent
REPORT_IN = ROOT / "多传感器融合扩展板课程论文初稿_补充磁力计椭球标定结果.docx"
REPORT_OUT = ROOT / "多传感器融合扩展板课程论文初稿_补充陀螺仪Allan方差结果.docx"
SUMMARY_CSV = ROOT / "data" / "analysis" / "gyro_allan_summary.csv"
AXIS_CSV = ROOT / "data" / "analysis" / "gyro_allan_axis_stats.csv"


def read_summary():
    data = {}
    with SUMMARY_CSV.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        for row in reader:
            data[row["item"]] = row
    return data


def read_axis_stats():
    with AXIS_CSV.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def val(summary, key, as_float=True):
    raw = summary[key]["value"]
    return float(raw) if as_float else raw


def find_para_startswith(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise ValueError("paragraph not found: " + prefix)


def insert_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def add_table_after(paragraph, rows, headers):
    tbl = paragraph._parent.add_table(rows=1, cols=len(headers), width=Inches(6.5))
    paragraph._p.addnext(tbl._tbl)
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    for row in rows:
        cells = tbl.add_row().cells
        for i, h in enumerate(headers):
            cells[i].text = str(row[h])
    after_p = OxmlElement("w:p")
    tbl._tbl.addnext(after_p)
    return tbl, Paragraph(after_p, paragraph._parent)


def main():
    summary = read_summary()
    axis_rows = read_axis_stats()

    source_file = val(summary, "source_file", as_float=False)
    samples = int(val(summary, "samples"))
    duration = val(summary, "duration")
    sample_rate = val(summary, "sample_rate")
    temp_min = val(summary, "temp_min")
    temp_max = val(summary, "temp_max")
    temp_range = val(summary, "temp_range")

    by_axis = {row["axis"]: row for row in axis_rows}

    work_dir = ROOT / "_doc_work" / "gyro_allan"
    if work_dir.exists():
        shutil.rmtree(work_dir, ignore_errors=True)
    work_dir.mkdir(parents=True, exist_ok=True)

    try:
        tmp_in = work_dir / "report.docx"
        tmp_out = work_dir / "report_out.docx"
        shutil.copy2(REPORT_IN, tmp_in)
        doc = Document(str(tmp_in))

        anchor = find_para_startswith(doc, "4.2")
        anchor = insert_after(
            anchor,
            "为进一步分析陀螺仪静态噪声与零偏稳定性，本实验新增陀螺仪 Allan 方差测试。实验数据来自 %s，板卡在整个采集过程中保持静止，ESP32 WiFi 关闭，采样率为 %.3f Hz，持续时间为 %.3f s，共获得 %d 组样本。实验期间温度范围为 %.3f°C 至 %.3f°C，温度变化 %.3f°C，说明测试环境温度相对稳定。"
            % (source_file, sample_rate, duration, samples, temp_min, temp_max, temp_range),
            "Normal",
        )
        anchor = insert_after(
            anchor,
            "静态统计结果表明，陀螺仪三轴均存在非零零偏。其中 X/Y/Z 轴平均零偏分别为 %.6f deg/s、%.6f deg/s 和 %.6f deg/s，标准差分别为 %.6f deg/s、%.6f deg/s 和 %.6f deg/s。Y 轴零偏明显大于 X、Z 两轴，若直接对角速度积分，将在长时间姿态估计中引入明显漂移，因此后续互补滤波、Mahony 等姿态融合算法需要先扣除陀螺仪零偏。"
            % (
                float(by_axis["x"]["bias_mean_dps"]),
                float(by_axis["y"]["bias_mean_dps"]),
                float(by_axis["z"]["bias_mean_dps"]),
                float(by_axis["x"]["std_dps"]),
                float(by_axis["y"]["std_dps"]),
                float(by_axis["z"]["std_dps"]),
            ),
            "Normal",
        )
        anchor = insert_after(
            anchor,
            "Allan 方差分析显示，三轴 Allan 偏差随平均时间 tau 增大先下降，随后在较长平均时间附近趋于平缓，符合 MEMS 陀螺仪短时间白噪声主导、长时间零偏慢漂移逐渐显现的典型特征。X/Y/Z 轴最小 Allan 偏差分别为 %.6f deg/s、%.6f deg/s 和 %.6f deg/s，对应平均时间分别为 %.3f s、%.3f s 和 %.3f s。该结果可作为后续滤波器参数选择和姿态融合误差分析的依据。"
            % (
                float(by_axis["x"]["allan_min_dps"]),
                float(by_axis["y"]["allan_min_dps"]),
                float(by_axis["z"]["allan_min_dps"]),
                float(by_axis["x"]["allan_min_tau_s"]),
                float(by_axis["y"]["allan_min_tau_s"]),
                float(by_axis["z"]["allan_min_tau_s"]),
            ),
            "Normal",
        )

        table_rows = []
        axis_name = {"x": "X", "y": "Y", "z": "Z"}
        for axis in ("x", "y", "z"):
            row = by_axis[axis]
            table_rows.append({
                "轴": axis_name[axis],
                "零偏(deg/s)": "%.6f" % float(row["bias_mean_dps"]),
                "标准差(deg/s)": "%.6f" % float(row["std_dps"]),
                "Allan最小值(deg/s)": "%.6f" % float(row["allan_min_dps"]),
                "对应tau(s)": "%.3f" % float(row["allan_min_tau_s"]),
                "ARW(deg/sqrt(h))": "%.6f" % float(row["arw_deg_per_sqrt_h"]),
            })
        p = insert_after(anchor, "")
        _, anchor = add_table_after(
            p,
            table_rows,
            ["轴", "零偏(deg/s)", "标准差(deg/s)", "Allan最小值(deg/s)", "对应tau(s)", "ARW(deg/sqrt(h))"],
        )

        anchor = insert_after(anchor, "图：静止状态下陀螺仪三轴输出与温度变化曲线（对应文件 gyro_allan_timeseries.png）", "Normal")
        anchor = insert_after(anchor, "图：静止状态下陀螺仪三轴零偏分布直方图（对应文件 gyro_allan_histogram.png）", "Normal")
        anchor = insert_after(anchor, "图：陀螺仪三轴 Allan 偏差曲线（对应文件 gyro_allan_deviation.png）", "Normal")

        exp_anchor = find_para_startswith(doc, "6.3")
        insert_after(
            exp_anchor,
            "陀螺仪 Allan 方差实验结果显示，在 30 min 静止采样条件下，X/Y/Z 轴零偏分别为 %.6f deg/s、%.6f deg/s 和 %.6f deg/s，标准差分别为 %.6f deg/s、%.6f deg/s 和 %.6f deg/s；三轴最小 Allan 偏差分别为 %.6f deg/s、%.6f deg/s 和 %.6f deg/s。该结果说明当前 IMU 陀螺仪存在可观测零偏，其中 Y 轴零偏最明显，后续姿态融合中应进行零偏补偿。"
            % (
                float(by_axis["x"]["bias_mean_dps"]),
                float(by_axis["y"]["bias_mean_dps"]),
                float(by_axis["z"]["bias_mean_dps"]),
                float(by_axis["x"]["std_dps"]),
                float(by_axis["y"]["std_dps"]),
                float(by_axis["z"]["std_dps"]),
                float(by_axis["x"]["allan_min_dps"]),
                float(by_axis["y"]["allan_min_dps"]),
                float(by_axis["z"]["allan_min_dps"]),
            ),
            "Normal",
        )

        doc.save(str(tmp_out))
        shutil.copy2(tmp_out, REPORT_OUT)
    finally:
        shutil.rmtree(work_dir, ignore_errors=True)

    print("wrote", REPORT_OUT)
    print("source", source_file)
    print("samples", samples)
    print("duration", duration)


if __name__ == "__main__":
    main()
