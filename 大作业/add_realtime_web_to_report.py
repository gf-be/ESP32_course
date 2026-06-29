# -*- coding: utf-8 -*-
from pathlib import Path
import shutil

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parent
REPORT_IN = ROOT / "多传感器融合扩展板课程论文初稿_补充BMP280实验结果.docx"
REPORT_OUT = ROOT / "多传感器融合扩展板课程论文初稿_补充实时姿态显示结果.docx"
WEB_IMAGE = ROOT / "photo" / "MahonyPI_web.png"


def find_para_by_prefixes(doc, prefixes):
    for prefix in prefixes:
        for p in doc.paragraphs:
            if p.text.strip().startswith(prefix):
                return p
    return doc.paragraphs[-1]


def insert_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def add_caption(paragraph, text):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.size = Pt(9)


def main():
    work_dir = ROOT / "_doc_work" / "realtime_web"
    if work_dir.exists():
        shutil.rmtree(work_dir, ignore_errors=True)
    work_dir.mkdir(parents=True, exist_ok=True)

    try:
        tmp_in = work_dir / "report.docx"
        tmp_out = work_dir / "report_out.docx"
        shutil.copy2(REPORT_IN, tmp_in)
        doc = Document(str(tmp_in))

        ch5 = find_para_by_prefixes(doc, ["5.", "5 ", "第5", "第五"])
        anchor = insert_after(
            ch5,
            "为满足课程要求中的“实时姿态显示”数据可视化任务，本项目进一步实现了 Python/Web 实时姿态板。电脑端程序 pc_mahony_serial_web.py 启动本地 Web 服务，并通过串口 raw REPL 将临时 MicroPython 程序发送至 ESP32 RAM；ESP32 端实时读取 MPU 系列 IMU 与 HMC5883L 磁力计，运行 Mahony PI 姿态融合算法，并将 roll、pitch、yaw、温度和更新频率回传至电脑端，网页端以 3D 板卡模型和实时曲线显示姿态变化。",
            "Normal",
        )
        anchor = insert_after(
            anchor,
            "该实现与单纯串口打印不同，传感器读取和 Mahony PI 融合计算均在 ESP32 端实时完成，Python 端只负责启动程序、接收实时姿态流并提供 Web 可视化界面。实验中网页显示的更新频率稳定在约 100 Hz，倾斜板卡时 3D 模型和 roll/pitch/yaw 曲线同步变化，可作为答辩现场实时演示材料。",
            "Normal",
        )

        ch63 = find_para_by_prefixes(doc, ["6.3", "6、3", "6 3", "6. 3", "6."])
        anchor = insert_after(
            ch63,
            "实时姿态显示实验中，ESP32 端 Mahony PI 程序持续输出 t_ms、roll、pitch、yaw、temp_c 和 update_hz，Web 页面以 3D 姿态板和曲线方式显示结果。截图中更新频率稳定为 100 Hz，说明实时姿态显示链路满足课程中“串口 + Python/Web，倾斜板子看角度变化”的可视化要求。",
            "Normal",
        )
        p = insert_after(anchor, "")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(WEB_IMAGE), width=Inches(5.8))
        cap = insert_after(p, "")
        add_caption(cap, "图：Mahony PI 实时姿态显示 Web 3D 姿态板运行结果")
        insert_after(
            cap,
            "图中 Web 面板显示 roll、pitch、yaw、温度和更新频率，并以 3D 板卡模型实时反映姿态变化。该实验补充验证了姿态融合算法不仅能离线分析，也能够部署到 ESP32 端实时运行并完成网页可视化展示。",
            "Normal",
        )

        ch67 = find_para_by_prefixes(doc, ["6.7", "6、7", "6 7", "6. 7", "6.6"])
        insert_after(
            ch67,
            "实时姿态显示功能进一步补充了系统工程化验证：ESP32 端负责传感器采集和 Mahony PI 姿态融合，电脑端负责 Web 渲染和数据显示。该结构将嵌入式实时计算与上位机可视化分离，既能展示算法实时性，也便于答辩现场观察板卡倾斜时的姿态角变化。",
            "Normal",
        )

        doc.save(str(tmp_out))
        shutil.copy2(tmp_out, REPORT_OUT)
    finally:
        shutil.rmtree(work_dir, ignore_errors=True)

    print("wrote", REPORT_OUT)
    print("image", WEB_IMAGE)


if __name__ == "__main__":
    main()
