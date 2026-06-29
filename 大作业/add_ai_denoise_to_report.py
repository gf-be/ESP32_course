# -*- coding: utf-8 -*-
from pathlib import Path
import csv
import shutil

from docx import Document
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parent
REPORT_IN = ROOT / "多传感器融合扩展板课程论文初稿_补充姿态融合实验结果.docx"
REPORT_OUT = ROOT / "多传感器融合扩展板课程论文初稿_补充AI去噪实验结果.docx"
METRICS_CSV = ROOT / "data" / "analysis" / "ai_denoise_metrics.csv"
SUMMARY_CSV = ROOT / "data" / "analysis" / "ai_denoise_summary.csv"
HISTORY_CSV = ROOT / "data" / "analysis" / "ai_denoise_training_history.csv"


METHOD_NAME = {
    "raw": "原始数据",
    "lowpass": "低通滤波",
    "kalman": "一阶卡尔曼",
    "cnn_1d": "1D-CNN",
}


def read_csv(path):
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def read_summary():
    rows = read_csv(SUMMARY_CSV)
    return {r["item"]: r["value"] for r in rows}


def find_para_by_prefixes(doc, prefixes):
    for prefix in prefixes:
        for p in doc.paragraphs:
            if p.text.strip().startswith(prefix):
                return p
    return doc.paragraphs[-1]


def insert_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def add_table_after(paragraph, rows, headers):
    tbl = paragraph._parent.add_table(rows=1, cols=len(headers), width=Inches(6.5))
    paragraph._p.addnext(tbl._tbl)
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    for row in rows:
        cells = tbl.add_row().cells
        for i, h in enumerate(headers):
            cells[i].text = str(row[h])
    after_p = OxmlElement("w:p")
    tbl._tbl.addnext(after_p)
    return tbl, Paragraph(after_p, paragraph._parent)


def metric(metrics, method, channel, field):
    for row in metrics:
        if row["method"] == method and row["channel"] == channel:
            return float(row[field])
    raise KeyError((method, channel, field))


def main():
    metrics = read_csv(METRICS_CSV)
    summary = read_summary()
    history = read_csv(HISTORY_CSV)

    source_file = summary.get("source_file", "")
    sample_rate = float(summary.get("sample_rate", "50"))
    window = int(float(summary.get("window", "64")))
    stride = int(float(summary.get("stride", "4")))
    epochs = int(float(summary.get("epochs", "18")))
    train0 = float(history[0]["train_loss"])
    val0 = float(history[0]["val_loss"])
    train_last = float(history[-1]["train_loss"])
    val_last = float(history[-1]["val_loss"])

    work_dir = ROOT / "_doc_work" / "ai_denoise"
    if work_dir.exists():
        shutil.rmtree(work_dir, ignore_errors=True)
    work_dir.mkdir(parents=True, exist_ok=True)

    try:
        tmp_in = work_dir / "report.docx"
        tmp_out = work_dir / "report_out.docx"
        shutil.copy2(REPORT_IN, tmp_in)
        doc = Document(str(tmp_in))

        ch55 = find_para_by_prefixes(doc, ["5.5", "5、5", "5 5", "5. 5", "5."])
        anchor = insert_after(
            ch55,
            "本节进一步开展 IMU 数据 AI 去噪实验。实验不重新采集额外数据，而是利用前述陀螺仪 Allan 方差实验中的长时间静止 IMU 数据构造训练样本，数据来源为 %s。输入特征为 ax、ay、az、gx、gy、gz 六通道时间窗口，采样率为 %.3f Hz，窗口长度为 %d 个采样点，步长为 %d 个采样点。"
            % (source_file, sample_rate, window, stride),
            "Normal",
        )
        anchor = insert_after(
            anchor,
            "由于静止实验中真实角速度理论值为 0，加速度模长理论上应稳定，因此本实验采用自监督方式构造训练目标：先对原始六通道信号进行零相位滑动均值平滑，将其作为伪干净标签，再训练 1D-CNN 从含噪窗口中预测窗口中心的六通道去噪值。该方案不依赖外部高精度传感器，适合在课程实验条件下验证 AI 去噪方法的有效性。",
            "Normal",
        )
        anchor = insert_after(
            anchor,
            "对比方法包括四类：原始数据、低通滤波、一阶卡尔曼滤波和 1D-CNN。低通滤波用于体现传统平滑方法的效果；一阶卡尔曼滤波用于体现基于状态估计的递推去噪能力；1D-CNN 则利用局部时间窗口内六通道信号的相关性学习非线性去噪映射。评价指标包括标准差、RMSE、SNR，以及由加速度计计算得到的 roll/pitch 姿态角抖动。",
            "Normal",
        )
        anchor = insert_after(
            anchor,
            "1D-CNN 训练共进行 %d 个 epoch，训练损失由 %.6f 下降至 %.6f，验证损失由 %.6f 下降至 %.6f，说明模型在静止 IMU 数据上能够有效学习到平滑后的信号特征。"
            % (epochs, train0, train_last, val0, val_last),
            "Normal",
        )

        jitter_rows = []
        for method in ("raw", "lowpass", "kalman", "cnn_1d"):
            jitter_rows.append({
                "方法": METHOD_NAME[method],
                "roll标准差(°)": "%.4f" % metric(metrics, method, "roll_from_acc", "std"),
                "pitch标准差(°)": "%.4f" % metric(metrics, method, "pitch_from_acc", "std"),
                "roll RMSE(°)": "%.4f" % metric(metrics, method, "roll_from_acc", "rmse"),
                "pitch RMSE(°)": "%.4f" % metric(metrics, method, "pitch_from_acc", "rmse"),
                "pitch SNR(dB)": "%.2f" % metric(metrics, method, "pitch_from_acc", "snr_db"),
            })
        p = insert_after(anchor, "")
        _, anchor = add_table_after(
            p,
            jitter_rows,
            ["方法", "roll标准差(°)", "pitch标准差(°)", "roll RMSE(°)", "pitch RMSE(°)", "pitch SNR(dB)"],
        )

        channel_rows = []
        for ch in ("ax_g", "ay_g", "az_g", "gx_dps", "gy_dps", "gz_dps"):
            channel_rows.append({
                "通道": ch,
                "原始std": "%.6f" % metric(metrics, "raw", ch, "std"),
                "低通std": "%.6f" % metric(metrics, "lowpass", ch, "std"),
                "卡尔曼std": "%.6f" % metric(metrics, "kalman", ch, "std"),
                "1D-CNN std": "%.6f" % metric(metrics, "cnn_1d", ch, "std"),
                "1D-CNN SNR(dB)": "%.2f" % metric(metrics, "cnn_1d", ch, "snr_db"),
            })
        p = insert_after(anchor, "")
        _, anchor = add_table_after(
            p,
            channel_rows,
            ["通道", "原始std", "低通std", "卡尔曼std", "1D-CNN std", "1D-CNN SNR(dB)"],
        )

        anchor = insert_after(anchor, "图：1D-CNN IMU 去噪模型训练损失曲线（对应文件 ai_denoise_training_curve.png）", "Normal")
        anchor = insert_after(anchor, "图：原始数据、低通滤波、一阶卡尔曼与 1D-CNN 去噪信号对比（对应文件 ai_denoise_signal_compare.png）", "Normal")
        anchor = insert_after(anchor, "图：不同去噪方法下由加速度计估计的姿态角抖动对比（对应文件 ai_denoise_attitude_jitter.png）", "Normal")

        ch66 = find_para_by_prefixes(doc, ["6.6", "6、6", "6 6", "6. 6", "6.5"])
        insert_after(
            ch66,
            "AI 去噪实验结果表明，四种方法中 1D-CNN 对静止 IMU 数据的噪声抑制效果最好。原始数据下 roll/pitch 姿态角抖动标准差分别为 %.4f° 和 %.4f°；低通滤波后下降至 %.4f° 和 %.4f°；一阶卡尔曼滤波后为 %.4f° 和 %.4f°；1D-CNN 进一步降低至 %.4f° 和 %.4f°。该结果说明 1D-CNN 能够利用六通道时间窗口中的局部相关性，比单通道递推滤波更有效地抑制随机噪声。"
            % (
                metric(metrics, "raw", "roll_from_acc", "std"),
                metric(metrics, "raw", "pitch_from_acc", "std"),
                metric(metrics, "lowpass", "roll_from_acc", "std"),
                metric(metrics, "lowpass", "pitch_from_acc", "std"),
                metric(metrics, "kalman", "roll_from_acc", "std"),
                metric(metrics, "kalman", "pitch_from_acc", "std"),
                metric(metrics, "cnn_1d", "roll_from_acc", "std"),
                metric(metrics, "cnn_1d", "pitch_from_acc", "std"),
            ),
            "Normal",
        )
        insert_after(
            ch66,
            "从六通道标准差看，1D-CNN 将 ax_g 标准差由 %.6f 降至 %.6f，将 gx_dps 标准差由 %.6f deg/s 降至 %.6f deg/s，将 gz_dps 标准差由 %.6f deg/s 降至 %.6f deg/s。与低通滤波相比，1D-CNN 在姿态角抖动和多个通道 SNR 上进一步改善；但其需要离线训练和更多计算资源，因此更适合作为精度增强模块或离线数据处理方法。"
            % (
                metric(metrics, "raw", "ax_g", "std"),
                metric(metrics, "cnn_1d", "ax_g", "std"),
                metric(metrics, "raw", "gx_dps", "std"),
                metric(metrics, "cnn_1d", "gx_dps", "std"),
                metric(metrics, "raw", "gz_dps", "std"),
                metric(metrics, "cnn_1d", "gz_dps", "std"),
            ),
            "Normal",
        )

        doc.save(str(tmp_out))
        shutil.copy2(tmp_out, REPORT_OUT)
    finally:
        shutil.rmtree(work_dir, ignore_errors=True)

    print("wrote", REPORT_OUT)
    print("source", source_file)
    print("train_loss", train0, train_last)
    print("val_loss", val0, val_last)


if __name__ == "__main__":
    main()
