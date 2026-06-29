# -*- coding: utf-8 -*-
from pathlib import Path
import csv
import shutil

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parent
REPORT_IN = ROOT / "多传感器融合扩展板课程论文初稿_补充六位置标定结果.docx"
REPORT_OUT = ROOT / "多传感器融合扩展板课程论文初稿_补充磁力计椭球标定结果.docx"
SUMMARY_CSV = ROOT / "data" / "analysis" / "mag_ellipsoid_summary.csv"
MATRIX_CSV = ROOT / "data" / "analysis" / "mag_ellipsoid_calibration_matrix.csv"
FIG_DIR = ROOT / "data" / "figures"


def read_summary():
    data = {}
    with SUMMARY_CSV.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        for row in reader:
            data[row["item"]] = row
    return data


def read_matrix():
    rows = []
    with MATRIX_CSV.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        for row in reader:
            rows.append(row)
    return rows


def value(data, key, as_float=True):
    raw = data[key]["value"]
    return float(raw) if as_float else raw


def find_para_startswith(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise ValueError("paragraph not found: " + prefix)


def insert_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def add_caption(paragraph, text):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.size = Pt(9)


def add_table_after(paragraph, rows, headers):
    tbl = paragraph._parent.add_table(rows=1, cols=len(headers), width=Inches(6.4))
    paragraph._p.addnext(tbl._tbl)
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    for row in rows:
        cells = tbl.add_row().cells
        for i, h in enumerate(headers):
            cells[i].text = str(row[h])
    after_p = OxmlElement("w:p")
    tbl._tbl.addnext(after_p)
    return tbl, Paragraph(after_p, paragraph._parent)


def add_picture_after(paragraph, path, caption):
    p = insert_after(paragraph, "")
    p.add_run().add_picture(str(path), width=Inches(5.7))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = insert_after(p, "")
    add_caption(cap, caption)
    return cap


def main():
    data = read_summary()
    matrix = read_matrix()

    samples = int(value(data, "samples"))
    source_file = value(data, "source_file", as_float=False)
    hx = value(data, "hard_iron_x")
    hy = value(data, "hard_iron_y")
    hz = value(data, "hard_iron_z")
    r1 = value(data, "ellipsoid_radius_1")
    r2 = value(data, "ellipsoid_radius_2")
    r3 = value(data, "ellipsoid_radius_3")
    imbalance = value(data, "axis_imbalance")
    raw_cv = value(data, "raw_radius_cv")
    cal_cv = value(data, "cal_radius_cv")
    improvement = value(data, "cv_improvement")
    raw_std = value(data, "raw_radius_std")
    cal_std = value(data, "cal_radius_std")

    work_dir = ROOT / "_doc_work" / "mag_ellipsoid"
    if work_dir.exists():
        shutil.rmtree(work_dir, ignore_errors=True)
    work_dir.mkdir(parents=True, exist_ok=True)

    try:
        tmp_in = work_dir / "report.docx"
        tmp_out = work_dir / "report_out.docx"
        shutil.copy2(REPORT_IN, tmp_in)
        doc = Document(str(tmp_in))

        anchor = find_para_startswith(doc, "4.2")
        anchor = insert_after(
            anchor,
            "在完成磁力计旋转干扰检查后，进一步进行磁力计椭球标定实验。实验数据来自 %s，共采集 %d 组样本，采样过程中关闭 ESP32 WiFi，并手持整块扩展板进行三维空间慢速旋转，覆盖正放、侧放、倒置和倾斜姿态，以满足椭球拟合对空间覆盖的要求。"
            % (source_file, samples),
            "Normal",
        )
        anchor = insert_after(
            anchor,
            "椭球拟合得到的 hard-iron 偏置为 X=%.3f、Y=%.3f、Z=%.3f raw count。三个主半径分别为 %.3f、%.3f、%.3f raw count，轴向不平衡比为 %.3f，说明磁力计存在可观测的固定偏置，但 soft-iron 椭球畸变较轻，当前采集覆盖质量可用于后续航向角补偿。"
            % (hx, hy, hz, r1, r2, r3, imbalance),
            "Normal",
        )
        anchor = insert_after(
            anchor,
            "校准前以中心化半径评价，半径变异系数为 %.4f，校准后为 %.4f，半径标准差由 %.3f raw count 变化为 %.3f raw count。由于原始数据本身接近球形，本次椭球校正的主要作用是消除 hard-iron 中心偏移，并提供可复用的 soft-iron 矫正矩阵。"
            % (raw_cv, cal_cv, raw_std, cal_std),
            "Normal",
        )

        table_rows = [
            {"项目": "样本数", "结果": str(samples), "单位": "rows"},
            {"项目": "hard-iron X", "结果": "%.3f" % hx, "单位": "raw count"},
            {"项目": "hard-iron Y", "结果": "%.3f" % hy, "单位": "raw count"},
            {"项目": "hard-iron Z", "结果": "%.3f" % hz, "单位": "raw count"},
            {"项目": "主半径 1/2/3", "结果": "%.3f / %.3f / %.3f" % (r1, r2, r3), "单位": "raw count"},
            {"项目": "轴向不平衡比", "结果": "%.3f" % imbalance, "单位": "ratio"},
            {"项目": "半径 CV 校准前", "结果": "%.4f" % raw_cv, "单位": "ratio"},
            {"项目": "半径 CV 校准后", "结果": "%.4f" % cal_cv, "单位": "ratio"},
            {"项目": "CV 改善倍数", "结果": "%.2f" % improvement, "单位": "times"},
        ]
        p = insert_after(anchor, "")
        _, anchor = add_table_after(p, table_rows, ["项目", "结果", "单位"])

        matrix_rows = []
        for row in matrix:
            matrix_rows.append({
                "矩阵行": row["row"],
                "m0": "%.6f" % float(row["m0"]),
                "m1": "%.6f" % float(row["m1"]),
                "m2": "%.6f" % float(row["m2"]),
            })
        p = insert_after(anchor, "soft-iron 矫正矩阵如下，实际使用时先减去 hard-iron 偏置，再左乘该矩阵完成椭球到球面的校正。", "Normal")
        _, anchor = add_table_after(p, matrix_rows, ["矩阵行", "m0", "m1", "m2"])

        cap = add_picture_after(
            anchor,
            FIG_DIR / "mag_ellipsoid_3d_before_after.png",
            "图：磁力计椭球标定前后三维点云对比",
        )
        cap = add_picture_after(
            cap,
            FIG_DIR / "mag_ellipsoid_projection_before_after.png",
            "图：磁力计椭球标定前后 XY/XZ/YZ 投影对比",
        )
        cap = add_picture_after(
            cap,
            FIG_DIR / "mag_ellipsoid_radius_hist.png",
            "图：磁力计标定前后半径分布对比",
        )

        exp_anchor = find_para_startswith(doc, "6.3")
        insert_after(
            exp_anchor,
            "磁力计椭球标定实验结果显示，HMC5883L 的 hard-iron 偏置约为 (%.3f, %.3f, %.3f) raw count，主半径不平衡比为 %.3f。校准前后半径 CV 为 %.4f -> %.4f。该结果说明当前 GY-273 与 ESP32 的安装位置虽然存在固定磁偏置，但整体软铁畸变较小，经过偏置扣除和椭球矩阵校正后可作为航向角解算的磁场输入。"
            % (hx, hy, hz, imbalance, raw_cv, cal_cv),
            "Normal",
        )

        doc.save(str(tmp_out))
        shutil.copy2(tmp_out, REPORT_OUT)
    finally:
        shutil.rmtree(work_dir, ignore_errors=True)

    print("wrote", REPORT_OUT)
    print("hard_iron", hx, hy, hz)
    print("axis_imbalance", imbalance)
    print("radius_cv", raw_cv, cal_cv)


if __name__ == "__main__":
    main()
