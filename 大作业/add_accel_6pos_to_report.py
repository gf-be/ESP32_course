# -*- coding: utf-8 -*-
from pathlib import Path
import csv
import shutil

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parent
REPORT_IN = ROOT / "多传感器融合扩展板课程论文初稿_补充PCB可实验性说明.docx"
REPORT_OUT = ROOT / "多传感器融合扩展板课程论文初稿_补充六位置标定结果.docx"
MEANS_CSV = ROOT / "data" / "analysis" / "accel_6pos_means.csv"
PARAMS_CSV = ROOT / "data" / "analysis" / "accel_6pos_calibration_params.csv"
FIG_DIR = ROOT / "data" / "figures"

PHOTO_MAP = [
    ("pos_x_up", "六位置标定x轴向上.jpg"),
    ("neg_x_up", "六位置标定x轴向下.jpg"),
    ("pos_y_up", "六位置标定y轴向上.jpg"),
    ("neg_y_up", "六位置标定y轴向下.jpg"),
    ("pos_z_up", "六位置标定z轴向上.jpg"),
    ("neg_z_up", "六位置标定z轴向下.jpg"),
]

LABELS = {
    "pos_x_up": "+X up",
    "neg_x_up": "-X up",
    "pos_y_up": "+Y up",
    "neg_y_up": "-Y up",
    "pos_z_up": "+Z up",
    "neg_z_up": "-Z up",
}

POSITION_CN = {
    "pos_x_up": "x 轴正向",
    "neg_x_up": "x 轴负向",
    "pos_y_up": "y 轴正向",
    "neg_y_up": "y 轴负向",
    "pos_z_up": "z 轴正向",
    "neg_z_up": "z 轴负向",
}


def read_dicts(path):
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def find_para_startswith(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise ValueError(f"paragraph not found: {prefix}")


def insert_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def add_caption(paragraph, text):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.size = Pt(9)


def add_table_after(paragraph, rows, headers):
    tbl = paragraph._parent.add_table(rows=1, cols=len(headers), width=Inches(6.4))
    paragraph._p.addnext(tbl._tbl)
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = tbl.add_row().cells
        for i, h in enumerate(headers):
            cells[i].text = str(row[h])
    tbl.style = "Table Grid"

    after_p = OxmlElement("w:p")
    tbl._tbl.addnext(after_p)
    return tbl, Paragraph(after_p, paragraph._parent)


def make_figures(means):
    FIG_DIR.mkdir(parents=True, exist_ok=True)

    plt.rcParams["font.sans-serif"] = ["Arial", "Times New Roman", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False

    labels = [LABELS[r["position"]] for r in means]
    raw_err = [float(r["raw_error_mg"]) for r in means]
    cal_err = [float(r["cal_error_mg"]) for r in means]
    raw_norm = [float(r["raw_norm_g"]) for r in means]
    cal_norm = [float(r["cal_norm_g"]) for r in means]

    x = range(len(labels))

    fig, ax = plt.subplots(figsize=(9, 4.8))
    width = 0.36
    ax.bar([i - width / 2 for i in x], raw_err, width=width, label="校准前", color="#f97316")
    ax.bar([i + width / 2 for i in x], cal_err, width=width, label="校准后", color="#2563eb")
    ax.set_xticks(list(x), labels, rotation=20)
    ax.set_ylabel("加速度模长误差 (mg)")
    ax.set_title("六位置标定：误差对比")
    ax.grid(True, axis="y", alpha=0.25)
    ax.legend()
    fig.tight_layout()
    err_path = FIG_DIR / "accel_6pos_error_compare.png"
    fig.savefig(err_path, dpi=180)
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(9, 4.8))
    ax.plot(labels, raw_norm, marker="o", label="校准前", color="#f97316")
    ax.plot(labels, cal_norm, marker="o", label="校准后", color="#2563eb")
    ax.axhline(1.0, linestyle="--", linewidth=0.9, color="black")
    ax.set_ylabel("加速度模长 (g)")
    ax.set_title("六位置标定：模长与 1g 理论值对比")
    ax.grid(True, alpha=0.25)
    ax.legend()
    fig.tight_layout()
    norm_path = FIG_DIR / "accel_6pos_norm_compare.png"
    fig.savefig(norm_path, dpi=180)
    plt.close(fig)

    return err_path, norm_path


def summarize(means):
    raw = [float(r["raw_error_mg"]) for r in means]
    cal = [float(r["cal_error_mg"]) for r in means]
    raw_mean = sum(raw) / len(raw)
    cal_mean = sum(cal) / len(cal)
    return {
        "raw_mean": raw_mean,
        "cal_mean": cal_mean,
        "raw_max": max(raw),
        "cal_max": max(cal),
        "improve": raw_mean / cal_mean if cal_mean else 0,
    }


def add_photo_grid(paragraph, doc):
    photos = []
    for label, name in PHOTO_MAP:
        path = ROOT / "photo" / name
        if path.exists():
            photos.append((POSITION_CN[label], path))
    if not photos:
        return None

    tbl = paragraph._parent.add_table(rows=2, cols=3, width=Inches(6.5))
    paragraph._p.addnext(tbl._tbl)
    tbl.style = "Table Grid"
    for idx, (label, path) in enumerate(photos[:6]):
        cell = tbl.rows[idx // 3].cells[idx % 3]
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run()
        run.add_picture(str(path), width=Inches(1.9))
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = cell.add_paragraph(label)
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    after_p = OxmlElement("w:p")
    tbl._tbl.addnext(after_p)
    return Paragraph(after_p, paragraph._parent)


def main():
    means = read_dicts(MEANS_CSV)
    params = read_dicts(PARAMS_CSV)
    err_fig, norm_fig = make_figures(means)
    summary = summarize(means)

    work_dir = ROOT / "_doc_work" / "accel_6pos"
    if work_dir.exists():
        shutil.rmtree(work_dir, ignore_errors=True)
    work_dir.mkdir(parents=True, exist_ok=True)

    try:
        tmp_in = work_dir / "report.docx"
        tmp_out = work_dir / "report_out.docx"
        shutil.copy2(REPORT_IN, tmp_in)
        doc = Document(str(tmp_in))

        anchor = find_para_startswith(doc, "4.2")
        anchor = insert_after(
            anchor,
            "为提高重力加速度标定精度，我们新增执行“六位置加速度计标定”实验。板子本体固定为已焊接状态，按 ±X、±Y、±Z 方向逐个放置，保持每个姿态约 1500 个采样点，并采用脚本自动过滤静态段，输出每向量方向均值与模长。"
        )

        def _scale_value(row):
            for key in ("scale", "scale_g_per_g"):
                if key in row and row[key] != "":
                    return float(row[key])
            raise KeyError("scale field not found")

        param_map = {r["axis"]: (float(r["bias_g"]), _scale_value(r)) for r in params}
        px_bias, px_scale = param_map["x"]
        py_bias, py_scale = param_map["y"]
        pz_bias, pz_scale = param_map["z"]
        anchor = insert_after(
            anchor,
            f"标定参数如下：X 轴 bias={px_bias:.6f} g, scale={px_scale:.6f}；"
            f"Y 轴 bias={py_bias:.6f} g, scale={py_scale:.6f}；Z 轴 bias={pz_bias:.6f} g, scale={pz_scale:.6f}。"
            "应用一维线性补偿公式：(raw - bias) / scale。",
            "Normal"
        )

        anchor = insert_after(
            anchor,
            f"六位置结果：校准前平均模长误差为 {summary['raw_mean']:.3f} mg，校准后降至 {summary['cal_mean']:.3f} mg，"
            f"误差提升比约为 {summary['improve']:.2f} 倍。单点最大误差由 {summary['raw_max']:.3f} mg 降至 {summary['cal_max']:.3f} mg。"
            "说明线性六位置标定对本次 IMU 的重力模长准确性提升显著。",
            "Normal"
        )

        table_rows = []
        for r in means:
            table_rows.append({
                "姿态": POSITION_CN[r["position"]],
                "采样数": r["samples"],
                "原始模长(g)": f"{float(r['raw_norm_g']):.6f}",
                "原始误差(mg)": f"{float(r['raw_error_mg']):.3f}",
                "校准后模长(g)": f"{float(r['cal_norm_g']):.6f}",
                "校准后误差(mg)": f"{float(r['cal_error_mg']):.3f}",
            })

        p = insert_after(anchor, "")
        _, anchor = add_table_after(
            p,
            table_rows,
            ["姿态", "采样数", "原始模长(g)", "原始误差(mg)", "校准后模长(g)", "校准后误差(mg)"]
        )

        p = insert_after(anchor, "")
        p.add_run().add_picture(str(err_fig), width=Inches(5.6))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = insert_after(p, "")
        add_caption(cap, "图：六位置标定前后加速度模长误差对比")

        p = insert_after(cap, "")
        p.add_run().add_picture(str(norm_fig), width=Inches(5.6))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap2 = insert_after(p, "")
        add_caption(cap2, "图：六位置标定前后加速度模长（与1g比较）")

        p = add_photo_grid(cap2, doc)
        if p is not None:
            insert_after(p, "图：六个标定姿态实拍（按 x、y、z 轴正负方向）", "Normal")

        exp_anchor = find_para_startswith(doc, "6.3")
        insert_after(
            exp_anchor,
            "六位置标定已完成。实验数据表明，标定后各姿态加速度模长误差均显著下降，平均由 12.935 mg 降至 3.691 mg，"
            "最大误差由 24.704 mg 降至 5.272 mg，说明该补偿方法可明显提高重力相关测量准确性。可用于后续姿态计算与零速补偿。"
        )

        doc.save(str(tmp_out))
        shutil.copy2(tmp_out, REPORT_OUT)
    finally:
        shutil.rmtree(work_dir, ignore_errors=True)

    print("raw_mean_mg", summary["raw_mean"])
    print("cal_mean_mg", summary["cal_mean"])
    print("raw_max_mg", summary["raw_max"])
    print("cal_max_mg", summary["cal_max"])
    print("improve", summary["improve"])
    print("figures", err_fig, norm_fig)
    print("wrote", REPORT_OUT)


if __name__ == "__main__":
    main()
