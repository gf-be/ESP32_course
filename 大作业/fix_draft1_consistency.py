# -*- coding: utf-8 -*-
"""
Make draft wording consistent while keeping placeholder values for unfinished
GPS, Madgwick, and simplified ESKF experiments.
"""

from pathlib import Path
import shutil

from docx import Document


ROOT = Path(__file__).resolve().parent
REPORT_IN = ROOT / "多传感器融合扩展板课程论文初稿1.docx"
REPORT_OUT = ROOT / "多传感器融合扩展板课程论文初稿1_口径统一占位保留.docx"


REPLACEMENTS = [
    (
        "并实现互补滤波、Madgwick/Mahony 滤波以及简化误差状态卡尔曼滤波。",
        "并已实现互补滤波与 Mahony PI 姿态融合；Madgwick 滤波和简化误差状态卡尔曼滤波作为后续补充实验保留占位。"
    ),
    (
        "实测结果显示，加速度计六位置标定后平均模长误差由 10.401 mg 降至 0.021 mg，提升约 497.17 倍；1 h 静态 Allan 方差记录的平均采样率为 200.000 Hz，gz 轴 ARW 为 0.005880 deg/sqrt(s)，BI 最小 Allan 偏差为 0.001145 deg/s；互补滤波在水平静止段 roll 标准差为 0.196 deg、pitch 标准差为 0.019 deg，并能在晃动后回到近水平稳定状态；GNSS 扩展测试中 NMEA 数据有效率为 82.11%。",
        "目前已完成的实测结果包括：加速度计六位置标定、磁力计椭球标定、陀螺仪 Allan 方差分析、互补滤波与 Mahony PI 姿态融合、1D-CNN IMU 去噪、BMP280 气压计测试、系统频率测试和 Web 实时姿态显示。文中涉及 GPS/GNSS、Madgwick、简化 ESKF 等结果的旧数据暂作为占位保留，后续完成对应实验后再统一替换为最终实测值。"
    ),
    (
        "结果表明，该平台已具备传感器采集、标定补偿、姿态融合和扩展数据记录能力，但磁力计、气压计、功耗、成本和 AI 去噪的完整终稿指标仍需结合最终硬件实测继续补齐。",
        "结果表明，该平台已具备传感器采集、标定补偿、姿态融合、AI 去噪、BMP280 气压测试和实时可视化能力；GPS 户外轨迹、Madgwick、简化 ESKF、功耗和成本等内容将在后续实验中继续补齐。"
    ),
    (
        "本文设计互补滤波、Madgwick/Mahony 滤波和简化 ESKF，实现对 roll、pitch、yaw 姿态角的实时估计，并从精度、稳定性、计算量和实时性等方面进行对比。",
        "本文已设计并验证互补滤波与 Mahony PI，实现对 roll、pitch、yaw 姿态角的实时估计；Madgwick 滤波和简化 ESKF 暂作为后续实验占位，完成后再加入同一评价体系进行对比。"
    ),
    (
        "板载 MPU6050、HMC5883L 和BMP280 气压传感器，并预留 GPS UART 接口；",
        "板载 MPU6050/MPU6500 兼容六轴 IMU、HMC5883L 和 BMP280 气压传感器，并预留 GPS UART 接口；"
    ),
    (
        "集成 MPU6050 六轴惯性测量单元、HMC5883L 三轴磁力计以及BMP280 气压传感器",
        "集成 MPU6050/MPU6500 兼容六轴惯性测量单元、HMC5883L 三轴磁力计以及 BMP280 气压传感器"
    ),
    (
        "集成 MPU6050 六轴惯性测量单元、HMC5883L 三轴磁力计以及 BMP280 气压传感器",
        "集成 MPU6050/MPU6500 兼容六轴惯性测量单元、HMC5883L 三轴磁力计以及 BMP280 气压传感器"
    ),
    (
        "气压计高度测试，当前仓库尚缺最终实测记录，因此在表中标为待补测。",
        "GPS 户外轨迹、功耗和单板成本当前尚缺最终实测记录；磁力计椭球标定、BMP280 气压/高度变化测试、AI 去噪和实时姿态显示已完成阶段性实测。"
    ),
    (
        "磁力计完整标定和气压计高度测试",
        "GPS 户外轨迹、功耗和成本"
    ),
    (
        "GNSS 扩展测试中 NMEA 数据有效率为 82.11%",
        "GNSS/GPS 户外轨迹实验结果暂以占位数据保留，待后续实测替换"
    ),
    (
        "Madgwick/Mahony 滤波",
        "Mahony PI 滤波（Madgwick 待补）"
    ),
    (
        "简化 ESKF",
        "简化 ESKF（待补）"
    ),
    (
        "完整终稿指标仍需结合最终硬件实测继续补齐",
        "GPS、Madgwick、简化 ESKF、功耗和成本等终稿指标仍需结合后续实测继续补齐"
    ),
]


def replace_in_paragraph(paragraph):
    text = paragraph.text
    new_text = text
    for old, new in REPLACEMENTS:
        new_text = new_text.replace(old, new)
    if new_text == text:
        return 0

    # Preserve paragraph style; rebuild runs conservatively.
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)
    return 1


def main():
    work_dir = ROOT / "_doc_work" / "draft1_consistency"
    if work_dir.exists():
        shutil.rmtree(work_dir, ignore_errors=True)
    work_dir.mkdir(parents=True, exist_ok=True)

    tmp_in = work_dir / "input.docx"
    tmp_out = work_dir / "output.docx"
    shutil.copy2(REPORT_IN, tmp_in)
    doc = Document(str(tmp_in))

    changed = 0
    for paragraph in doc.paragraphs:
        changed += replace_in_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    changed += replace_in_paragraph(paragraph)

    doc.save(str(tmp_out))
    shutil.copy2(tmp_out, REPORT_OUT)
    shutil.rmtree(work_dir, ignore_errors=True)

    print("wrote", REPORT_OUT)
    print("changed_blocks", changed)


if __name__ == "__main__":
    main()
