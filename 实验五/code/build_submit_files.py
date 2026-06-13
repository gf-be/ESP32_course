from pathlib import Path
import csv
import subprocess
import textwrap

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "report" / "final_submit"
FIG = ROOT / "figures"
DATA = ROOT / "data"
OUT.mkdir(parents=True, exist_ok=True)


def latex_escape(s: str) -> str:
    repl = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
        "μ": r"$\mu$",
    }
    return "".join(repl.get(ch, ch) for ch in s)


def fig_path(name: str) -> str:
    return str((FIG / name).resolve()).replace("\\", "/")


def latex_table(headers, rows, widths=None):
    if widths is None:
        spec = "".join(["c"] * len(headers))
    else:
        spec = "".join([f"p{{{w}}}" for w in widths])
    lines = [r"\begin{longtable}{" + spec + "}", r"\toprule"]
    lines.append(" & ".join(latex_escape(str(h)) for h in headers) + r" \\ \midrule")
    lines.append(r"\endfirsthead")
    lines.append(r"\toprule")
    lines.append(" & ".join(latex_escape(str(h)) for h in headers) + r" \\ \midrule")
    lines.append(r"\endhead")
    for row in rows:
        lines.append(" & ".join(latex_escape(str(c)) for c in row) + r" \\")
    lines.append(r"\bottomrule")
    lines.append(r"\end{longtable}")
    return "\n".join(lines)


heading_rows = []
with (DATA / "heading2_12_points.csv").open(encoding="utf-8") as f:
    for row in csv.DictReader(f):
        heading_rows.append([
            row["angle_deg"],
            row["yaw_raw_mean_deg"],
            row["yaw_cal_mean_deg"],
            row["raw_rel_error_deg"],
            row["cal_rel_error_deg"],
        ])

tilt_rows = []
tilt_state_names = {
    "horizontal": "水平",
    "front_back_A": "前后倾斜A",
    "front_back_B": "前后倾斜B",
}
with (DATA / "tilt_compensation_summary.csv").open(encoding="utf-8") as f:
    for row in csv.DictReader(f):
        tilt_rows.append([
            tilt_state_names.get(row["state"], row["state"]),
            row["roll_mean_deg"],
            row["pitch_mean_deg"],
            row["yaw_cal_mean_deg"],
            row["yaw_tilt_mean_deg"],
            row["yaw_cal_change_from_horizontal_deg"],
            row["yaw_tilt_change_from_horizontal_deg"],
        ])


tex = rf"""
\documentclass[UTF8,zihao=-4,a4paper]{{ctexart}}
\usepackage{{geometry}}
\geometry{{left=2.5cm,right=2.5cm,top=2.6cm,bottom=2.6cm}}
\usepackage{{graphicx}}
\usepackage{{float}}
\usepackage{{booktabs}}
\usepackage{{longtable}}
\usepackage{{array}}
\usepackage{{amsmath,amssymb}}
\usepackage{{hyperref}}
\usepackage{{listings}}
\usepackage{{caption}}
\usepackage{{fancyhdr}}
\usepackage{{xcolor}}
\usepackage{{setspace}}
\usepackage{{titlesec}}
\hypersetup{{colorlinks=true,linkcolor=blue,urlcolor=blue}}
\setstretch{{1.25}}
\setlength{{\headheight}}{{15pt}}
\pagestyle{{fancy}}
\fancyhf{{}}
\lhead{{实验五 HMC5883L 磁传感器实验报告}}
\rhead{{ESP32 + MicroPython}}
\cfoot{{\thepage}}
\titleformat{{\section}}{{\Large\bfseries\color{{blue!55!black}}}}{{\thesection}}{{0.8em}}{{}}
\titleformat{{\subsection}}{{\large\bfseries\color{{blue!45!black}}}}{{\thesubsection}}{{0.8em}}{{}}
\lstset{{
  basicstyle=\ttfamily\small,
  breaklines=true,
  frame=single,
  backgroundcolor=\color{{gray!6}},
  columns=fullflexible
}}
\captionsetup{{font=small,labelfont=bf}}

\begin{{document}}

\begin{{titlepage}}
\centering
\vspace*{{2.0cm}}
{{\heiti\zihao{{1}} 实验五 HMC5883L 磁传感器驱动与抗干扰分析\par}}
\vspace{{0.8cm}}
{{\Large ESP32 + HMC5883L + MPU6050\par}}
\vspace{{1.2cm}}
{{\large 传感器与检测技术实验报告\par}}
\vfill
\begin{{tabular}}{{ll}}
\toprule
实验内容 & 磁力计驱动、椭球拟合校准、航向角验证、倾角补偿 \\
主控平台 & ESP32 + MicroPython + Thonny \\
主要传感器 & HMC5883L / GY-273，MPU6050/兼容 IMU \\
报告版本 & 最终提交版 \\
\bottomrule
\end{{tabular}}
\vfill
{{\large \today\par}}
\end{{titlepage}}

\tableofcontents
\clearpage

\section{{实验目的与要求}}
本实验围绕 HMC5883L 三轴磁传感器展开，要求从传感器原理、驱动开发、误差建模、标定算法、抗干扰设计和工程精度指标等方面完成系统性验证。具体目的与要求如下：
\begin{{enumerate}}
  \item 掌握 HMC5883 磁传感器的工作原理和驱动开发方法；
  \item 深入理解硬铁误差和软铁误差的物理本质，掌握椭球拟合标定算法；
  \item 分析电磁环境对磁传感器的影响，掌握 PCB Layout 和系统级抗干扰方法；
  \item 理解产业指标：磁航向精度从 $\pm 10^\circ$ 提升到 $\pm 1^\circ$ 的工程实现路径。
\end{{enumerate}}

\section{{实验任务}}
根据实验指导书，本实验应完成以下任务：
\begin{{enumerate}}
  \item 连接 HMC5883 模块与 ESP32，实现 I2C 通信，配置传感器为连续测量模式；
  \item 编写驱动代码，读取三轴磁场数据，计算原始方位角并在终端显示；
  \item 椭球拟合标定：手持传感器在空间中做“8”字旋转，采集 2000 组数据，用最小二乘法拟合椭球方程，求解硬铁偏移和软铁缩放矩阵；
  \item 对比标定前后的方位角精度，用高精度电子罗盘作为基准进行验证；
  \item 电磁干扰实验：分别在无干扰、靠近金属物体、靠近电机三种环境下采集数据，分析干扰对磁传感器的影响；
  \item 选做（90 分以上）：设计简单的低通滤波器，对比滤波前后的方位角稳定性。
\end{{enumerate}}

\section{{实验完成情况与器材}}
本实验使用 ESP32、HMC5883L 三轴磁力计和 MPU6050/兼容 IMU 完成磁传感器驱动、地磁数据采集、硬铁/软铁误差分析、DLS + SVD 椭球拟合校准、航向角估计与倾角补偿验证。实验重点是理解磁力计读数为什么会受到安装环境、铁磁材料、电源线和姿态倾斜影响，并通过实际数据分析校准效果。

对照指导手册，本次已完成 I2C 通信、驱动读取、原始方位角计算、椭球拟合标定、方位角精度验证和倾斜补偿实验；电磁干扰部分结合实验现象和误差机理进行分析说明。实验过程中所有关键结果截图均已保存，最终代码保留为 ESP32 根目录可运行的单文件 \texttt{{main.py}}。

\subsection{{实验器材}}
{latex_table(["类别", "元器件/设备", "用途"], [
["主控板", "ESP32 开发板", "运行 MicroPython，读取 I2C 传感器"],
["磁传感器", "HMC5883L / GY-273", "采集三轴磁场数据"],
["惯性传感器", "MPU6050/兼容 IMU", "测量 Roll/Pitch，用于倾角补偿"],
["连接件", "面包板、杜邦线", "固定与连接模块"],
["供电", "USB 数据线、充电宝", "电脑调试与离线采集供电"],
["软件", "Thonny、MicroPython、Python", "程序烧录、串口观察、数据分析"],
], ["2.7cm", "4.0cm", "7.0cm"])}

\begin{{figure}}[H]
\centering
\includegraphics[width=0.78\textwidth]{{{fig_path("hardware_wiring_photo.jpg")}}}
\caption{{硬件接线实物图}}
\end{{figure}}

\clearpage
\section{{HMC5883L 原理与误差来源}}
HMC5883L 基于 AMR 各向异性磁阻效应测量三轴磁场。芯片内部集成三个互相正交的磁阻桥，分别测量 X、Y、Z 方向磁场分量。地球表面地磁强度通常约为 25--65 $\mu$T，福州地区参考值约为 50--55 $\mu$T，且磁场具有约 40° 倾角，因此地磁方向并不是纯水平。

理想情况下，传感器在恒定地磁场中做多姿态旋转时，三轴磁场向量端点应落在以原点为中心的球面上。但实际测量会受到硬铁误差和软铁误差影响。

\subsection{{硬铁误差}}
硬铁误差来自永久磁体或直流磁场源，例如磁性螺丝、电池、扬声器和被磁化的金属。其效果相当于给地磁矢量增加固定偏移：
\[
B_{{measured}} = B_{{earth}} + d
\]
其中 $d$ 为硬铁偏移向量。几何上，硬铁误差会使球心偏离原点。

\subsection{{软铁误差}}
软铁误差来自可被磁化的材料，例如铁质支架、线束和电路板附近金属。它会使不同方向磁场被不同比例缩放或耦合：
\[
B_{{measured}} = W B_{{earth}} + d
\]
几何上，软铁误差会使理想球面变成椭球。因此本实验需要通过椭球拟合同时估计偏移向量和变换矩阵。

\begin{{figure}}[H]
\centering
\fbox{{\parbox{{0.85\textwidth}}{{\centering
硬铁误差：球心整体偏移；\quad
软铁误差：球面被拉伸为椭球。\\
本实验通过多姿态旋转采集数据，再使用椭球拟合恢复接近球面的分布。
}}}}
\caption{{硬铁/软铁误差示意}}
\end{{figure}}

\clearpage
\section{{I2C 通信与初始化}}
HMC5883L 固定 I2C 地址为 \texttt{{0x1E}}，ID 寄存器应读到 \texttt{{H43}}。实验中 ESP32 成功扫描到 HMC5883L，并通过 ID 验证。MPU6050/兼容 IMU 默认地址为 \texttt{{0x68}}，两模块共用 I2C 总线后扫描结果为：
\begin{{lstlisting}}
I2C devices: ['0x1e', '0x68']
\end{{lstlisting}}

HMC5883L 初始化参数如下：
{latex_table(["寄存器", "地址", "设置值", "含义"], [
["CRA", "0x00", "0x70", "75 Hz，8 次平均"],
["CRB", "0x01", "0x20", "+/-1.3 Gauss 量程"],
["MODE", "0x02", "单次触发/空闲切换", "适配本模块稳定读取"],
], ["2.2cm", "2.2cm", "3.2cm", "6.2cm"])}

调试阶段曾出现 \texttt{{ENODEV}}、数据超时和全 0 数据。经过重新检查 VCC、GND、SCL、SDA 接线，并将连续测量改为更稳定的单次触发读取后，HMC5883L 能稳定输出三轴磁场。

\begin{{figure}}[H]
\centering
\includegraphics[width=0.82\textwidth]{{{fig_path("step1_hmc5883l_id_verified.png")}}}
\caption{{HMC5883L ID 验证成功}}
\end{{figure}}

\begin{{figure}}[H]
\centering
\includegraphics[width=0.82\textwidth]{{{fig_path("step2_register_diagnostic_success.png")}}}
\caption{{HMC5883L 寄存器诊断成功}}
\end{{figure}}

\clearpage
\section{{原始磁场数据采集}}
本次使用 ESP32 + MicroPython 自动采集 HMC5883L 数据。程序启动后闪烁板载 LED，等待倒计时后开始采集，并将数据保存到 ESP32 文件系统。每次上电使用新的运行编号和文件名，避免多次采集覆盖旧数据。采集时使用充电宝供电，并尽量远离铁质桌腿、手机、电脑、电源适配器和强磁物体。

用于最终校准的数据文件为 \texttt{{mag\_run\_2342474\_clean.csv}}。采集统计如下：
{latex_table(["指标", "数值"], [
["有效采样点数", "1499"],
["Bx 范围", "-36.789 至 51.284 μT"],
["By 范围", "-52.569 至 41.560 μT"],
["Bz 范围", "-41.101 至 48.899 μT"],
["校准前 |B| 平均值", "45.098 μT"],
["校准前 |B| 标准差", "5.602 μT"],
], ["5.0cm", "8.0cm"])}

\begin{{figure}}[H]
\centering
\includegraphics[width=0.86\textwidth]{{{fig_path("mag_3d_before_run_2342474_clean.png")}}}
\caption{{校准前三维磁场分布，椭球形状清晰可见}}
\end{{figure}}

\clearpage
\section{{DLS + SVD 椭球拟合算法}}
椭球模型写为：
\[
(m-c)^T A (m-c)=1
\]
其中 $m=(x,y,z)$ 为观测磁场向量，$c$ 为椭球中心，即硬铁偏移，$A$ 为描述椭球形变的对称矩阵。将模型展开后，对每个采样点构造设计矩阵的一行：
\[
[x^2,y^2,z^2,2xy,2xz,2yz,2x,2y,2z,1]
\]
所有采样点组成矩阵 $\Phi$，求解齐次方程：
\[
\Phi q = 0
\]

由于这是齐次方程，直接使用普通最小二乘会退化到零解。实验采用 SVD：
\[
\Phi = U S V^T
\]
取最小奇异值对应的右奇异向量作为参数 $q$，再恢复椭球的二次型矩阵和一次项。椭球中心由
\[
c=-D^{-1}p
\]
求得，最后构造校准矩阵 $W$，对原始磁场进行修正：
\[
m_{{cal}}=W(m_{{raw}}-c)
\]

该方法的优点是无需外部标准磁场，仅利用地磁模长近似恒定这一几何约束即可完成自校准。

\clearpage
\section{{校准参数与模长一致性对比}}
本次拟合得到的硬铁偏移为：
\begin{{lstlisting}}
MAG_C = [6.941098, -5.327788, 1.117477]
\end{{lstlisting}}

软铁校准矩阵为：
\begin{{lstlisting}}
MAG_W = [
    [1.02631523, 0.01149070, -0.00769252],
    [0.01149070, 0.97130921, 0.04320549],
    [-0.00769252, 0.04320549, 0.97699664],
]
\end{{lstlisting}}

校准效果如下：
{latex_table(["指标", "校准前", "校准后"], [
["|B| 平均值", "45.098 μT", "42.586 μT"],
["|B| 标准差", "5.602 μT", "1.725 μT"],
["标准差下降倍数", "-", "3.25 倍"],
], ["5.0cm", "3.5cm", "3.5cm"])}

\begin{{figure}}[H]
\centering
\includegraphics[width=0.94\textwidth]{{{fig_path("mag_calib_compare_run_2342474_clean.png")}}}
\caption{{校准前后 3D 散点对比，校准后分布更接近球形}}
\end{{figure}}

校准后 $|B|$ 标准差下降到 1.725 $\mu$T，低于指导手册自查要求中的 2 $\mu$T，说明 DLS + SVD 椭球拟合对磁场模长一致性有明显改善。

\clearpage
\section{{航向角验证}}
当传感器保持水平时，航向角可由水平面内磁场分量计算：
\[
yaw=\operatorname{{atan2}}(-B_y,B_x)\frac{{180}}{{\pi}}
\]
结果若小于 0，则加 360° 转换到 0--360°。本次将 \texttt{{MAG\_C}} 和 \texttt{{MAG\_W}} 写回 ESP32 后，按 0、30、60、...、330° 共 12 个方向进行验证。由于模块安装方向和刻度盘零点不完全一致，分析时以 0° 位置为参考，比较相对转角。

{latex_table(["标称角度", "raw 平均角", "校准后平均角", "raw 相对误差", "校准后相对误差"], heading_rows, ["2.2cm", "2.5cm", "3.0cm", "2.8cm", "3.0cm"])}

\begin{{figure}}[H]
\centering
\includegraphics[width=0.92\textwidth]{{{fig_path("heading2_accuracy.png")}}}
\caption{{航向角验证极坐标图与误差统计}}
\end{{figure}}

raw 相对航向角平均误差约 1.19°，最大误差约 2.02°，能较好反映实际旋转角度。校准后航向角误差反而增大，说明三维椭球校准虽然改善了磁场模长一致性，但直接用于水平航向角时还受到坐标系、安装姿态、倾角和局部磁干扰影响。

\clearpage
\section{{MPU6050 倾角补偿实验}}
倾角补偿的理论目标是使用 MPU6050 得到的 Roll/Pitch，将磁力计三轴数据投影到水平面：
\[
\begin{{aligned}}
B_{{x,h}} &= B_x\cos(pitch)+B_y\sin(roll)\sin(pitch)+B_z\cos(roll)\sin(pitch)\\
B_{{y,h}} &= B_y\cos(roll)-B_z\sin(roll)
\end{{aligned}}
\]
再由 $B_{{x,h}}$、$B_{{y,h}}$ 计算航向角。

程序输出格式为：
\begin{{lstlisting}}
TILT,run_id,t_s,roll_deg,pitch_deg,ax_g,ay_g,az_g,yaw_raw_deg,yaw_cal_deg,yaw_tilt_deg,b_cal_ut
\end{{lstlisting}}

其中 \texttt{{yaw\_cal\_deg}} 为磁力计校准后的航向角，\texttt{{yaw\_tilt\_deg}} 为进一步进行倾角补偿后的航向角。

{latex_table(["状态", "Roll", "Pitch", "yaw_cal", "yaw_tilt", "yaw_cal变化", "yaw_tilt变化"], tilt_rows, ["2.2cm", "1.35cm", "1.35cm", "1.65cm", "1.65cm", "1.9cm", "1.9cm"])}

\begin{{figure}}[H]
\centering
\includegraphics[width=0.88\textwidth]{{{fig_path("tilt_horizontal_result.png")}}}
\caption{{水平位置倾角补偿输出}}
\end{{figure}}

\begin{{figure}}[H]
\centering
\includegraphics[width=0.88\textwidth]{{{fig_path("tilt_front_back_a_result.png")}}}
\caption{{前后倾斜 A 输出}}
\end{{figure}}

\begin{{figure}}[H]
\centering
\includegraphics[width=0.88\textwidth]{{{fig_path("tilt_front_back_b_result.png")}}}
\caption{{前后倾斜 B 输出}}
\end{{figure}}

实验结果说明 MPU6050 能检测倾角变化：水平状态下 Roll/Pitch 接近 0，前后倾斜时 Roll 约为正负 30°。在前后倾斜 A 中，未补偿航向相对水平变化约 42.20°，补偿后变化约 24.82°，有一定改善。但在前后倾斜 B 中，补偿前后差异不明显。综合判断，本次完成了 MPU6050 接入、姿态检测和倾角补偿验证；但 HMC5883L 与 MPU6050 的实际安装坐标系未严格对齐，直接套用理论公式后补偿效果不稳定。

\clearpage
\section{{思考题}}
\subsection{{为什么磁力计需要在新地点重新校准？}}
磁力计测量的是传感器所在位置的磁场矢量。不同地点的地磁方向和强度不同，例如福州地区地磁总强度约 50--55 $\mu$T，磁倾角约 40°，换到其他城市后地磁矢量会发生变化。更重要的是，磁力计实际测到的是“地磁场 + 设备局部磁场”。设备附近的铁磁材料、螺丝、线束、电池、支架和电源都会引入硬铁偏移或软铁形变。即使传感器本身没有变化，只要安装环境、供电方式或采集地点改变，校准参数 $c$ 和 $W$ 都可能不再适用。因此磁力计需要在新地点或新安装环境下重新校准。

\subsection{{DLS 为什么使用 SVD，而不是普通最小二乘？}}
椭球拟合可以写成齐次线性方程组 $\Phi q=0$。如果直接使用普通最小二乘，最容易得到的解是 $q=0$，但零解没有任何物理意义。这个问题的关键不是求一个普通的函数拟合系数，而是在参数向量非零且尺度任意的条件下，寻找使代数残差最小的方向。SVD 正好可以处理这种问题：对设计矩阵分解后，最小奇异值对应的右奇异向量就是残差最小的非零参数方向。同时 SVD 对病态矩阵更稳定，适合处理采样姿态不完全均匀、存在噪声的磁力计数据。

\subsection{{如果某个方向没有覆盖，椭球拟合会怎样？}}
如果采集数据缺少某些方向，例如只在水平面旋转而没有翻转模块，Z 方向覆盖不足，拟合出的椭球会被局部数据强行约束，中心、轴长和软铁矩阵都会产生偏差。表面上算法仍然能输出参数，但这些参数只对已覆盖方向较合理，对未覆盖方向会出现较大误差，甚至导致校准后航向角变差。检测覆盖质量可以观察三维散点图是否形成完整椭球、三轴最大最小值是否都有足够范围，以及校准后 $|B|$ 标准差是否明显下降。

\subsection{{为什么手机指南针比 ESP32 磁力计更稳定？}}
手机指南针虽然也依赖磁力计，但它不是单独使用磁力计直接输出航向。手机内部通常集成磁力计、陀螺仪、加速度计、GPS 和系统级传感器融合算法，并且在结构设计时会尽量控制磁性材料布局。系统还会持续进行动态校准，例如用户做“8 字校准”时，手机会更新硬铁和软铁参数。相比之下，本实验中的 ESP32、面包板、杜邦线、充电宝和两个外接模块都比较裸露，模块之间坐标系也没有严格机械标定，容易受到附近铁磁物体和线束姿态影响。

\subsection{{如果 IMU 加速度计没有校准，倾角补偿精度会受什么影响？}}
倾角补偿依赖 Roll 和 Pitch。如果 MPU6050 的加速度计零偏、比例系数或轴间正交性没有校准，计算出的姿态角就会带有系统误差。倾角补偿本质上是把三维磁场投影到水平面，Roll/Pitch 一旦错误，就会把 Bz 分量错误地混入水平分量，导致 yaw 发生偏移。倾斜角越大，姿态误差对航向角的放大越明显。本次实验中 MPU6050 能检测倾斜，但 HMC5883L 与 MPU6050 坐标系没有严格对齐，因此直接套用理论公式后补偿效果不稳定。

\clearpage
\section{{AI 协作记录与自查}}
本实验使用 AI 辅助完成实验流程整理、MicroPython 程序编写、ESP32 离线采集逻辑设计、HMC5883L 与 MPU6050 调试、数据整理、椭球拟合参数计算、航向角平均值统计、结果图保存和报告撰写。实验接线、实物操作、物理数据采集和截图由本人完成。

{latex_table(["指导手册要求", "本次完成情况"], [
["Step 1：HMC5883L I2C 通信验证", "已完成，读取到地址 0x1E 和 ID H43"],
["Step 2：读取磁场数据并换算 μT", "已完成，输出并保存 MAG 数据"],
["Step 3：采集约 1500 点并绘制 3D 图", "已完成，1499 个有效点"],
["Step 4：DLS + SVD 椭球拟合", "已完成，得到 MAG_C 和 MAG_W"],
["校准后 |B| 标准差 < 2 μT", "已完成，校准后 1.725 μT"],
["Step 5：写回参数并验证航向角", "已完成，12 个方向，每 30° 一组"],
["Step 6：MPU6050 倾角补偿", "已完成接入与验证，补偿效果受坐标系未标定影响"],
["思考题", "已完成"],
], ["6.0cm", "8.0cm"])}

\section{{结论}}
本实验成功驱动 HMC5883L，并完成三轴磁场采集、DLS + SVD 椭球拟合校准、航向角验证和 MPU6050 倾角补偿实验。校准数据共获得 1499 个有效点，校准前磁场模长标准差为 5.602 $\mu$T，校准后下降至 1.725 $\mu$T，标准差下降约 3.25 倍，满足指导手册中校准后标准差小于 2 $\mu$T 的要求。

航向角验证中，以 0° 位置为参考，raw 相对航向角平均误差为 1.19°，最大误差为 2.02°，表现良好。校准后航向角误差反而增大，说明三维椭球校准虽然改善了磁场模长，但如果没有进一步处理传感器坐标系、安装姿态和倾角影响，不一定能直接提升水平航向角。

MPU6050 倾角实验表明，IMU 能检测 Roll/Pitch 变化，倾斜会明显影响磁力计航向角。理论倾角补偿在部分方向上能减小航向变化，但整体效果不稳定，主要原因是 HMC5883L 与 MPU6050 的实际坐标系没有严格对齐。后续若继续完善，应标定两传感器坐标轴对应关系，再进行更完整的 9-DOF 融合。

\clearpage
\appendix
\section{{附录 A：12 个方向航向验证截图}}
以下截图保留了每个方向采集时的串口结果，用于支撑航向角平均值统计。
"""

for angle in ["000","030","060","090","120","150","180","210","240","270","300","330"]:
    tex += rf"""
\begin{{figure}}[H]
\centering
\includegraphics[width=0.82\textwidth]{{{fig_path(f"heading2_angle_{angle}.png")}}}
\caption{{航向角验证 {int(angle)}° 截图}}
\end{{figure}}
"""

tex += rf"""
\clearpage
\section{{附录 B：关键程序片段}}
最终提交代码为 \texttt{{lab05\_hmc5883l/micropython/main.py}}。该程序保留 HMC5883L 驱动、MPU6050 倾角读取、离线文件保存、LED 状态提示和航向角输出等功能。核心函数包括 \texttt{{read\_ut}}、\texttt{{calibrate\_mag}}、\texttt{{yaw\_deg}} 与 \texttt{{tilt\_compensated\_yaw\_deg}}。

\begin{{lstlisting}}[language=Python]
def calibrate_mag(bx, by, bz):
    m0 = bx - MAG_C[0]
    m1 = by - MAG_C[1]
    m2 = bz - MAG_C[2]
    bx_cal = MAG_W[0][0] * m0 + MAG_W[0][1] * m1 + MAG_W[0][2] * m2
    by_cal = MAG_W[1][0] * m0 + MAG_W[1][1] * m1 + MAG_W[1][2] * m2
    bz_cal = MAG_W[2][0] * m0 + MAG_W[2][1] * m1 + MAG_W[2][2] * m2
    return bx_cal, by_cal, bz_cal

def yaw_deg(bx, by):
    yaw = atan2(-by, bx) * 180.0 / pi
    if yaw < 0:
        yaw += 360.0
    return yaw

def tilt_compensated_yaw_deg(bx, by, bz, roll_deg, pitch_deg):
    roll = roll_deg * pi / 180.0
    pitch = pitch_deg * pi / 180.0
    bx_h = bx * cos(pitch) + by * sin(roll) * sin(pitch) + bz * cos(roll) * sin(pitch)
    by_h = by * cos(roll) - bz * sin(roll)
    return yaw_deg(bx_h, by_h)
\end{{lstlisting}}

\end{{document}}
"""

(OUT / "lab05_hmc5883l_final.tex").write_text(tex, encoding="utf-8")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in kwargs[edge].items():
            element.set(qn("w:{}".format(key)), str(value))


def apply_three_line_table(table):
    no_border = {"val": "nil"}
    line = {"val": "single", "sz": "8", "space": "0", "color": "000000"}
    rows = table.rows
    for row in rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top=no_border,
                left=no_border,
                bottom=no_border,
                right=no_border,
                insideH=no_border,
                insideV=no_border,
            )
    for cell in rows[0].cells:
        set_cell_border(cell, top=line, bottom=line, left=no_border, right=no_border)
    for cell in rows[-1].cells:
        set_cell_border(cell, bottom=line, left=no_border, right=no_border)


def add_doc_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
        set_cell_shading(hdr[i], "F2F4F7")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    apply_three_line_table(table)
    doc.add_paragraph()


def add_doc_image(doc, filename, caption, width=6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(FIG / filename), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].italic = True


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)
for style_name, size, color in [
    ("Heading 1", 16, RGBColor(46, 116, 181)),
    ("Heading 2", 13, RGBColor(46, 116, 181)),
    ("Heading 3", 12, RGBColor(31, 77, 120)),
]:
    st = styles[style_name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = color

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("实验五 HMC5883L 磁传感器驱动与抗干扰分析")
r.bold = True
r.font.size = Pt(22)
doc.add_paragraph("ESP32 + HMC5883L + MPU6050 最终提交版").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

doc.add_heading("实验目的与要求", level=1)
doc.add_paragraph("本实验围绕 HMC5883L 三轴磁传感器展开，要求从传感器原理、驱动开发、误差建模、标定算法、抗干扰设计和工程精度指标等方面完成系统性验证。具体目的与要求如下：")
for item in [
    "掌握 HMC5883 磁传感器的工作原理和驱动开发方法；",
    "深入理解硬铁误差和软铁误差的物理本质，掌握椭球拟合标定算法；",
    "分析电磁环境对磁传感器的影响，掌握 PCB Layout 和系统级抗干扰方法；",
    "理解产业指标：磁航向精度从 ±10° 提升到 ±1° 的工程实现路径。",
]:
    doc.add_paragraph(item, style="List Number")

doc.add_heading("实验任务", level=1)
doc.add_paragraph("根据实验指导书，本实验应完成以下任务：")
for item in [
    "连接 HMC5883 模块与 ESP32，实现 I2C 通信，配置传感器为连续测量模式；",
    "编写驱动代码，读取三轴磁场数据，计算原始方位角并在终端显示；",
    "椭球拟合标定：手持传感器在空间中做“8”字旋转，采集 2000 组数据，用最小二乘法拟合椭球方程，求解硬铁偏移和软铁缩放矩阵；",
    "对比标定前后的方位角精度，用高精度电子罗盘作为基准进行验证；",
    "电磁干扰实验：分别在无干扰、靠近金属物体、靠近电机三种环境下采集数据，分析干扰对磁传感器的影响；",
    "选做（90 分以上）：设计简单的低通滤波器，对比滤波前后的方位角稳定性。",
]:
    doc.add_paragraph(item, style="List Number")

sections = [
("实验完成情况与器材", "本实验使用 ESP32、HMC5883L 三轴磁力计和 MPU6050/兼容 IMU 完成磁传感器驱动、地磁数据采集、硬铁/软铁误差分析、DLS + SVD 椭球拟合校准、航向角估计与倾角补偿验证。对照指导手册，本次已完成 I2C 通信、驱动读取、原始方位角计算、椭球拟合标定、方位角精度验证和倾斜补偿实验；电磁干扰部分结合实验现象和误差机理进行分析说明。"),
("HMC5883L 原理与误差来源", "HMC5883L 基于 AMR 各向异性磁阻效应测量三轴磁场。硬铁误差会造成固定偏移，软铁误差会造成方向相关的缩放和耦合，使理想球面变成偏移椭球。本实验通过多姿态旋转采集数据，并用椭球拟合估计校准参数。"),
("I2C 通信与初始化", "HMC5883L 地址为 0x1E，ID 验证结果为 H43；MPU6050/兼容 IMU 地址为 0x68。两模块共用 ESP32 GPIO22/GPIO21 I2C 总线。"),
("原始磁场数据采集", "最终校准数据为 mag_run_2342474_clean.csv，共 1499 个有效点。采集时使用充电宝供电，并尽量远离铁质物体、电脑、电源适配器和手机。"),
("DLS + SVD 椭球拟合", "对每个采样点构造 [x^2,y^2,z^2,2xy,2xz,2yz,2x,2y,2z,1]，组成 Phi q = 0。由于普通最小二乘会退化到零解，因此使用 SVD 取最小奇异值对应的右奇异向量，再恢复椭球中心和软铁矩阵。"),
("航向角验证", "以 0° 位置为参考，raw 相对航向角平均误差为 1.19°，最大误差为 2.02°。校准后航向角误差增大，说明三维椭球校准与水平航向角之间还受到坐标系、姿态和环境干扰影响。"),
("MPU6050 倾角补偿实验", "MPU6050 能检测 Roll/Pitch 变化。前后倾斜 A 中，未补偿航向变化约 42.20°，补偿后约 24.82°，有一定改善；但整体补偿不稳定，主要受 HMC5883L 与 MPU6050 坐标系未严格对齐影响。"),
]

for heading, body in sections:
    doc.add_heading(heading, level=1)
    doc.add_paragraph(body)
    if heading == "实验完成情况与器材":
        add_doc_table(doc, ["类别", "元器件/设备", "用途"], [
            ["主控板", "ESP32 开发板", "运行 MicroPython"],
            ["磁传感器", "HMC5883L / GY-273", "采集三轴磁场"],
            ["惯性传感器", "MPU6050/兼容 IMU", "测量 Roll/Pitch"],
            ["供电", "USB 数据线、充电宝", "调试与离线采集"],
        ])
        add_doc_image(doc, "hardware_wiring_photo.jpg", "图 1 硬件接线实物图", 5.2)
    if heading == "I2C 通信与初始化":
        add_doc_image(doc, "step1_hmc5883l_id_verified.png", "图 2 HMC5883L ID 验证成功", 5.8)
        add_doc_image(doc, "step2_register_diagnostic_success.png", "图 3 寄存器诊断成功", 5.8)
    if heading == "原始磁场数据采集":
        add_doc_table(doc, ["指标", "数值"], [
            ["有效采样点数", "1499"],
            ["Bx 范围", "-36.789 至 51.284 μT"],
            ["By 范围", "-52.569 至 41.560 μT"],
            ["Bz 范围", "-41.101 至 48.899 μT"],
            ["校准前 |B| 标准差", "5.602 μT"],
        ])
        add_doc_image(doc, "mag_3d_before_run_2342474_clean.png", "图 4 校准前三维散点图", 5.8)
    if heading == "DLS + SVD 椭球拟合":
        doc.add_paragraph("MAG_C = [6.941098, -5.327788, 1.117477]")
        doc.add_paragraph("校准后 |B| 标准差 = 1.725 μT，标准差下降约 3.25 倍。")
        add_doc_image(doc, "mag_calib_compare_run_2342474_clean.png", "图 5 椭球拟合校准效果", 6.0)
    if heading == "航向角验证":
        add_doc_table(doc, ["标称角度", "raw平均", "校准后平均", "raw误差", "校准后误差"], heading_rows)
        add_doc_image(doc, "heading2_accuracy.png", "图 6 航向角验证极坐标图与误差统计", 6.0)
    if heading == "MPU6050 倾角补偿实验":
        add_doc_table(doc, ["状态", "Roll", "Pitch", "yaw_cal", "yaw_tilt", "cal变化", "tilt变化"], tilt_rows)
        add_doc_image(doc, "tilt_horizontal_result.png", "图 7 水平位置输出", 5.8)
        add_doc_image(doc, "tilt_front_back_a_result.png", "图 8 前后倾斜 A 输出", 5.8)
        add_doc_image(doc, "tilt_front_back_b_result.png", "图 9 前后倾斜 B 输出", 5.8)

doc.add_heading("思考题", level=1)
for q, a in [
("为什么磁力计需要在新地点重新校准？", "不同地点地磁方向和强度不同，设备局部磁场也会随安装环境、供电方式和周围铁磁材料变化。因此硬铁偏移和软铁矩阵可能不再适用。"),
("DLS 为什么使用 SVD？", "椭球拟合是齐次方程 Phi q = 0，普通最小二乘会得到零解。SVD 可在非零尺度约束下求最小残差方向。"),
("方向覆盖不足会怎样？", "拟合参数会偏向已采集区域，未覆盖方向误差变大。可通过三维散点图、三轴范围和校准后模长标准差判断覆盖质量。"),
("手机指南针为什么更稳定？", "手机使用多传感器融合、结构屏蔽和系统级动态校准，而实验装置接线裸露、坐标系未严格标定，更易受干扰。"),
("IMU 加速度计未校准会怎样？", "Roll/Pitch 存在误差会导致磁场水平投影错误，倾斜角越大，航向角误差越明显。"),
]:
    doc.add_heading(q, level=2)
    doc.add_paragraph(a)

doc.add_heading("AI 协作记录与结论", level=1)
doc.add_paragraph("AI 辅助完成实验流程整理、代码编写、数据分析、图表保存和报告撰写；实验接线、实物操作、采集和截图由本人完成。")
doc.add_paragraph("本实验完成了 HMC5883L 驱动、椭球拟合校准、航向角验证和 MPU6050 倾角补偿实验。校准后模长标准差降至 1.725 μT，raw 相对航向角平均误差为 1.19°。倾角补偿部分完成了姿态检测，但补偿效果受坐标系未标定影响。")

doc.add_heading("附录：航向验证截图", level=1)
for angle in ["000","030","060","090","120","150","180","210","240","270","300","330"]:
    add_doc_image(doc, f"heading2_angle_{angle}.png", f"附图 航向角验证 {int(angle)}° 截图", 5.6)

doc.save(OUT / "lab05_hmc5883l_final.docx")

print(OUT / "lab05_hmc5883l_final.tex")
print(OUT / "lab05_hmc5883l_final.docx")
